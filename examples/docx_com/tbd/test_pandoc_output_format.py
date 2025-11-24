"""Test to see what pandoc-crossref actually outputs."""

import base64
import subprocess

import pytest
from docx import Document

from paradoc.io.word.com_api import is_word_com_available


@pytest.mark.skipif(not is_word_com_available(), reason="COM automation only if Word COM is available")
def test_pandoc_crossref_output(tmp_path):
    """Check what pandoc-crossref actually outputs for cross-references."""

    test_dir = tmp_path / "pandoc_test"
    test_dir.mkdir()

    # Create test image
    png_data = base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mP8/5+hHgAHggJ/PchI7wAAAABJRU5ErkJggg=="
    )

    images_dir = test_dir / "images"
    images_dir.mkdir()

    img_path = images_dir / "fig1.png"
    img_path.write_bytes(png_data)

    # Simple markdown with cross-reference
    md_content = """# Section 1

This is a reference to [@fig:test].

![Test figure](images/fig1.png){#fig:test}
"""

    md_file = test_dir / "test.md"
    md_file.write_text(md_content, encoding="utf-8")

    # Create minimal metadata file
    metadata_content = """---
chapters: true
chaptersDepth: 1
numberSections: true
sectionsDepth: 2
figPrefix: "Figure"
tblPrefix: "Table"
figureTitle: "Figure"
tableTitle: "Table"
---
"""
    metadata_file = test_dir / "metadata.yaml"
    metadata_file.write_text(metadata_content, encoding="utf-8")

    output_file = test_dir / "output.docx"

    # Run pandoc with pandoc-crossref
    cmd = [
        "pandoc",
        str(md_file),
        "-o",
        str(output_file),
        "--filter",
        "pandoc-crossref",
        f"--metadata-file={metadata_file}",
        f"--resource-path={test_dir}",
    ]

    print("\n" + "=" * 80)
    print("RUNNING PANDOC")
    print("=" * 80)
    print(f"Command: {' '.join(cmd)}")

    result = subprocess.run(cmd, capture_output=True, text=True)

    if result.returncode != 0:
        print(f"STDERR: {result.stderr}")
        pytest.fail(f"Pandoc failed with return code {result.returncode}")

    print("✓ Pandoc completed successfully")

    # Read the document and check what text was generated
    doc = Document(str(output_file))

    print("\n" + "=" * 80)
    print("DOCUMENT PARAGRAPHS")
    print("=" * 80)

    for i, para in enumerate(doc.paragraphs):
        text = para.text
        if text.strip():
            print(f"\nPara {i}: {text}")

            # Check for figure references
            if "fig" in text.lower() and "reference" in text.lower():
                print("  ^^^ THIS IS THE REFERENCE PARAGRAPH")
                print(f"  Raw text: '{text}'")

                # Check all runs
                print("  Runs in this paragraph:")
                for j, run in enumerate(para.runs):
                    print(f"    Run {j}: '{run.text}'")

    print("\n" + "=" * 80)
    print("ANALYSIS")
    print("=" * 80)

    # Find the reference paragraph
    ref_para = None
    for para in doc.paragraphs:
        if "reference to" in para.text.lower():
            ref_para = para
            break

    assert ref_para is not None, "Could not find reference paragraph"

    ref_text = ref_para.text
    print(f"\nReference text: '{ref_text}'")

    # Check what format pandoc-crossref used
    if "Figure 1" in ref_text or "Figure1" in ref_text:
        print("✓ Found 'Figure' with number")

        # Check exact format
        if " 1" in ref_text:
            print("  Format: 'Figure 1' (with space)")
        elif "Figure1" in ref_text:
            print("  Format: 'Figure1' (NO space)")

        if "-1" in ref_text or " 1-1" in ref_text:
            print("  Has chapter numbering: '1-1'")
