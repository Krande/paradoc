import logging
import os
import pathlib
import shutil
from dataclasses import dataclass

import pypandoc
from docx import Document
from docxcompose.composer import Composer

from .utils import (
    apply_custom_styles_to_docx,
    close_word_docs_by_name,
    docx_update,
    fix_headers_after_compose,
    get_list_of_files,
)

MY_DOCX_TMPL = pathlib.Path(__file__).resolve().absolute().parent / "resources" / "template.docx"
MY_DOCX_TMPL_BLANK = pathlib.Path(__file__).resolve().absolute().parent / "resources" / "template_blank.docx"


@dataclass
class MdFile:
    path: pathlib.Path
    is_appendix: bool
    new_file: pathlib.Path
    build_file: pathlib.Path


@dataclass
class Formatting:
    is_appendix: bool
    paragraph_style_map: dict
    table_format: str


class OneDoc:
    """

    Work dir should be structured with 2 directories '00-main' and '01-app' representing the Main and Appendix part
    of your report.


    :param source_dir:
    :param export_format:
    :param clean_build_dir:
    :param create_dirs: If not exists create default main and app dirs.
    :param variable_map:
    :param table_format: Optional override table format
    :param paragraph_style_map: Optional override paragraph format
    :param appendix_heading_map: Optional override appendix formats
    """

    default_app_map = {
        "Heading 1": "Appendix",
        "Heading 2": "Appendix X.1",
        "Heading 3": "Appendix X.1.1",
        "Heading 4": "Appendix X.1.1.1",
    }

    def __init__(
        self,
        source_dir=None,
        export_format="docx",
        main_prefix="00-main",
        app_prefix="01-app",
        clean_build_dir=True,
        create_dirs=True,
        **kwargs,
    ):
        self.source_dir = pathlib.Path().resolve().absolute() if source_dir is None else pathlib.Path(source_dir)
        self.work_dir = kwargs.get("work_dir", pathlib.Path("").resolve().absolute())
        self._main_prefix = main_prefix
        self._app_prefix = app_prefix
        self.export_format = export_format
        self.variables = dict()
        self.tables = dict()
        self.equations = dict()

        # Style info: https://python-docx.readthedocs.io/en/latest/user/styles-using.html
        self.table_format = kwargs.get("table_format", "Grid Table 1 Light")
        self.paragraph_style_map = kwargs.get(
            "paragraph_style_map",
            {
                # "Normal": "Normal Indent",
                "First Paragraph": "Normal Indent",
                "Body Text": "Normal Indent",
                "Compact": "Normal Indent",
            },
        )

        self.appendix_heading_map = kwargs.get(
            "appendix_heading_map",
            OneDoc.default_app_map,
        )

        self.md_files_main = []
        self.md_files_app = []

        report_dir = self.source_dir / main_prefix
        os.makedirs(report_dir, exist_ok=True)

        for md_file in get_list_of_files(report_dir, ".md"):
            is_appendix = True if app_prefix in md_file else False
            md_file = pathlib.Path(md_file)
            new_file = self.build_dir / md_file.relative_to(report_dir).with_suffix(".docx")
            build_file = self.build_dir / md_file.relative_to(report_dir)

            md_file = MdFile(
                path=md_file,
                is_appendix=is_appendix,
                new_file=new_file,
                build_file=build_file,
            )
            if is_appendix:
                self.md_files_app.append(md_file)
            else:
                self.md_files_main.append(md_file)

        if create_dirs is True:
            os.makedirs(self.main_dir, exist_ok=True)
            os.makedirs(self.app_dir, exist_ok=True)

        if clean_build_dir is True:
            shutil.rmtree(self.build_dir, ignore_errors=True)

    def compile(self, output_name, auto_open=False, metadata_file=None):
        """

        :param output_name: Name of output document
        :param auto_open:
        :param metadata_file:
        """
        from .utils import variable_sub

        dest_file = (self.dist_dir / output_name).with_suffix(f".{self.export_format}").resolve().absolute()

        logging.debug(f'Compiling report to "{dest_file}"')
        os.makedirs(self.build_dir, exist_ok=True)
        os.makedirs(self.dist_dir, exist_ok=True)

        for mdf in self.md_files_main + self.md_files_app:
            md_file = mdf.path
            os.makedirs(mdf.new_file.parent, exist_ok=True)

            # Substitute parameters/tables in the creation of the document
            with open(md_file, "r") as f:
                tmp_md_doc = f.read()
                tmp_md_doc = variable_sub(tmp_md_doc, self.tables)
                tmp_md_doc = variable_sub(tmp_md_doc, self.variables)
                tmp_md_doc = variable_sub(tmp_md_doc, self.equations)

            with open(mdf.build_file, "w") as f:
                f.write(tmp_md_doc)

            metadata_file = self.source_dir / "metadata.yaml" if metadata_file is None else metadata_file
            if metadata_file.exists() is False:
                with open(metadata_file, "w") as f:
                    f.write('linkReferences: true\nnameInLink: true\nfigPrefix: "Figure"\ntblPrefix: "Table"')

            pypandoc.convert_file(
                str(mdf.build_file),
                self.export_format,
                outputfile=str(mdf.new_file),
                format="markdown",
                extra_args=[
                    "-M2GB",
                    "+RTS",
                    "-K64m",
                    "-RTS",
                    f"--resource-path={md_file.parent}",
                    f"--metadata-file={metadata_file}"
                    # f"--reference-doc={MY_DOCX_TMPL}",
                ],
                filters=["pandoc-crossref"],
                encoding="utf8",
            )

        # Main Document - Format Style
        composer_main = Composer(Document(MY_DOCX_TMPL))
        composer_main.doc.add_page_break()

        for i, md in enumerate(self.md_files_main):
            doc_in = Document(str(md.new_file))
            doc_in.add_page_break()
            composer_main.append(doc_in)
            logging.info(f"Added {md.new_file}")

        main_format = Formatting(False, self.paragraph_style_map, self.table_format)
        _ = apply_custom_styles_to_docx(composer_main.doc, main_format)
        composer_main.doc.add_page_break()

        # Appendix - Format Style
        composer_app = Composer(Document(MY_DOCX_TMPL_BLANK))
        for i, md in enumerate(self.md_files_app):
            doc_in = Document(str(md.new_file))
            doc_in.add_page_break()
            composer_app.append(doc_in)
            logging.info(f"Added {md.new_file}")

        app_paragraph_style = dict()
        app_paragraph_style.update(self.appendix_heading_map)
        app_paragraph_style.update(self.paragraph_style_map)

        app_format = Formatting(True, app_paragraph_style, self.table_format)
        _ = apply_custom_styles_to_docx(composer_app.doc, app_format)

        composer_main.append(composer_app.doc)

        fix_headers_after_compose(composer_main.doc)

        print("Close Existing Word documents")
        close_word_docs_by_name([output_name, f"{output_name}.docx"])

        print(f'Saving Composed Document to "{dest_file}"')
        composer_main.save(dest_file)

        docx_update(str(dest_file))

        if auto_open is True:
            os.startfile(dest_file)

    @property
    def main_dir(self):
        return self.source_dir / self._main_prefix

    @property
    def app_dir(self):
        return self.source_dir / self._app_prefix

    @property
    def build_dir(self):
        return self.work_dir / "temp" / "_build"

    @property
    def dist_dir(self):
        return self.work_dir / "temp" / "_dist"
