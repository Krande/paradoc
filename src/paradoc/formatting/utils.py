import logging

from docx import Document
from docx.shared import Pt
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph
from typing import Dict
from .concepts import Formatting, TableFormat
from paradoc.concepts import Table
from typing import Union


def add_indented_normal(doc):
    from docx.enum.style import WD_STYLE_TYPE
    from docx.shared import Mm, Pt

    styles = doc.styles
    style = styles.add_style("Normal indent", WD_STYLE_TYPE.PARAGRAPH)
    style.base_style = styles["Normal"]

    paragraph_format = style.paragraph_format
    paragraph_format.left_indent = Mm(0.25)
    paragraph_format.space_before = Pt(12)
    paragraph_format.widow_control = True

    return style


def format_paragraph(pg, document, paragraph_style_map: dict):
    from docx.shared import Mm

    style_name = pg.style.name
    logging.debug(style_name)
    if style_name == "Compact":  # Is a bullet point list
        new_style_name = paragraph_style_map[pg.style.name]
        new_style = document.styles[new_style_name]
        pg.style = new_style
        pg.paragraph_format.left_indent = Mm(25)

    elif style_name in paragraph_style_map.keys():
        new_style_name = paragraph_style_map[pg.style.name]

        if new_style_name not in document.styles:
            styles = "".join([x.name + "\n" for x in document.styles])
            raise ValueError(
                f'The requested style "{new_style_name}" does not exist in style_doc.\n'
                "Note! Style names are CAPS sensitive.\n"
                f"Available styles are:\n{styles}"
            )

        new_style = document.styles[new_style_name]
        pg.style = new_style

        logging.debug(f'Changed paragraph style "{pg.style}" to "{new_style_name}"')
    else:
        if style_name not in document.styles:
            logging.info(f'StyleDoc missing style "{style_name}"')


def apply_custom_styles_to_docx(doc, doc_format: Formatting = None, style_doc=None):
    from paradoc import MY_DOCX_TMPL
    from paradoc.io.word.utils import iter_block_items

    document = style_doc if style_doc is not None else Document(MY_DOCX_TMPL)
    prev_table = False
    refs = dict()

    for block in iter_block_items(doc):
        if type(block) == Paragraph:
            if prev_table and len(block.runs) > 0:
                block.runs[0].text = "\n" + block.runs[0].text
                prev_table = False
                block.paragraph_format.space_before = None
            if block.style.name in ("Image Caption", "Table Caption"):
                ref_ = format_captions(block, doc_format)
                refs.update(ref_)
            else:
                format_paragraph(block, document, doc_format.paragraph_style_map)

        elif type(block) == DocxTable:
            if doc_format.table_format:
                format_table(block, document, doc_format.table_format)
            prev_table = True

    return refs


def get_table_ref(docx_table: DocxTable, tables: Dict[str, Table]) -> Union[Table, None]:
    cell0 = docx_table.rows[1].cells[0].paragraphs[0]
    cell0_str = cell0.text
    for key, tbl in tables.items():
        if key == cell0_str:
            df = tbl.df
            col_name = df.columns[0]
            cell0.text = str(df.iloc[0, df.columns.get_loc(col_name)])
            return tbl
    return None


def format_table(tbl: DocxTable, document, tbl_format: TableFormat):
    new_tbl_style = document.styles[tbl_format.style]
    tbl.style = new_tbl_style
    logging.info(f'Changed Table style from "{tbl.style}" to "{new_tbl_style}"')
    # tbl.paragraph_format.space_after = Pt(12)
    for i, row in enumerate(tbl.rows):
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    # run.style = document.styles["Normal"]
                    font.name = tbl_format.font_style
                    font.size = Pt(tbl_format.font_size)
                    if i == 0:
                        font.bold = True
                    else:
                        font.bold = False


def fix_headers_after_compose(doc: Document):
    from paradoc import OneDoc
    from paradoc.io.word.utils import delete_paragraph, iter_block_items

    pg_rem = []
    for pg in iter_block_items(doc):
        if type(pg) == Paragraph:
            if pg.style.name in ("Image Caption", "Table Caption"):
                continue
            else:
                if pg.style.name in list(OneDoc.default_app_map.values())[1:]:
                    pg.insert_paragraph_before(pg.text, style=pg.style.name)
                    pg_rem.append(pg)

    for pg in pg_rem:
        delete_paragraph(pg)


def format_captions(pg, doc_format: Formatting):
    from paradoc.references import insert_caption_into_runs

    ref_dict = dict()
    style_name = pg.style.name
    logging.debug(style_name)
    tmp_split = pg.text.split(":")
    prefix = tmp_split[0].strip()
    if style_name == "Image Caption":
        ref_dict[prefix] = insert_caption_into_runs(pg, "Figure", doc_format)
    elif style_name == "Table Caption":
        ref_dict[prefix] = insert_caption_into_runs(pg, "Table", doc_format)
    else:
        raise ValueError("Not possible")

    return ref_dict
