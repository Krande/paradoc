from dataclasses import dataclass

from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from paradoc.common import Figure, Table
from paradoc.config import create_logger

from .references import add_seq_reference

logger = create_logger()


@dataclass
class DocXTableRef:
    table_ref: Table = None
    docx_table: DocxTable = None
    docx_caption: Paragraph = None
    docx_following_pg: Paragraph = None
    is_appendix = False
    document_index: int = None
    actual_bookmark_name: str = None  # Store the actual Word-style bookmark name

    def is_complete(self):
        docx_attr = [self.docx_caption, self.docx_table, self.docx_following_pg]
        return all([x is not None for x in docx_attr])

    def get_content_cell0_pg(self) -> Paragraph:
        tbl = self.docx_table
        return tbl.rows[1].cells[0].paragraphs[0]

    def format_table(self, is_appendix, restart_caption_numbering=False):
        tbl = self.docx_table
        tbl_format = self.table_ref.format

        # Format content of table
        tbl.style = tbl_format.style
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

        logger.info(f'Changed Table style from "{tbl.style}" to "{tbl_format.style}"')
        for i, row in enumerate(tbl.rows):
            for cell in row.cells:
                # https://python-docx.readthedocs.io/en/latest/api/enum/WdCellVerticalAlignment.html
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    # assert isinstance(paragraph, Paragraph)
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                    for run in paragraph.runs:
                        font = run.font
                        font.name = tbl_format.font_style
                        font.size = Pt(tbl_format.font_size)
                        if i == 0:
                            font.bold = True
                        else:
                            font.bold = False
        tbl.autofit = True

        # Format table Caption
        self.docx_caption.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        rebuild_caption(self.docx_caption, "Table", self.table_ref.caption, is_appendix, restart_caption_numbering)

        # Add bookmark around the caption number for Word cross-references
        if self.table_ref.link_name_override or self.table_ref.name:
            from .references import add_bookmark_around_seq_field

            table_id = self.table_ref.link_name_override or self.table_ref.name
            bookmark_name = f"tbl:{table_id}"
            # Capture the actual bookmark name that was created
            self.actual_bookmark_name = add_bookmark_around_seq_field(self.docx_caption, bookmark_name)

        for run in self.docx_caption.runs:
            run.font.name = tbl_format.font_style
            # run.font.size = tbl_format.font_size

        # Fix formatting before Table
        self.docx_caption.paragraph_format.space_before = Pt(18)

        # Fix formatting after Table
        self.docx_following_pg.paragraph_format.space_before = Pt(22)

    def substitute_back_temp_var(self):
        pg_0 = self.get_content_cell0_pg()

        df = self.table_ref.df
        col_name = df.columns[0]
        res = df.iloc[0, df.columns.get_loc(col_name)]
        fmt = self.table_ref.format.float_fmt

        use_decimals = True
        if len(self.docx_table.rows) > 1:
            try:
                row2 = self.docx_table.rows[2].cells[0].paragraphs[0]
            except IndexError as e:
                logger.error(f"Second row not used by table. Using first error: '{e}'")
                row2 = self.docx_table.rows[1].cells[0].paragraphs[0]

            if "." not in row2.text:
                use_decimals = False
        if fmt is None or use_decimals is False:
            pg_0.text = f"{res}"
        else:
            pg_0.text = f"{res:{fmt}}"


@dataclass
class DocXFigureRef:
    figure_ref: Figure = None
    docx_figure: Paragraph = None
    docx_caption: Paragraph = None
    docx_following_pg: Paragraph = None
    is_appendix = False
    document_index: int = None
    actual_bookmark_name: str = None  # Store the actual Word-style bookmark name

    def format_figure(self, is_appendix, restart_caption_numbering):
        figure_format = self.figure_ref.format
        self.docx_caption.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Rebuild caption without bookmark
        rebuild_caption(self.docx_caption, "Figure", self.figure_ref.caption, is_appendix, restart_caption_numbering)

        # Add bookmark around the caption number (SEQ field) for Word cross-references
        if self.figure_ref.reference:
            from .references import add_bookmark_around_seq_field

            bookmark_name = f"fig:{self.figure_ref.reference}"
            # Capture the actual bookmark name that was created
            self.actual_bookmark_name = add_bookmark_around_seq_field(self.docx_caption, bookmark_name)

        for run in self.docx_caption.runs:
            run.font.name = figure_format.font_style


def rebuild_caption(caption: Paragraph, caption_prefix, caption_str, is_appendix, should_restart=False):
    caption.clear()
    caption.runs.clear()

    run = caption.add_run()

    heading_ref = '"Appendix"' if is_appendix is True else '"Heading 1"'

    # Build the SEQ field instruction
    # Format: SEQ Figure \* ARABIC \s 1
    # where Figure/Table is the identifier, \* ARABIC is the format, \s 1 restarts numbering
    seq_instruction = f"SEQ {caption_prefix} \\* ARABIC"

    if should_restart:
        # \r switch restarts numbering at heading level
        if is_appendix:
            seq_instruction += ' \\r "Appendix X.1"'
        else:
            seq_instruction += ' \\r "Heading 1"'
    else:
        # \s switch continues numbering from heading level
        seq_instruction += ' \\s 1'


    seq1 = caption._element._new_r()
    seq1.text = caption_prefix + " "

    add_seq_reference(seq1, f"STYLEREF \\s {heading_ref} \\n", run._parent)
    run._element.addprevious(seq1)

    stroke = caption._element._new_r()
    new_run = Run(stroke, run._parent)
    new_run.text = "-"
    run._element.addprevious(stroke)

    seq2 = caption._element._new_r()
    add_seq_reference(seq2, seq_instruction, run._parent)
    run._element.addprevious(seq2)

    fin = caption._element._new_r()
    fin_run = Run(fin, run._parent)
    fin_run.text = ": " + caption_str
    run._element.addprevious(fin)
