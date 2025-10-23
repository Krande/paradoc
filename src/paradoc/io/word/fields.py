"""Word field operations (SEQ, REF, STYLEREF)."""

from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.text.run import Run


def add_seq_reference(run_in, seq: str, parent):
    """Add a SEQ field to a run.

    Args:
        run_in: The run element to add the field to
        seq: The SEQ field instruction (e.g., "SEQ Figure \\* ARABIC \\s 1")
        parent: The parent element

    Returns:
        The new run with the SEQ field
    """
    new_run = Run(run_in, parent)
    r = new_run._r
    fldChar = OxmlElement("w:fldChar")
    fldChar.set(qn("w:fldCharType"), "begin")
    r.append(fldChar)
    instrText = OxmlElement("w:instrText")
    instrText.text = seq
    r.append(instrText)
    fldChar = OxmlElement("w:fldChar")
    fldChar.set(qn("w:fldCharType"), "end")
    r.append(fldChar)
    return new_run


def add_table_reference(paragraph: Paragraph, seq: str = " SEQ Table \\* ARABIC \\s 1"):
    """Add a table reference SEQ field to a paragraph.

    Args:
        paragraph: The paragraph to add the field to
        seq: The SEQ field instruction

    Returns:
        The run containing the SEQ field
    """
    run = paragraph.add_run()
    r = run._r
    fldChar = OxmlElement("w:fldChar")
    fldChar.set(qn("w:fldCharType"), "begin")
    r.append(fldChar)
    instrText = OxmlElement("w:instrText")
    instrText.text = seq
    r.append(instrText)
    fldChar = OxmlElement("w:fldChar")
    fldChar.set(qn("w:fldCharType"), "end")
    r.append(fldChar)

    return run


def add_ref_field_to_paragraph(paragraph: Paragraph, bookmark_name: str, label: str = "Figure"):
    """Add a REF field to a paragraph using the add_run API.

    Args:
        paragraph: The paragraph to add the REF field to
        bookmark_name: The name of the bookmark to reference
        label: The label text to show (e.g., "Figure", "Table")
    """
    # Create field begin
    run1 = paragraph.add_run()
    r1 = run1._r
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    r1.append(fldChar1)

    # Create field instruction
    run2 = paragraph.add_run()
    r2 = run2._r
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = f" REF {bookmark_name} \\h "
    r2.append(instrText)

    # Create field separator
    run3 = paragraph.add_run()
    r3 = run3._r
    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "separate")
    r3.append(fldChar3)

    # Create field result (placeholder)
    run4 = paragraph.add_run(f"{label} 1")

    # Create field end
    run5 = paragraph.add_run()
    r5 = run5._r
    fldChar5 = OxmlElement("w:fldChar")
    fldChar5.set(qn("w:fldCharType"), "end")
    r5.append(fldChar5)


def create_ref_field_runs(p_element, bookmark_name: str, label: str = "Figure"):
    """Create REF field runs and append them to a paragraph element.

    This creates the complete REF field structure with begin, instruction,
    separator, result, and end by directly manipulating XML.

    Args:
        p_element: The paragraph XML element (w:p)
        bookmark_name: The name of the bookmark to reference
        label: The label prefix to include (e.g., "Figure", "Table", "Eq")
    """
    # Add space before the field
    r_label = OxmlElement("w:r")
    t_label = OxmlElement("w:t")
    t_label.set(qn("xml:space"), "preserve")
    t_label.text = " "
    r_label.append(t_label)
    p_element.append(r_label)

    # Run 1: Field begin
    r1 = OxmlElement("w:r")
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    r1.append(fldChar1)
    p_element.append(r1)

    # Run 2: Field instruction
    r2 = OxmlElement("w:r")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    # Use Word's cross-reference format: REF bookmark \h \* MERGEFORMAT
    # \h creates a hyperlink, \* MERGEFORMAT preserves formatting
    instrText.text = f" REF {bookmark_name} \\h "
    r2.append(instrText)
    p_element.append(r2)

    # Run 3: Field separator
    r3 = OxmlElement("w:r")
    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "separate")
    r3.append(fldChar3)
    p_element.append(r3)

    # Run 4: Field result (placeholder text that will be replaced when field is updated)
    r4 = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "1-1"  # Placeholder - will be updated by Word to show full numbering
    r4.append(t)
    p_element.append(r4)

    # Run 5: Field end
    r5 = OxmlElement("w:r")
    fldChar5 = OxmlElement("w:fldChar")
    fldChar5.set(qn("w:fldCharType"), "end")
    r5.append(fldChar5)
    p_element.append(r5)


def create_text_run(p_element, text: str):
    """Create a text run and append it to a paragraph element.

    Args:
        p_element: The paragraph XML element (w:p)
        text: The text content for the run
    """
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = text
    r.append(t)
    p_element.append(r)


def append_ref_to_paragraph(paragraph: Paragraph, ref_name: str, text: str = ""):
    """Legacy function - Add a REF field to a paragraph with optional prefix text.

    Args:
        paragraph: The paragraph to add the REF field to
        ref_name: The bookmark name to reference
        text: Optional text to add before the field
    """
    # run 1
    run = paragraph.add_run(text)
    r = run._r
    fldChar = OxmlElement("w:fldChar")
    fldChar.set(qn("w:fldCharType"), "begin")
    r.append(fldChar)
    # run 2
    run = paragraph.add_run()
    r = run._r
    instrText = OxmlElement("w:instrText")
    instrText.text = "REF " + ref_name + " \\h"
    r.append(instrText)
    # run 3
    run = paragraph.add_run()
    r = run._r
    fldChar = OxmlElement("w:fldChar")
    fldChar.set(qn("w:fldCharType"), "end")
    r.append(fldChar)

