"""Data classes for Word document references."""

import random
from dataclasses import dataclass

from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph

from paradoc.common import Figure, Table
from paradoc.config import create_logger

logger = create_logger()


@dataclass
class DocXTableRef:
    """Reference to a table in a Word document with associated metadata."""

    table_ref: Table = None
    docx_table: DocxTable = None
    docx_caption: Paragraph = None
    docx_following_pg: Paragraph = None
    is_appendix: bool = False
    document_index: int = None
    actual_bookmark_name: str = None  # Store the actual Word-style bookmark name

    def is_complete(self) -> bool:
        """Check if all required document elements are present."""
        docx_attr = [self.docx_caption, self.docx_table, self.docx_following_pg]
        return all([x is not None for x in docx_attr])

    def get_content_cell0_pg(self) -> Paragraph:
        """Get the first paragraph of the first content cell (row 1, col 0)."""
        tbl = self.docx_table
        return tbl.rows[1].cells[0].paragraphs[0]

    def format_table(self, is_appendix: bool, restart_caption_numbering: bool = False, reference_helper=None):
        """Format the table and its caption with proper styling and numbering.

        Args:
            is_appendix: Whether this table is in the appendix
            restart_caption_numbering: Whether to restart caption numbering
            reference_helper: Optional ReferenceHelper instance for managing cross-references
        """
        from .bookmarks import add_bookmark_around_seq_field
        from .captions import rebuild_caption

        tbl = self.docx_table
        tbl_format = self.table_ref.format

        # Format content of table
        tbl.style = tbl_format.style
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

        logger.info(f'Changed Table style from "{tbl.style}" to "{tbl_format.style}"')
        for i, row in enumerate(tbl.rows):
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                    for run in paragraph.runs:
                        font = run.font
                        font.name = tbl_format.font_style
                        font.size = Pt(tbl_format.font_size)
                        if i == 0:
                            font.bold = True
                        else:
                            font.bold = False
        tbl.autofit = True

        # Format table Caption
        self.docx_caption.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        rebuild_caption(self.docx_caption, "Table", self.table_ref.caption, is_appendix, restart_caption_numbering)

        # Add bookmark around the caption number for Word cross-references
        if self.table_ref.link_name_override or self.table_ref.name:
            table_id = self.table_ref.link_name_override or self.table_ref.name

            # Use ReferenceHelper if provided
            if reference_helper:
                # Register the table and get Word-style bookmark
                bookmark_name = reference_helper.register_table(table_id, self.docx_caption)
                self.actual_bookmark_name = bookmark_name
                # Apply the bookmark to the caption paragraph
                add_bookmark_around_seq_field(self.docx_caption, bookmark_name)
            else:
                # Fallback to old method
                bookmark_name = f"tbl:{table_id}"
                # Capture the actual bookmark name that was created
                self.actual_bookmark_name = add_bookmark_around_seq_field(self.docx_caption, bookmark_name)

        for run in self.docx_caption.runs:
            run.font.name = tbl_format.font_style

        # Fix formatting before Table
        self.docx_caption.paragraph_format.space_before = Pt(18)

        # Fix formatting after Table
        self.docx_following_pg.paragraph_format.space_before = Pt(22)

    def substitute_back_temp_var(self):
        """Substitute temporary variable back with actual data value."""
        pg_0 = self.get_content_cell0_pg()

        df = self.table_ref.df
        col_name = df.columns[0]
        res = df.iloc[0, df.columns.get_loc(col_name)]
        fmt = self.table_ref.format.float_fmt

        use_decimals = True
        if len(self.docx_table.rows) > 1:
            try:
                row2 = self.docx_table.rows[2].cells[0].paragraphs[0]
            except IndexError as e:
                logger.error(f"Second row not used by table. Using first error: '{e}'")
                row2 = self.docx_table.rows[1].cells[0].paragraphs[0]

            if "." not in row2.text:
                use_decimals = False
        if fmt is None or use_decimals is False:
            pg_0.text = f"{res}"
        else:
            pg_0.text = f"{res:{fmt}}"


@dataclass
class DocXFigureRef:
    """Reference to a figure in a Word document with associated metadata."""

    figure_ref: Figure = None
    docx_figure: Paragraph = None
    docx_caption: Paragraph = None
    docx_following_pg: Paragraph = None
    is_appendix: bool = False
    document_index: int = None
    actual_bookmark_name: str = None  # Store the actual Word-style bookmark name

    def format_figure(self, is_appendix: bool, restart_caption_numbering: bool, reference_helper=None):
        """Format the figure caption with proper styling and numbering.

        Args:
            is_appendix: Whether this figure is in the appendix
            restart_caption_numbering: Whether to restart caption numbering
            reference_helper: Optional ReferenceHelper instance for managing cross-references
        """
        from .bookmarks import add_bookmark_around_seq_field
        from .captions import rebuild_caption

        figure_format = self.figure_ref.format
        self.docx_caption.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Rebuild caption without bookmark
        rebuild_caption(self.docx_caption, "Figure", self.figure_ref.caption, is_appendix, restart_caption_numbering)

        # Add bookmark around the caption number (SEQ field) for Word cross-references
        if self.figure_ref.reference:
            # Use ReferenceHelper if provided
            if reference_helper:
                # Register the figure and get Word-style bookmark
                bookmark_name = reference_helper.register_figure(self.figure_ref.reference, self.docx_caption)
                self.actual_bookmark_name = bookmark_name
                # Apply the bookmark to the caption paragraph
                add_bookmark_around_seq_field(self.docx_caption, bookmark_name)
            else:
                # Fallback to old method
                bookmark_name = f"fig:{self.figure_ref.reference}"
                # Capture the actual bookmark name that was created
                self.actual_bookmark_name = add_bookmark_around_seq_field(self.docx_caption, bookmark_name)

        for run in self.docx_caption.runs:
            run.font.name = figure_format.font_style


@dataclass
class DocXEquationRef:
    """Reference to an equation in a Word document with associated metadata."""

    semantic_id: str = None  # e.g., "maxwell_equation" from "eq:maxwell_equation"
    docx_equation: Paragraph = None  # The paragraph containing the equation
    docx_caption: Paragraph = None  # The caption paragraph (not used for inline captions)
    is_appendix: bool = False
    document_index: int = None
    actual_bookmark_name: str = None  # Store the actual Word-style bookmark name

    def format_equation(self, is_appendix: bool, restart_caption_numbering: bool = False, reference_helper=None):
        """Format the equation by replacing pandoc-crossref's number with proper SEQ field caption.

        Pandoc-crossref adds a number like "(1)" to equations. This method:
        1. Finds and removes the pandoc-crossref number (e.g., "(1)", "(2)")
        2. Replaces it with proper caption "(Eq. 1-1)" with SEQ fields
        3. Wraps a bookmark around just the caption for cross-referencing

        The result is an equation like:
        E = mc^2                                                (Eq. 1-1)

        Args:
            is_appendix: Whether this equation is in the appendix
            restart_caption_numbering: Whether to restart caption numbering
            reference_helper: Optional ReferenceHelper instance for managing cross-references
        """
        import re
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        from .fields import add_seq_reference
        from .bookmarks import generate_word_bookmark_name

        if not self.docx_equation or not self.semantic_id:
            return

        # Get the equation paragraph
        eq_para = self.docx_equation
        p_element = eq_para._p

        # Find and remove pandoc-crossref's equation number (e.g., "(1)", "(2)")
        # Pandoc-crossref places the number INSIDE the oMath element, typically in m:t elements
        # We need to find and remove those text elements that contain just a number

        logger.debug(f"[Equation] Searching for pandoc number in equation {self.semantic_id}")

        # Look for the number inside oMath elements
        # The structure is typically: oMath > ... > m:r > m:t with text like "(1)"
        # Pandoc-crossref places this at the VERY END of the math element
        math_elements = p_element.findall('.//m:oMath', namespaces={'m': 'http://schemas.openxmlformats.org/officeDocument/2006/math'})

        for math_elem in math_elements:
            # Find all m:r (math runs) within this oMath
            math_runs = math_elem.findall('.//m:r', namespaces={'m': 'http://schemas.openxmlformats.org/officeDocument/2006/math'})

            logger.debug(f"[Equation] Found {len(math_runs)} math runs in oMath element")

            # Pandoc-crossref splits the number across multiple runs at the end: '(', '1', ')'
            # We need to check if the last 3 runs match this pattern
            if len(math_runs) >= 3:
                # Extract text from last 3 runs
                last_3_texts = []
                for m_run in math_runs[-3:]:
                    m_t = m_run.find('.//m:t', namespaces={'m': 'http://schemas.openxmlformats.org/officeDocument/2006/math'})
                    if m_t is not None and m_t.text:
                        last_3_texts.append(m_t.text.strip())
                    else:
                        last_3_texts.append('')

                logger.debug(f"[Equation] Last 3 math run texts: {last_3_texts}")

                # Check if pattern is: '(', digit, ')'
                if (len(last_3_texts) == 3 and
                    last_3_texts[0] == '(' and
                    last_3_texts[1].isdigit() and
                    last_3_texts[2] == ')'):

                    logger.debug(f"[Equation] Found pandoc number pattern at end: {last_3_texts} - removing last 3 runs")
                    # Remove the last 3 math runs
                    for m_run in math_runs[-3:]:
                        m_run.getparent().remove(m_run)
                else:
                    logger.debug(f"[Equation] No pandoc number pattern found at end")
            else:
                logger.debug(f"[Equation] Not enough math runs to contain pandoc number pattern")

        # Also check for numbers in regular text runs outside oMath (fallback)
        runs = p_element.findall('.//w:r', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
        pandoc_number_run = None

        for run in runs:
            # Skip runs that contain math elements (oMath) - we already handled those
            if run.find('.//m:oMath', namespaces={'m': 'http://schemas.openxmlformats.org/officeDocument/2006/math'}):
                continue

            t_elements = run.findall('.//w:t', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
            for t_elem in t_elements:
                if t_elem.text:
                    text = t_elem.text.strip()
                    # Match patterns like "(1)", " (1) ", "(2)", etc.
                    if re.match(r'^\(?\s*\d+\s*\)?$', text):
                        logger.debug(f"[Equation] Found pandoc number in regular run: '{t_elem.text}' in run")
                        pandoc_number_run = run
                        break
            if pandoc_number_run:
                break

        if not pandoc_number_run:
            logger.debug(f"[Equation] No pandoc number found in regular runs for equation {self.semantic_id}")

        # Remove the old eq: bookmark if it exists
        old_bookmark_name = f"eq:{self.semantic_id}"
        bookmark_starts = p_element.findall('.//w:bookmarkStart', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
        bookmark_ends = p_element.findall('.//w:bookmarkEnd', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})

        for bm_start in bookmark_starts:
            if bm_start.get(qn('w:name')) == old_bookmark_name:
                bm_id = bm_start.get(qn('w:id'))
                # Remove this bookmark start
                bm_start.getparent().remove(bm_start)
                # Remove matching bookmark end
                for bm_end in bookmark_ends:
                    if bm_end.get(qn('w:id')) == bm_id:
                        bm_end.getparent().remove(bm_end)
                        break
                break

        # If we found pandoc's number, we'll replace it in place
        # Otherwise, add at the end of the paragraph
        insertion_point = pandoc_number_run if pandoc_number_run else None

        # Build the caption with SEQ fields
        # Format: "(Eq. STYLEREF-SEQ)" but bookmark wraps only "Eq. STYLEREF-SEQ" (without parentheses)
        heading_ref = '"Appendix"' if is_appendix else '"Heading 1"'

        # Create runs for the caption
        # Opening parenthesis (OUTSIDE bookmark)
        open_paren_run_elem = p_element._new_r()
        t_elem = OxmlElement('w:t')
        t_elem.text = '('
        open_paren_run_elem.append(t_elem)

        # Use ReferenceHelper if provided to get Word-style bookmark
        if reference_helper:
            bookmark_name = reference_helper.register_equation(self.semantic_id, eq_para)
            self.actual_bookmark_name = bookmark_name
        else:
            bookmark_name, bookmark_id_str = generate_word_bookmark_name()
            self.actual_bookmark_name = bookmark_name
            bookmark_id_str = str(random.randint(1, 999999))

        # Bookmark start
        bookmark_start = OxmlElement('w:bookmarkStart')
        bookmark_start.set(qn('w:id'), bookmark_id_str if 'bookmark_id_str' in locals() else str(random.randint(1, 999999)))
        bookmark_start.set(qn('w:name'), bookmark_name)

        # "Eq. " text
        eq_label_run_elem = p_element._new_r()
        t_elem = OxmlElement('w:t')
        t_elem.set(qn('xml:space'), 'preserve')
        t_elem.text = 'Eq. '
        eq_label_run_elem.append(t_elem)

        # STYLEREF field for chapter number
        styleref_run_elem = p_element._new_r()
        add_seq_reference(styleref_run_elem, f"STYLEREF \\s {heading_ref} \\n", eq_para)

        # Hyphen separator
        hyphen_run_elem = p_element._new_r()
        t_elem = OxmlElement('w:t')
        t_elem.text = '-'
        hyphen_run_elem.append(t_elem)

        # SEQ field for equation number
        if restart_caption_numbering:
            seq_instruction = f"SEQ Equation \\* ARABIC \\r 1 \\s 1"
        else:
            seq_instruction = f"SEQ Equation \\* ARABIC \\s 1"

        seq_run_elem = p_element._new_r()
        add_seq_reference(seq_run_elem, seq_instruction, eq_para)

        # Bookmark end (wraps "Eq. 1-1" without parentheses)
        bookmark_end = OxmlElement('w:bookmarkEnd')
        bookmark_end.set(qn('w:id'), bookmark_start.get(qn('w:id')))

        # Closing parenthesis (OUTSIDE bookmark)
        close_paren_run_elem = p_element._new_r()
        t_elem = OxmlElement('w:t')
        t_elem.text = ')'
        close_paren_run_elem.append(t_elem)

        if insertion_point:
            # Replace pandoc's number with our caption
            # Insert all caption elements before the pandoc number run
            insertion_point.addprevious(open_paren_run_elem)
            insertion_point.addprevious(bookmark_start)
            insertion_point.addprevious(eq_label_run_elem)
            insertion_point.addprevious(styleref_run_elem)
            insertion_point.addprevious(hyphen_run_elem)
            insertion_point.addprevious(seq_run_elem)
            insertion_point.addprevious(bookmark_end)
            insertion_point.addprevious(close_paren_run_elem)
            # Remove the old pandoc number
            insertion_point.getparent().remove(insertion_point)
        else:
            # No pandoc number found, append to end
            p_element.append(open_paren_run_elem)
            p_element.append(bookmark_start)
            p_element.append(eq_label_run_elem)
            p_element.append(styleref_run_elem)
            p_element.append(hyphen_run_elem)
            p_element.append(seq_run_elem)
            p_element.append(bookmark_end)
            p_element.append(close_paren_run_elem)

