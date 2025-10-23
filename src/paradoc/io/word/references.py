import re

from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from .utils import iter_block_items


def _normalize_bookmark_name(name: str) -> str:
    """Convert a bookmark name to Word-compatible format.

    Word doesn't like colons in bookmark names for REF fields.
    Convert 'fig:test_figure' to '_Reffig_test_figure'.

    Args:
        name: The bookmark name (e.g., "fig:test_figure")

    Returns:
        Word-compatible bookmark name (e.g., "_Reffig_test_figure")
    """
    # Replace colons with underscores
    normalized = name.replace(':', '_')

    # Add _Ref prefix if not already present (Word's convention)
    if not normalized.startswith('_Ref'):
        normalized = '_Ref' + normalized

    return normalized


def resolve_references(document):
    refs = dict()
    fig_re = re.compile(
        r"(?:Figure\s(?P<number>[0-9]{0,5})\s*)",
        re.MULTILINE | re.DOTALL | re.IGNORECASE,
    )
    tbl_re = re.compile(r"(?:Table\s(?P<number>[0-9]{0,5})\s*)", re.MULTILINE | re.DOTALL | re.IGNORECASE)

    # Fix references
    for block in iter_block_items(document):
        if type(block) is Paragraph:
            if block.style.name in ("Image Caption", "Table Caption"):
                continue
            if "Figure" in block.text or "Table" in block.text:
                for m in fig_re.finditer(block.text):
                    d = m.groupdict()
                    n = d["number"]
                    figref = f"Figure {n}"
                    if figref in refs.keys():
                        fref = refs[figref]
                        pg_ref = fref[1]

                for m in tbl_re.finditer(block.text):
                    d = m.groupdict()
                    n = d["number"]
                    tblref = f"Table {n}"
                    if tblref in refs.keys():
                        tref = refs[tblref]
                        pg_ref = tref[1]
                        parent = pg_ref._p
                        # ref_id = parent.id
                        print(parent)


def add_bookmarkStart(paragraph: Paragraph, _id):
    name = "_Ref_id_num_" + _id
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:name"), name)
    start.set(qn("w:id"), str(_id))
    paragraph._p.append(start)
    return name


def append_ref_to_paragraph(paragraph: Paragraph, ref_name, text=""):
    # run 1
    run = paragraph.add_run(text)
    r = run._r
    fldChar = OxmlElement("w:fldChar")
    fldChar.set(qn("w:fldCharType"), "begin")
    r.append(fldChar)
    # run 2
    run = paragraph.add_run()
    r = run._r
    instrText = OxmlElement("w:instrText")
    instrText.text = "REF " + ref_name + " \\h"
    r.append(instrText)
    # run 3
    run = paragraph.add_run()
    r = run._r
    fldChar = OxmlElement("w:fldChar")
    fldChar.set(qn("w:fldCharType"), "end")
    r.append(fldChar)


def add_bookmark(paragraph: Paragraph, bookmark_text, bookmark_name):
    run = paragraph.add_run()
    tag = run._r  # for reference the following also works: tag =  document.element.xpath('//w:r')[-1]
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), "0")
    start.set(qn("w:name"), bookmark_name)
    tag.append(start)

    text = OxmlElement("w:r")
    text.text = bookmark_text
    tag.append(text)

    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), "0")
    end.set(qn("w:name"), bookmark_name)
    tag.append(end)


def add_bookmark_around_seq_field(paragraph: Paragraph, bookmark_name: str) -> str:
    """Add a bookmark around the caption numbering sequence.

    For proper cross-referencing, the bookmark needs to wrap around the entire
    numbering sequence (STYLEREF + hyphen + SEQ) so that REF fields show "1-1"
    instead of just "1".

    Caption structure:
    - STYLEREF field (chapter number: "1")
    - Hyphen text run: "-"
    - SEQ field (figure/table number: "1")
    - Caption text: ": caption"

    Args:
        paragraph: The caption paragraph containing the numbering fields
        bookmark_name: The name of the bookmark (e.g., "fig:test_figure")

    Returns:
        The actual bookmark name that was created (e.g., "_Ref306075071")
    """
    # Generate Word-style random bookmark name
    import random
    random_id = random.randint(100000000, 999999999)
    word_style_name = f"_Ref{random_id}"

    # Generate a unique bookmark ID
    bookmark_id = str(random.randint(1, 999999))

    p_element = paragraph._p
    runs = list(p_element.findall(qn('w:r')))

    # Find the STYLEREF field (chapter number) - this is the start of the numbering
    styleref_begin_idx = None
    seq_end_idx = None

    for idx, run in enumerate(runs):
        # Look for field characters
        fld_chars = run.findall(qn('w:fldChar'))
        for fld_char in fld_chars:
            fld_type = fld_char.get(qn('w:fldCharType'))

            if fld_type == 'begin' and styleref_begin_idx is None:
                # Check if this is the STYLEREF field
                for check_idx in range(idx, min(idx + 3, len(runs))):
                    instr_texts = runs[check_idx].findall(qn('w:instrText'))
                    for instr in instr_texts:
                        if instr.text and 'STYLEREF' in instr.text:
                            styleref_begin_idx = idx
                            break
                    if styleref_begin_idx is not None:
                        break

            elif fld_type == 'end' and styleref_begin_idx is not None:
                # Check if there's a SEQ field after this
                # Look ahead for SEQ field
                for check_idx in range(idx + 1, min(idx + 10, len(runs))):
                    check_run_fld_chars = runs[check_idx].findall(qn('w:fldChar'))
                    for check_fld_char in check_run_fld_chars:
                        check_fld_type = check_fld_char.get(qn('w:fldCharType'))
                        if check_fld_type == 'end':
                            # Check if this is a SEQ field end
                            for seq_check_idx in range(max(0, check_idx - 3), check_idx):
                                seq_instr_texts = runs[seq_check_idx].findall(qn('w:instrText'))
                                for seq_instr in seq_instr_texts:
                                    if seq_instr.text and 'SEQ' in seq_instr.text and 'STYLEREF' not in seq_instr.text:
                                        seq_end_idx = check_idx
                                        break
                                if seq_end_idx is not None:
                                    break
                            if seq_end_idx is not None:
                                break
                    if seq_end_idx is not None:
                        break
                break

        if seq_end_idx is not None:
            break

    if styleref_begin_idx is not None and seq_end_idx is not None:
        # Found the full numbering sequence - wrap bookmark around STYLEREF + hyphen + SEQ
        start = OxmlElement("w:bookmarkStart")
        start.set(qn("w:id"), bookmark_id)
        start.set(qn("w:name"), word_style_name)
        runs[styleref_begin_idx].addprevious(start)

        end = OxmlElement("w:bookmarkEnd")
        end.set(qn("w:id"), bookmark_id)
        runs[seq_end_idx].addnext(end)

        return word_style_name

    # Fallback: try to find just the SEQ field if STYLEREF approach didn't work
    for idx, run in enumerate(runs):
        fld_chars = run.findall(qn('w:fldChar'))
        has_begin = False
        has_end = False

        for fld_char in fld_chars:
            fld_type = fld_char.get(qn('w:fldCharType'))
            if fld_type == 'begin':
                has_begin = True
            elif fld_type == 'end':
                has_end = True

        if has_begin and has_end:
            instr_texts = run.findall(qn('w:instrText'))
            for instr in instr_texts:
                if instr.text and 'SEQ' in instr.text and 'STYLEREF' not in instr.text:
                    start = OxmlElement("w:bookmarkStart")
                    start.set(qn("w:id"), bookmark_id)
                    start.set(qn("w:name"), word_style_name)
                    run.addprevious(start)

                    end = OxmlElement("w:bookmarkEnd")
                    end.set(qn("w:id"), bookmark_id)
                    run.addnext(end)

                    return word_style_name

    # Final fallback to paragraph bookmark
    return add_bookmark_to_caption(paragraph, bookmark_name)


def add_bookmark_to_caption(paragraph: Paragraph, bookmark_name: str) -> str:
    """Add a bookmark to a caption paragraph for cross-referencing.

    The bookmark wraps around the entire caption paragraph so that
    Word's cross-reference feature can reference it properly.

    Args:
        paragraph: The caption paragraph to add the bookmark to
        bookmark_name: The name of the bookmark (e.g., "fig:test_figure")

    Returns:
        The actual bookmark name that was created
    """
    # Generate Word-style random bookmark name
    import random
    random_id = random.randint(100000000, 999999999)
    word_style_name = f"_Ref{random_id}"

    # Generate a unique bookmark ID
    # Important: Both bookmarkStart and bookmarkEnd must use the SAME ID
    bookmark_id = str(random.randint(1, 999999))

    # Add bookmark start at the beginning of the paragraph
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), bookmark_id)
    start.set(qn("w:name"), word_style_name)
    paragraph._p.insert(0, start)

    # Add bookmark end at the end of the paragraph
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), bookmark_id)
    paragraph._p.append(end)

    return word_style_name


def insert_caption(pg: Paragraph, prefix, run, text, is_appendix: bool):
    heading_ref = "Appendix" if is_appendix is True else '"Heading 1"'

    seq1 = pg._element._new_r()
    add_seq_reference(seq1, f"STYLEREF \\s {heading_ref} \\n", run._parent)
    run._element.addprevious(seq1)
    stroke = pg._element._new_r()
    new_run = Run(stroke, run._parent)
    new_run.text = "-"
    run._element.addprevious(stroke)
    seq2 = pg._element._new_r()
    add_seq_reference(seq2, f"SEQ {prefix} \\* ARABIC \\s 1", run._parent)
    run._element.addprevious(seq2)
    fin = pg._element._new_r()
    fin_run = Run(fin, run._parent)
    fin_run.text = ": " + text
    run._element.addprevious(fin)


def insert_caption_into_runs(pg: Paragraph, prefix: str, is_appendix: bool):
    tmp_split = pg.text.split(":")
    prefix_old = tmp_split[0].strip()
    text = tmp_split[-1].strip()
    srun = pg.runs[0]
    if len(pg.runs) > 1:
        run = pg.runs[1]
        tmp_str = pg.runs[0].text
        pg.runs[0].text = f"{prefix} "
        insert_caption(pg, prefix, run, tmp_str.split(":")[-1].strip(), is_appendix)
    else:
        srun.text = f"{prefix} "
        run = pg.add_run()
        insert_caption(pg, prefix, run, text, is_appendix)

    return srun, pg, prefix_old


def add_seq_reference(run_in, seq, parent):
    from docx.text.run import Run

    new_run = Run(run_in, parent)
    r = new_run._r
    fldChar = OxmlElement("w:fldChar")
    fldChar.set(qn("w:fldCharType"), "begin")
    r.append(fldChar)
    instrText = OxmlElement("w:instrText")
    instrText.text = seq
    r.append(instrText)
    fldChar = OxmlElement("w:fldChar")
    fldChar.set(qn("w:fldCharType"), "end")
    r.append(fldChar)
    return new_run


def add_table_reference(paragraph, seq=" SEQ Table \\* ARABIC \\s 1"):
    run = paragraph.add_run()
    r = run._r
    fldChar = OxmlElement("w:fldChar")
    fldChar.set(qn("w:fldCharType"), "begin")
    r.append(fldChar)
    instrText = OxmlElement("w:instrText")
    instrText.text = seq
    r.append(instrText)
    fldChar = OxmlElement("w:fldChar")
    fldChar.set(qn("w:fldCharType"), "end")
    r.append(fldChar)

    return run


def convert_figure_references_to_ref_fields(document, figures):
    """Convert plain text figure references to Word REF fields.

    This function finds paragraphs that contain figure references (like "Figure 1" or "fig. 1")
    and converts them to proper Word REF fields that point to the bookmarked captions.

    Args:
        document: The Word document
        figures: List of DocXFigureRef objects containing figure information
    """
    # Build a mapping from semantic reference IDs to actual bookmark names
    # and a list of bookmark names in order
    ref_id_to_bookmark = {}
    bookmarks_in_order = []

    for fig in figures:
        if hasattr(fig, 'figure_ref') and fig.figure_ref.reference:
            ref_id = fig.figure_ref.reference
            # Use the actual bookmark name that was created (e.g., "_Ref306075071")
            # instead of trying to normalize the semantic name
            if hasattr(fig, 'actual_bookmark_name') and fig.actual_bookmark_name:
                actual_bookmark = fig.actual_bookmark_name
                ref_id_to_bookmark[ref_id] = actual_bookmark
                bookmarks_in_order.append(actual_bookmark)
            else:
                # Fallback to old behavior if actual_bookmark_name not set
                bookmark_name = f"fig:{ref_id}"
                normalized_name = _normalize_bookmark_name(bookmark_name)
                ref_id_to_bookmark[ref_id] = normalized_name
                bookmarks_in_order.append(normalized_name)

    if not bookmarks_in_order:
        return  # No figures to process

    # Pattern to match figure references from pandoc-crossref
    # Match: "Figure X", "fig. X", "Figure X-Y" etc.
    # Note: SEQ fields show as "-" before evaluation, so we match any number pattern
    fig_ref_pattern = re.compile(r'\b((?:Figure|fig\.)\s+([\d\-]+))\b', re.IGNORECASE)

    # Track figure numbers to bookmark mapping
    # Extract figure numbers from bookmarks for mapping
    figure_num_to_bookmark = {}
    for idx, bookmark in enumerate(bookmarks_in_order):
        # Map figure index (1-based) to bookmark
        figure_num_to_bookmark[idx + 1] = bookmark

    # Iterate through all paragraphs
    for block in iter_block_items(document):
        if not isinstance(block, Paragraph):
            continue

        # Skip caption paragraphs
        if block.style.name in ("Image Caption", "Table Caption", "Captioned Figure"):
            continue

        # Check if paragraph contains figure references
        if not re.search(fig_ref_pattern, block.text):
            continue

        # Process the paragraph to replace text with REF fields
        original_text = block.text

        # Find all matches in the paragraph text
        matches = list(fig_ref_pattern.finditer(original_text))
        if not matches:
            continue

        # Store paragraph element before clearing
        p_element = block._p

        # Clear all runs in the paragraph by removing them from the XML
        for run in list(block.runs):  # Use list() to avoid modification during iteration
            p_element.remove(run._element)

        # Also remove hyperlink elements (pandoc-crossref creates these)
        for child in list(p_element):
            if child.tag == qn('w:hyperlink'):
                p_element.remove(child)

        # Rebuild the paragraph with text and REF fields
        # We manually create and append run elements to ensure correct order
        last_pos = 0
        for match in matches:
            # Extract the figure number from the match
            figure_num_str = match.group(2)  # The number part (e.g., "1" from "Figure 1")

            # Try to parse the figure number
            # If it's "-" (unevaluated SEQ field), assume it's the first figure
            try:
                if figure_num_str == "-":
                    figure_num = 1  # Default to first figure
                else:
                    figure_num = int(figure_num_str)
            except ValueError:
                figure_num = 1  # Fallback to first figure

            # Map figure number to bookmark (1-based indexing)
            # Figure 1 -> bookmarks_in_order[0], Figure 2 -> bookmarks_in_order[1], etc.
            bookmark_idx = figure_num - 1

            if 0 <= bookmark_idx < len(bookmarks_in_order):
                bookmark_name = bookmarks_in_order[bookmark_idx]
            else:
                # If figure number is out of range, use the last bookmark
                bookmark_name = bookmarks_in_order[-1] if bookmarks_in_order else None

            if bookmark_name is None:
                # No bookmark available, just add remaining text and break
                if last_pos < len(original_text):
                    _add_text_run(p_element, original_text[last_pos:])
                break


            # Add text before the reference
            if match.start() > last_pos:
                before_text = original_text[last_pos:match.start()]
                _add_text_run(p_element, before_text)

            # Add REF field with "Figure" label
            _add_ref_field_runs(p_element, bookmark_name, label="Figure")

            last_pos = match.end()

        # Add remaining text after the last reference
        if last_pos < len(original_text):
            after_text = original_text[last_pos:]
            _add_text_run(p_element, after_text)




def _add_text_run(p_element, text):
    """Add a text run to a paragraph element by appending to XML.

    Args:
        p_element: The paragraph XML element (w:p)
        text: The text content for the run
    """
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = text
    r.append(t)

    # Simply append to the paragraph - runs should come after pPr
    p_element.append(r)


def _add_ref_field_runs(p_element, bookmark_name, label="Figure"):
    """Add REF field runs to a paragraph element by appending to XML.

    This creates the complete REF field structure with begin, instruction, separator, result, and end.

    Args:
        p_element: The paragraph XML element (w:p)
        bookmark_name: The name of the bookmark to reference
        label: The label prefix to include (e.g., "Figure", "Table", "Eq")
    """
    # Add space before the label to ensure proper spacing
    r_space_before = OxmlElement("w:r")
    t_space_before = OxmlElement("w:t")
    t_space_before.set(qn("xml:space"), "preserve")
    t_space_before.text = " "
    r_space_before.append(t_space_before)
    p_element.append(r_space_before)

    # Add the label prefix as regular text before the field
    r_label = OxmlElement("w:r")
    t_label = OxmlElement("w:t")
    t_label.set(qn("xml:space"), "preserve")
    t_label.text = f"{label} "
    r_label.append(t_label)
    p_element.append(r_label)

    # Run 1: Field begin
    r1 = OxmlElement("w:r")
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    r1.append(fldChar1)
    p_element.append(r1)

    # Run 2: Field instruction
    r2 = OxmlElement("w:r")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    # Use Word's cross-reference format: REF bookmark \h \* MERGEFORMAT
    # \h creates a hyperlink, \* MERGEFORMAT preserves formatting
    instrText.text = f" REF {bookmark_name} \\h "
    r2.append(instrText)
    p_element.append(r2)

    # Run 3: Field separator
    r3 = OxmlElement("w:r")
    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "separate")
    r3.append(fldChar3)
    p_element.append(r3)

    # Run 4: Field result (placeholder text that will be replaced when field is updated)
    r4 = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "1-1"  # Placeholder - will be updated by Word to show full numbering
    r4.append(t)
    p_element.append(r4)

    # Run 5: Field end
    r5 = OxmlElement("w:r")
    fldChar5 = OxmlElement("w:fldChar")
    fldChar5.set(qn("w:fldCharType"), "end")
    r5.append(fldChar5)
    p_element.append(r5)

    # Add a space after the REF field
    r_space = OxmlElement("w:r")
    t_space = OxmlElement("w:t")
    t_space.set(qn("xml:space"), "preserve")
    t_space.text = " "
    r_space.append(t_space)
    p_element.append(r_space)


def add_ref_field_to_paragraph(paragraph: Paragraph, bookmark_name: str):
    """Add a REF field to a paragraph.

    Args:
        paragraph: The paragraph to add the REF field to
        bookmark_name: The name of the bookmark to reference
    """
    # Create field begin
    run1 = paragraph.add_run()
    r1 = run1._r
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    r1.append(fldChar1)

    # Create field instruction
    run2 = paragraph.add_run()
    r2 = run2._r
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = f" REF {bookmark_name} \\h "
    r2.append(instrText)

    # Create field separator
    run3 = paragraph.add_run()
    r3 = run3._r
    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "separate")
    r3.append(fldChar3)

    # Create field result (placeholder)
    run4 = paragraph.add_run("Figure 1")

    # Create field end
    run5 = paragraph.add_run()
    r5 = run5._r
    fldChar5 = OxmlElement("w:fldChar")
    fldChar5.set(qn("w:fldCharType"), "end")
    r5.append(fldChar5)


def convert_table_references_to_ref_fields(document, tables):
    """Convert plain text table references to Word REF fields.

    This function finds paragraphs that contain table references (like "Table 1" or "tbl. 1")
    and converts them to proper Word REF fields that point to the bookmarked captions.

    Args:
        document: The Word document
        tables: List of DocXTableRef objects containing table information
    """
    # Build a mapping from semantic reference IDs to actual bookmark names
    ref_id_to_bookmark = {}
    bookmarks_in_order = []

    for tbl in tables:
        if hasattr(tbl, 'table_ref') and hasattr(tbl.table_ref, 'link_name_override'):
            ref_id = tbl.table_ref.link_name_override or tbl.table_ref.name
            if hasattr(tbl, 'actual_bookmark_name') and tbl.actual_bookmark_name:
                actual_bookmark = tbl.actual_bookmark_name
                ref_id_to_bookmark[ref_id] = actual_bookmark
                bookmarks_in_order.append(actual_bookmark)

    if not bookmarks_in_order:
        return  # No tables to process

    # Pattern to match table references from pandoc-crossref
    tbl_ref_pattern = re.compile(r'\b((?:Table|tbl\.)\s+([\d\-]+))\b', re.IGNORECASE)

    # Iterate through all paragraphs
    for block in iter_block_items(document):
        if not isinstance(block, Paragraph):
            continue

        # Skip caption paragraphs
        if block.style.name in ("Image Caption", "Table Caption", "Captioned Figure"):
            continue

        # Check if paragraph contains table references
        if not re.search(tbl_ref_pattern, block.text):
            continue

        # Process the paragraph to replace text with REF fields
        original_text = block.text
        matches = list(tbl_ref_pattern.finditer(original_text))
        if not matches:
            continue

        # Store paragraph element before clearing
        p_element = block._p

        # Clear all runs and hyperlinks
        for run in list(block.runs):
            p_element.remove(run._element)
        for child in list(p_element):
            if child.tag == qn('w:hyperlink'):
                p_element.remove(child)

        # Rebuild the paragraph with text and REF fields
        last_pos = 0
        for match in matches:
            table_num_str = match.group(2)
            try:
                if table_num_str == "-":
                    table_num = 1
                else:
                    table_num = int(table_num_str)
            except ValueError:
                table_num = 1

            bookmark_idx = table_num - 1
            if 0 <= bookmark_idx < len(bookmarks_in_order):
                bookmark_name = bookmarks_in_order[bookmark_idx]
            else:
                bookmark_name = bookmarks_in_order[-1] if bookmarks_in_order else None

            if bookmark_name is None:
                if last_pos < len(original_text):
                    _add_text_run(p_element, original_text[last_pos:])
                break

            # Add text before the reference
            if match.start() > last_pos:
                before_text = original_text[last_pos:match.start()]
                _add_text_run(p_element, before_text)

            # Add REF field with "Table" label
            _add_ref_field_runs(p_element, bookmark_name, label="Table")

            last_pos = match.end()

        # Add remaining text after the last reference
        if last_pos < len(original_text):
            after_text = original_text[last_pos:]
            _add_text_run(p_element, after_text)


def convert_equation_references_to_ref_fields(document, equations):
    """Convert plain text equation references to Word REF fields.

    This function finds paragraphs that contain equation references (like "Eq 1" or "eq. 1")
    and converts them to proper Word REF fields that point to the bookmarked captions.

    Args:
        document: The Word document
        equations: List of equation objects containing equation information
    """
    # Build a mapping from semantic reference IDs to actual bookmark names
    ref_id_to_bookmark = {}
    bookmarks_in_order = []

    # Note: This assumes equations will have similar structure to figures/tables
    # You may need to adjust based on how equations are actually stored
    for eq in equations:
        if hasattr(eq, 'reference') and eq.reference:
            ref_id = eq.reference
            if hasattr(eq, 'actual_bookmark_name') and eq.actual_bookmark_name:
                actual_bookmark = eq.actual_bookmark_name
                ref_id_to_bookmark[ref_id] = actual_bookmark
                bookmarks_in_order.append(actual_bookmark)

    if not bookmarks_in_order:
        return  # No equations to process

    # Pattern to match equation references from pandoc-crossref
    # Match: "Eq X", "eq. X", "Equation X", etc.
    eq_ref_pattern = re.compile(r'\b((?:Eq(?:uation)?|eq\.)\s+([\d\-]+))\b', re.IGNORECASE)

    # Iterate through all paragraphs
    for block in iter_block_items(document):
        if not isinstance(block, Paragraph):
            continue

        # Skip caption paragraphs
        if block.style.name in ("Image Caption", "Table Caption", "Captioned Figure"):
            continue

        # Check if paragraph contains equation references
        if not re.search(eq_ref_pattern, block.text):
            continue

        # Process the paragraph to replace text with REF fields
        original_text = block.text
        matches = list(eq_ref_pattern.finditer(original_text))
        if not matches:
            continue

        # Store paragraph element before clearing
        p_element = block._p

        # Clear all runs and hyperlinks
        for run in list(block.runs):
            p_element.remove(run._element)
        for child in list(p_element):
            if child.tag == qn('w:hyperlink'):
                p_element.remove(child)

        # Rebuild the paragraph with text and REF fields
        last_pos = 0
        for match in matches:
            eq_num_str = match.group(2)
            try:
                if eq_num_str == "-":
                    eq_num = 1
                else:
                    eq_num = int(eq_num_str)
            except ValueError:
                eq_num = 1

            bookmark_idx = eq_num - 1
            if 0 <= bookmark_idx < len(bookmarks_in_order):
                bookmark_name = bookmarks_in_order[bookmark_idx]
            else:
                bookmark_name = bookmarks_in_order[-1] if bookmarks_in_order else None

            if bookmark_name is None:
                if last_pos < len(original_text):
                    _add_text_run(p_element, original_text[last_pos:])
                break

            # Add text before the reference
            if match.start() > last_pos:
                before_text = original_text[last_pos:match.start()]
                _add_text_run(p_element, before_text)

            # Add REF field with "Eq" label
            _add_ref_field_runs(p_element, bookmark_name, label="Eq")

            last_pos = match.end()

        # Add remaining text after the last reference
        if last_pos < len(original_text):
            after_text = original_text[last_pos:]
            _add_text_run(p_element, after_text)
