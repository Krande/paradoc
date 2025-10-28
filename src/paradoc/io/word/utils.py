import os
import pathlib
from typing import List

import pypandoc
from docx.document import Document as ProxyDocument
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

from paradoc.common import MY_DOCX_TMPL, MarkDownFile
from paradoc.config import logger
from paradoc.io.word.compose.composer import Composer
from paradoc.utils import get_list_of_files


def fix_bookmark_ids(document):
    """Fix bookmark ID mismatches that occur during docxcompose merging.

    docxcompose renumbers bookmarkStart IDs to avoid conflicts when merging documents,
    but it doesn't update the corresponding bookmarkEnd IDs, causing Word to treat
    them as invalid bookmarks and show "Error! Not a valid bookmark self-reference".

    This function finds all bookmarkStart elements and ensures their corresponding
    bookmarkEnd elements have matching IDs.

    Args:
        document: The Word document to fix
    """
    from docx.oxml.ns import qn

    # Get the document body element
    body = document._element.body

    # Find all bookmarkStart elements and build a mapping of name -> ID
    bookmark_name_to_id = {}

    for elem in body.iter():
        if elem.tag == qn("w:bookmarkStart"):
            bm_id = elem.get(qn("w:id"))
            bm_name = elem.get(qn("w:name"))
            if bm_id and bm_name:
                bookmark_name_to_id[bm_name] = bm_id

    # Now find all bookmarkEnd elements and update their IDs to match
    # Note: bookmarkEnd doesn't have a name attribute, so we need to track
    # the order and match them up

    # Alternative approach: Track bookmark starts and ends by name using a different strategy
    # Build a mapping by finding start/end pairs in order
    bookmark_starts = {}

    for elem in body.iter():
        if elem.tag == qn("w:bookmarkStart"):
            bm_id = elem.get(qn("w:id"))
            bm_name = elem.get(qn("w:name"))
            if bm_id and bm_name:
                bookmark_starts[bm_id] = bm_name

    # Now update bookmarkEnd elements to match their corresponding start IDs
    # We need to track which starts have been closed
    open_bookmarks = {}  # Maps old_id -> new_id for currently open bookmarks

    for paragraph in body.iter(qn("w:p")):
        for elem in paragraph:
            if elem.tag == qn("w:bookmarkStart"):
                old_id = elem.get(qn("w:id"))
                bm_name = elem.get(qn("w:name"))
                if old_id and bm_name:
                    # Track this bookmark as open
                    open_bookmarks[old_id] = (old_id, bm_name)

            elif elem.tag == qn("w:bookmarkEnd"):
                end_id = elem.get(qn("w:id"))
                if end_id and end_id in open_bookmarks:
                    # This end matches a start we've seen
                    start_id, bm_name = open_bookmarks[end_id]
                    # ID is already correct, just close it
                    del open_bookmarks[end_id]
                else:
                    # Try to find the matching start by looking for an open bookmark
                    # that hasn't been closed yet
                    if len(open_bookmarks) == 1:
                        # Only one open bookmark, this end must be for it
                        old_end_id = end_id
                        start_id, bm_name = list(open_bookmarks.values())[0]
                        # Update the end ID to match the start
                        elem.set(qn("w:id"), start_id)
                        # Remove from open bookmarks (use the old key)
                        for k, v in list(open_bookmarks.items()):
                            if v == (start_id, bm_name):
                                del open_bookmarks[k]
                                break


def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None


def iter_block_items(parent):
    """
    Yield each paragraph and table child within *parent*, in document order.
    Each returned value is an instance of either Table or Paragraph. *parent*
    would most commonly be a reference to a main Document object, but
    also works for a _Cell object, which itself can contain paragraphs and tables.
    """
    from docx.document import Document

    if isinstance(parent, Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("something's not right")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)
        else:
            logger.debug(f"Unrecognized child element type {type(child)}")


def convert_markdown_dir_to_docx(source, dest, dest_format, extra_args, style_doc=None):
    """

    :param source:
    :param dest:
    :param dest_format:
    :param extra_args:
    :param style_doc:
    :return:
    """
    from docx import Document
    from paradoc.io.word.compose.composer import Composer

    build_dir = source / "_build"
    if style_doc is not None:
        document = Document(str(style_doc))
        document.add_page_break()
        composer = Composer(document)
    else:
        composer = None
    files = []
    for md_file in get_list_of_files(source, ".md"):
        if "_build" in md_file or "_dist" in md_file:
            continue
        md_file = pathlib.Path(md_file)
        new_file = build_dir / md_file.parent.name / md_file.with_suffix(".docx").name
        os.makedirs(new_file.parent, exist_ok=True)

        output = pypandoc.convert_file(
            str(md_file),
            dest_format,
            format="markdown",
            outputfile=str(new_file),
            extra_args=extra_args,
            filters=["pandoc-crossref"],
            sandbox=False,
        )
        logger.info(output)
        files.append(str(new_file))

    # for i in range(0, len(files)):
    #     doc = Document(files[i])
    #     doc.add_page_break()
    #     if composer is None:
    #         composer = Composer(doc)
    #     else:
    #         composer.append(doc)
    #
    #     logger.info(f"Added {files[i]}")

    composer.save(str(dest))


def get_from_doc_by_index(index: int, doc: ProxyDocument):
    for i, block in enumerate(iter_block_items(doc)):
        if i == index:
            return block


def add_to_composer(source_doc, md_files: List[MarkDownFile]) -> Composer:
    from docx import Document

    composer_doc = Composer(Document(str(source_doc)))
    if source_doc == MY_DOCX_TMPL:
        composer_doc.doc.add_page_break()
    for i, md in enumerate(md_files):
        doc_in = Document(str(md.new_file))
        doc_in.add_page_break()
        composer_doc.append(doc_in)
        logger.info(f"Added {md.new_file}")
    return composer_doc
