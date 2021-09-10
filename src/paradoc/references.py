import logging

from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def resolve_references(document):
    import re

    from docx.text.paragraph import Paragraph

    from .utils import iter_block_items

    refs = dict()
    fig_re = re.compile(
        r"(?:Figure\s(?P<number>[0-9]{0,5})\s*)",
        re.MULTILINE | re.DOTALL | re.IGNORECASE,
    )
    tbl_re = re.compile(r"(?:Table\s(?P<number>[0-9]{0,5})\s*)", re.MULTILINE | re.DOTALL | re.IGNORECASE)

    # Fix references
    for block in iter_block_items(document):
        if type(block) == Paragraph:
            if block.style.name in ("Image Caption", "Table Caption"):
                continue
            if "Figure" in block.text or "Table" in block.text:
                for m in fig_re.finditer(block.text):
                    d = m.groupdict()
                    n = d["number"]
                    figref = f"Figure {n}"
                    if figref in refs.keys():
                        fref = refs[figref]
                        pg_ref = fref[1]

                for m in tbl_re.finditer(block.text):
                    d = m.groupdict()
                    n = d["number"]
                    tblref = f"Table {n}"
                    if tblref in refs.keys():
                        tref = refs[tblref]
                        pg_ref = tref[1]
                        parent = pg_ref._p
                        # ref_id = parent.id
                        print(parent)


def add_bookmarkStart(paragraph, _id):
    name = "_Ref_id_num_" + _id
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:name"), name)
    start.set(qn("w:id"), str(_id))
    paragraph._p.append(start)
    return name


def append_ref_to_paragraph(paragraph, refName, text=""):
    # run 1
    run = paragraph.add_run(text)
    r = run._r
    fldChar = OxmlElement("w:fldChar")
    fldChar.set(qn("w:fldCharType"), "begin")
    r.append(fldChar)
    # run 2
    run = paragraph.add_run()
    r = run._r
    instrText = OxmlElement("w:instrText")
    instrText.text = "REF " + refName + " \\h"
    r.append(instrText)
    # run 3
    run = paragraph.add_run()
    r = run._r
    fldChar = OxmlElement("w:fldChar")
    fldChar.set(qn("w:fldCharType"), "end")
    r.append(fldChar)


def add_bookmark(paragraph, bookmark_text, bookmark_name):
    """

    :param paragraph:
    :param bookmark_text:
    :param bookmark_name:
    :return:
    """
    from docx.oxml.ns import qn
    from docx.oxml.shared import OxmlElement

    run = paragraph.add_run()
    tag = run._r  # for reference the following also works: tag =  document.element.xpath('//w:r')[-1]
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), "0")
    start.set(qn("w:name"), bookmark_name)
    tag.append(start)

    text = OxmlElement("w:r")
    text.text = bookmark_text
    tag.append(text)

    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), "0")
    end.set(qn("w:name"), bookmark_name)
    tag.append(end)


def insert_caption(pg, prefix, run, text, doc_format):
    """

    :param pg:
    :param prefix:
    :param run:
    :param text:
    :param doc_format:
    :type doc_format: paradoc.Formatting
    :return:
    """
    from docx.text.run import Run

    heading_ref = "Appendix" if doc_format.is_appendix is True else '"Heading 1"'

    seq1 = pg._element._new_r()
    add_seq_reference(seq1, f"STYLEREF \\s {heading_ref} \\n", run._parent)
    run._element.addprevious(seq1)
    stroke = pg._element._new_r()
    new_run = Run(stroke, run._parent)
    new_run.text = "-"
    run._element.addprevious(stroke)
    seq2 = pg._element._new_r()
    add_seq_reference(seq2, f"SEQ {prefix} \\* ARABIC \\s 1", run._parent)
    run._element.addprevious(seq2)
    fin = pg._element._new_r()
    fin_run = Run(fin, run._parent)
    fin_run.text = ": " + text
    run._element.addprevious(fin)


def insert_caption_into_runs(pg, prefix, doc_format):
    """

    :param pg:
    :param prefix:
    :param doc_format:
    :type doc_format: paradoc.Formatting
    :return:
    """

    tmp_split = pg.text.split(":")
    prefix_old = tmp_split[0].strip()
    text = tmp_split[-1].strip()
    srun = pg.runs[0]
    if len(pg.runs) > 1:
        run = pg.runs[1]
        tmp_str = pg.runs[0].text
        pg.runs[0].text = f"{prefix} "
        insert_caption(pg, prefix, run, tmp_str.split(":")[-1].strip(), doc_format)
    else:
        srun.text = f"{prefix} "
        run = pg.add_run()
        insert_caption(pg, prefix, run, text, doc_format)

    return srun, pg, prefix_old


def format_captions(pg, doc_format):
    """

    :param pg:
    :param doc_format:
    :type doc_format: paradoc.Formatting
    :return:
    """
    ref_dict = dict()
    style_name = pg.style.name
    logging.debug(style_name)
    tmp_split = pg.text.split(":")
    prefix = tmp_split[0].strip()
    if style_name == "Image Caption":
        ref_dict[prefix] = insert_caption_into_runs(pg, "Figure", doc_format)
    elif style_name == "Table Caption":
        ref_dict[prefix] = insert_caption_into_runs(pg, "Table", doc_format)
    else:
        raise ValueError("Not possible")

    return ref_dict


def add_seq_reference(run_in, seq, parent):
    from docx.text.run import Run

    new_run = Run(run_in, parent)
    r = new_run._r
    fldChar = OxmlElement("w:fldChar")
    fldChar.set(qn("w:fldCharType"), "begin")
    r.append(fldChar)
    instrText = OxmlElement("w:instrText")
    instrText.text = seq
    r.append(instrText)
    fldChar = OxmlElement("w:fldChar")
    fldChar.set(qn("w:fldCharType"), "end")
    r.append(fldChar)
    return new_run
