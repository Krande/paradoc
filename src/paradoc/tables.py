import logging

from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def add_table_reference(paragraph, seq=" SEQ Table \\* ARABIC \\s 1"):
    run = paragraph.add_run()
    r = run._r
    fldChar = OxmlElement("w:fldChar")
    fldChar.set(qn("w:fldCharType"), "begin")
    r.append(fldChar)
    instrText = OxmlElement("w:instrText")
    instrText.text = seq
    r.append(instrText)
    fldChar = OxmlElement("w:fldChar")
    fldChar.set(qn("w:fldCharType"), "end")
    r.append(fldChar)

    return run


def format_table(tbl, document, table_format):
    from docx.shared import Pt

    new_tbl_style = document.styles[table_format]
    tbl.style = new_tbl_style
    logging.info(f'Changed Table style from "{tbl.style}" to "{new_tbl_style}"')
    # tbl.paragraph_format.space_after = Pt(12)
    for i, row in enumerate(tbl.rows):
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    # run.style = document.styles["Normal"]
                    font.name = "Arial"
                    font.size = Pt(12)
                    if i == 0:
                        font.bold = True
                    else:
                        font.bold = False
