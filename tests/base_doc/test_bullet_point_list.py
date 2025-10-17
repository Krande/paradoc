import os

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from paradoc import OneDoc


def test_bullet_point_list(files_dir, tmp_path):
    source_dir = files_dir / "doc_bullet_points"
    dest = tmp_path / "bullet_points/bullet_points.docx"

    one = OneDoc(source_dir, work_dir=dest.parent)
    one.compile(dest.stem)


def test_bullet_point_list_from_scratch(tmp_path):
    def create_list(paragraph, list_type):
        p = paragraph._p  # access to xml paragraph element
        pPr = p.get_or_add_pPr()  # access paragraph properties
        numPr = OxmlElement("w:numPr")  # create number properties element
        numId = OxmlElement("w:numId")  # create numId element - sets bullet type
        numId.set(qn("w:val"), list_type)  # set list type/indentation
        numPr.append(numId)  # add bullet type to number properties list
        pPr.append(numPr)  # add number properties to paragraph

    ordered = "5"
    unordered = "1"

    document = Document()

    paragraph = document.add_paragraph("Hello", "List Paragraph")
    create_list(paragraph, unordered)

    paragraph = document.add_paragraph("Hello Again", "List Paragraph")
    create_list(paragraph, unordered)

    paragraph = document.add_paragraph("Goodbye", "List Paragraph")
    create_list(paragraph, unordered)

    paragraph = document.add_paragraph("Hello", "List Paragraph")
    create_list(paragraph, ordered)

    paragraph = document.add_paragraph("Hello Again", "List Paragraph")
    create_list(paragraph, ordered)

    paragraph = document.add_paragraph("Goodbye", "List Paragraph")
    create_list(paragraph, ordered)

    document.save(tmp_path / "bullet_list_demo.docx")
