"""Debug script to examine bookmark issues in cross-reference test."""

import base64
import pathlib
from docx import Document
from docx.oxml.ns import qn


def analyze_docx_bookmarks(docx_path: pathlib.Path, label: str):
    """Analyze bookmarks in a docx file."""
    print(f"\n{'='*80}")
    print(f"ANALYZING: {label}")
    print(f"File: {docx_path}")
    print(f"{'='*80}")

    # Open as Document to use python-docx
    doc = Document(str(docx_path))

    # Find all bookmarks
    bookmarks = []
    for para in doc.paragraphs:
        p_element = para._p
        for child in p_element:
            if child.tag == qn("w:bookmarkStart"):
                bm_id = child.get(qn("w:id"))
                bm_name = child.get(qn("w:name"))
                # Find the paragraph text
                para_text = para.text[:60] if para.text else "(empty)"
                bookmarks.append({"id": bm_id, "name": bm_name, "para_text": para_text})

    print(f"\nFound {len(bookmarks)} bookmarks:")
    for bm in bookmarks:
        print(f"  ID={bm['id']:6s}  Name={bm['name']:30s}  Para: {bm['para_text']}")

    # Also check for bookmark ends
    bookmark_ends = []
    for para in doc.paragraphs:
        p_element = para._p
        for child in p_element:
            if child.tag == qn("w:bookmarkEnd"):
                end_id = child.get(qn("w:id"))
                bookmark_ends.append(end_id)

    print(f"\nFound {len(bookmark_ends)} bookmark ends")
    print(f"Bookmark end IDs: {', '.join(sorted(set(bookmark_ends)))}")

    # Check for ID mismatches
    start_ids = set(bm["id"] for bm in bookmarks)
    end_ids = set(bookmark_ends)

    missing_ends = start_ids - end_ids
    missing_starts = end_ids - start_ids

    if missing_ends:
        print(f"\n⚠️  WARNING: Bookmark starts without matching ends: {missing_ends}")
    if missing_starts:
        print(f"\n⚠️  WARNING: Bookmark ends without matching starts: {missing_starts}")

    # Find all REF fields - need to look deeper into run structure
    ref_fields = []
    for para in doc.paragraphs:
        p_element = para._p
        # Look at all descendants, not just direct children
        for elem in p_element.iter():
            if elem.tag == qn("w:instrText"):
                if elem.text and " REF " in elem.text:
                    ref_fields.append({"instr": elem.text, "para_text": para.text[:60]})

    print(f"\nFound {len(ref_fields)} REF fields:")
    for rf in ref_fields:
        print(f"  {rf['instr']:40s}  Para: {rf['para_text']}")

    return bookmarks, bookmark_ends, ref_fields


def main():
    """Run the test and analyze intermediate files."""
    import tempfile
    from paradoc import OneDoc

    # Create temporary directory
    tmp_path = pathlib.Path(tempfile.mkdtemp())
    print(f"Working in: {tmp_path}")

    try:
        source_dir = tmp_path / "test_doc"
        main_dir = source_dir / "00-main"
        main_dir.mkdir(parents=True)

        # Create test images
        png_data = base64.b64decode(
            "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mP8/5+hHgAHggJ/PchI7wAAAABJRU5ErkJggg=="
        )

        images_dir = main_dir / "images"
        images_dir.mkdir(parents=True)

        for i in range(3):
            img_path = images_dir / f"fig{i+1}.png"
            img_path.write_bytes(png_data)

        # Create markdown with 2 sections and 2 figures
        md_content = """# Section 1: Introduction

As shown in [@fig:trends], the historical trends demonstrate a clear pattern.

![Historical trends visualization](images/fig1.png){#fig:trends}

# Section 2: Analysis

As shown in [@fig:analysis], the analysis reveals important insights. Compare this with [@fig:trends] from the previous section.

![Analysis results visualization](images/fig2.png){#fig:analysis}
"""

        md_file = main_dir / "test.md"
        md_file.write_text(md_content, encoding="utf-8")

        # Create metadata.yaml
        metadata_content = """---
figureTitle: "Figure"
tableTitle: "Table"
figPrefix: "Figure"
tblPrefix: "Table"
---"""
        metadata_file = source_dir / "metadata.yaml"
        metadata_file.write_text(metadata_content, encoding="utf-8")

        # Compile the document
        work_dir = tmp_path / "work"
        one = OneDoc(source_dir, work_dir=work_dir)
        one.compile("test_crossref", auto_open=False, export_format="docx", update_docx_with_com=False)

        output_file = work_dir / "_dist" / "test_crossref.docx"

        # Analyze intermediate files
        build_dir = work_dir / "_build"

        print("\n" + "=" * 80)
        print("STEP 1: Examining individual MD -> DOCX files")
        print("=" * 80)

        # Find all .docx files in build directory
        docx_files = list(build_dir.glob("*.docx"))
        for docx_file in docx_files:
            analyze_docx_bookmarks(docx_file, f"Individual file: {docx_file.name}")

        print("\n" + "=" * 80)
        print("STEP 2: Examining final composed DOCX")
        print("=" * 80)

        analyze_docx_bookmarks(output_file, "Final composed document")

        # Keep temp files for manual inspection
        print(f"\n{'='*80}")
        print(f"Temporary files kept at: {tmp_path}")
        print(f"Build dir: {build_dir}")
        print(f"Output file: {output_file}")
        print(f"{'='*80}")

    except Exception as e:
        print(f"\nERROR: {e}")
        import traceback

        traceback.print_exc()


if __name__ == "__main__":
    main()
