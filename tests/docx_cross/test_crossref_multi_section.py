"""Test cross-reference numbering across multiple sections."""

import base64
import platform

import pytest
from docx import Document

from paradoc import OneDoc


def test_figure_crossref_numbering_across_sections(tmp_path):
    """Verify figure cross-references show correct chapter-section numbering.

    This test ensures that:
    1. Figures in different sections get correct numbering (1-1, 2-1, etc.)
    2. Cross-references to these figures display the correct numbers
    3. The "Figure" label is included in all cross-references
    """
    source_dir = tmp_path / "test_doc"
    main_dir = source_dir / "00-main"
    main_dir.mkdir(parents=True)

    # Create test images
    png_data = base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mP8/5+hHgAHggJ/PchI7wAAAABJRU5ErkJggg=="
    )

    images_dir = main_dir / "images"
    images_dir.mkdir(parents=True)

    img1_path = images_dir / "fig1.png"
    img1_path.write_bytes(png_data)

    img2_path = images_dir / "fig2.png"
    img2_path.write_bytes(png_data)

    # Create markdown with two sections, each with a figure
    md_content = """# Section 1

This is the first section with a reference to [@fig:first_figure].

![First Figure Caption](images/fig1.png){#fig:first_figure}

Another reference to the first figure: [@fig:first_figure].

# Section 2

This is the second section with a reference to [@fig:second_figure].

![Second Figure Caption](images/fig2.png){#fig:second_figure}

Cross-reference to second figure: [@fig:second_figure].

Also reference the first figure from section 1: [@fig:first_figure].
"""

    md_file = main_dir / "test.md"
    md_file.write_text(md_content, encoding="utf-8")

    # Compile
    work_dir = tmp_path / "work"
    one = OneDoc(source_dir, work_dir=work_dir)
    one.compile("test_output", auto_open=False, export_format="docx")

    output_file = work_dir / "_dist" / "test_output.docx"

    # Re-open document after field update
    doc = Document(str(output_file))

    # Find all paragraphs with "Figure" references
    figure_references = {}
    for para in doc.paragraphs:
        text = para.text
        if "Figure" in text and any(keyword in text.lower() for keyword in ["reference", "cross-reference"]):
            # Extract which figure is being referenced based on context
            if "first section" in text.lower() or ("section 1" in text.lower() and "first figure" in text.lower()):
                figure_references.setdefault("first_in_section1", []).append(text)
            elif "second section" in text.lower():
                figure_references.setdefault("second_in_section2", []).append(text)
            elif "section 1" in text.lower():
                figure_references.setdefault("first_from_section2", []).append(text)
            else:
                # Generic references
                if "@fig:first_figure" in md_content and "first" in text.lower():
                    figure_references.setdefault("first_general", []).append(text)
                elif "@fig:second_figure" in md_content and "second" in text.lower():
                    figure_references.setdefault("second_general", []).append(text)

    print("\n=== Cross-Reference Analysis ===")
    for para in doc.paragraphs:
        text = para.text
        if "Figure" in text and ("reference" in text.lower() or "section" in text.lower()):
            print(f"Para: {text}")

    # Verify all references include "Figure" label
    all_refs = []
    for para in doc.paragraphs:
        text = para.text
        if any(word in text.lower() for word in ["reference", "cross-reference"]) and "Figure" in text:
            all_refs.append(text)
            # Each reference should have "Figure" followed by a number
            assert "Figure" in text, f"Reference missing 'Figure' label: {text}"

            # Should have format "Figure X-Y" where X and Y are digits
            import re

            fig_match = re.search(r"Figure\s+(\d+)-(\d+)", text)
            assert fig_match is not None, f"Reference doesn't have correct numbering format: {text}"

            chapter_num = int(fig_match.group(1))
            fig_num = int(fig_match.group(2))

            print(f"  Found: Figure {chapter_num}-{fig_num}")

            # Verify the numbering makes sense (chapter 1 or 2, figure 1)
            assert chapter_num in [1, 2], f"Chapter number should be 1 or 2, got {chapter_num}"
            assert fig_num == 1, f"Figure number should be 1, got {fig_num}"

    # Should have found at least 4 references (2 for first figure, 2 for second figure, 1 cross-ref)
    assert len(all_refs) >= 4, f"Expected at least 4 references with Figure label, found {len(all_refs)}"

    print(
        f"\n✓ All {len(all_refs)} cross-references have correct format with 'Figure' label and chapter-section numbering"
    )

@pytest.mark.skipif(platform.system() != "Windows", reason="COM automation only available on Windows")
def test_multiple_figures_in_same_section(tmp_path):
    """Test that multiple figures in the same section get sequential numbering."""
    source_dir = tmp_path / "test_doc"
    main_dir = source_dir / "00-main"
    main_dir.mkdir(parents=True)

    # Create test images
    png_data = base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mP8/5+hHgAHggJ/PchI7wAAAABJRU5ErkJggg=="
    )

    images_dir = main_dir / "images"
    images_dir.mkdir(parents=True)

    for i in range(3):
        img_path = images_dir / f"fig{i+1}.png"
        img_path.write_bytes(png_data)

    # Create markdown with multiple figures in one section
    md_content = """# Section 1

Reference to first: [@fig:fig1].

![Figure 1 Caption](images/fig1.png){#fig:fig1}

Reference to second: [@fig:fig2].

![Figure 2 Caption](images/fig2.png){#fig:fig2}

Reference to third: [@fig:fig3].

![Figure 3 Caption](images/fig3.png){#fig:fig3}

All references: [@fig:fig1], [@fig:fig2], and [@fig:fig3].
"""

    md_file = main_dir / "test.md"
    md_file.write_text(md_content, encoding="utf-8")

    # Compile
    work_dir = tmp_path / "work"
    one = OneDoc(source_dir, work_dir=work_dir)
    one.compile("test_output", auto_open=False, export_format="docx")

    output_file = work_dir / "_dist" / "test_output.docx"

    # Re-open document after field update
    doc = Document(str(output_file))

    print("\n=== Multiple Figures Test ===")

    import re

    figure_numbers_found = []

    for para in doc.paragraphs:
        text = para.text
        if "Figure" in text and "reference" in text.lower():
            print(f"Para: {text}")
            # Extract all figure numbers from this paragraph
            fig_matches = re.findall(r"Figure\s+(\d+)-(\d+)", text)
            for chapter, fig_num in fig_matches:
                figure_numbers_found.append(f"{chapter}-{fig_num}")
                print(f"  Found: Figure {chapter}-{fig_num}")

    # Should have found references to 1-1, 1-2, and 1-3
    # (The "All references" line should have all three)
    assert len(figure_numbers_found) >= 6, f"Expected at least 6 figure references, found {len(figure_numbers_found)}"

    # Verify we have the expected numbering (all in chapter 1)
    unique_numbers = set(figure_numbers_found)
    assert "1-1" in unique_numbers, "Should have Figure 1-1"
    assert "1-2" in unique_numbers, "Should have Figure 1-2"
    assert "1-3" in unique_numbers, "Should have Figure 1-3"

    print(f"\n✓ Found all expected figure numbers: {sorted(unique_numbers)}")
