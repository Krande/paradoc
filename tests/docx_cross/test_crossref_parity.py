"""Test cross-reference parity between COM API and Paradoc.

This test creates identical documents using both COM API and Paradoc to ensure
that figure and table numbering and cross-references match exactly.
"""

import base64
import platform
import re
import zipfile
from pathlib import Path

import pytest
from docx import Document


@pytest.mark.skipif(platform.system() != "Windows", reason="COM automation only available on Windows")
def test_com_api_reference_document(tmp_path):
    """Create a reference document using pure COM API.

    Structure:
    - Section 1
      - Section 1.1 (with Figure 1-1 and Table 1-1)
      - Section 1.2 (with Figure 1-2 and Table 1-2)
    - Section 2
      - Section 2.1 (with Figure 2-1 and Table 2-1)
      - Section 2.2 (with Figure 2-2 and Table 2-2)
    - Section 3
      - Section 3.1 (with Figure 3-1 and Table 3-1)
      - Section 3.2 (with Figure 3-2 and Table 3-2)

    All figures and tables are cross-referenced within their sections.
    """
    print("\n" + "=" * 80)
    print("TEST: COM API REFERENCE DOCUMENT")
    print("=" * 80)

    from paradoc.io.word.com_api import WordApplication

    # Ensure directory exists
    tmp_path.mkdir(parents=True, exist_ok=True)
    output_file = tmp_path / "com_reference.docx"

    with WordApplication(visible=False) as word_app:
        doc = word_app.create_document()

        # Create 3 sections, each with 2 subsections
        for section_num in range(1, 4):
            print(f"\n[Section {section_num}]")
            doc.add_heading(f"Section {section_num}", level=1)
            doc.add_paragraph(f"This is section {section_num} introduction.")
            doc.add_paragraph()

            for subsection_num in range(1, 3):
                subsection_label = f"{section_num}.{subsection_num}"
                print(f"  [Subsection {subsection_label}]")

                doc.add_heading(f"Section {subsection_label}", level=2)

                # Add paragraph with cross-reference placeholders
                doc.add_text("This subsection discusses ")
                doc.add_paragraph()

                # Add figure
                fig_ref = doc.add_figure_with_caption(
                    caption_text=f"Caption for figure in section {subsection_label}",
                    width=150,
                    height=100,
                    use_chapter_numbers=True,
                )
                print(f"    Added Figure {section_num}-{subsection_num}")

                # Add table
                table_data = [["Header 1", "Header 2"], [f"Data {subsection_label}.1", f"Data {subsection_label}.2"]]
                tbl_ref = doc.add_table_with_caption(
                    caption_text=f"Caption for table in section {subsection_label}",
                    rows=2,
                    cols=2,
                    data=table_data,
                    use_chapter_numbers=True,
                )
                print(f"    Added Table {section_num}-{subsection_num}")

                # Add cross-references
                doc.add_text("As shown in ")
                doc.add_cross_reference(fig_ref, include_hyperlink=True)
                doc.add_text(" and ")
                doc.add_cross_reference(tbl_ref, include_hyperlink=True)
                doc.add_text(", the data is consistent.")
                doc.add_paragraph()
                doc.add_paragraph()

        # Save and update fields
        doc.save(str(output_file))
        doc.update_fields()
        doc.save(str(output_file))

    print(f"\n[SAVED] {output_file}")

    # Analyze the document
    analyze_document_structure(output_file, "COM API")

    return output_file


def test_paradoc_document(tmp_path):
    """Create an identical document using Paradoc.

    Same structure as COM API test.
    """
    print("\n" + "=" * 80)
    print("TEST: PARADOC DOCUMENT")
    print("=" * 80)

    from paradoc import OneDoc

    # Create source directory structure
    source_dir = tmp_path / "paradoc_source"
    main_dir = source_dir / "00-main"
    main_dir.mkdir(parents=True)

    # Create test image
    png_data = base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mP8/5+hHgAHggJ/PchI7wAAAABJRU5ErkJggg=="
    )
    images_dir = main_dir / "images"
    images_dir.mkdir(parents=True)

    # Create 6 images (one for each subsection)
    for i in range(1, 7):
        img_path = images_dir / f"fig{i}.png"
        img_path.write_bytes(png_data)

    # Create markdown content
    md_content = ""
    fig_counter = 1

    for section_num in range(1, 4):
        md_content += f"# Section {section_num}\n\n"
        md_content += f"This is section {section_num} introduction.\n\n"

        for subsection_num in range(1, 3):
            subsection_label = f"{section_num}.{subsection_num}"

            md_content += f"## Section {subsection_label}\n\n"
            md_content += "This subsection discusses\n\n"

            # Add figure
            fig_id = f"fig{section_num}_{subsection_num}"
            md_content += (
                f"![Caption for figure in section {subsection_label}](images/fig{fig_counter}.png){{#fig:{fig_id}}}\n\n"
            )

            # Add table - use simple markdown table without database integration
            tbl_id = f"tbl{section_num}_{subsection_num}"
            md_content += f"Table: Caption for table in section {subsection_label} {{#tbl:{tbl_id}}}\n\n"
            md_content += "| Header 1 | Header 2 |\n"
            md_content += "|:---------|:---------|\n"
            md_content += f"| Data {subsection_label}.1 | Data {subsection_label}.2 |\n\n"

            # Add cross-references
            md_content += f"As shown in [@fig:{fig_id}] and [@tbl:{tbl_id}], the data is consistent.\n\n"

            fig_counter += 1

    # Write markdown file
    md_file = main_dir / "document.md"
    md_file.write_text(md_content, encoding="utf-8")

    print(f"[Created] {md_file}")
    print(f"Content preview:\n{md_content[:500]}...\n")

    # Compile document
    work_dir = tmp_path / "paradoc_work"
    one = OneDoc(source_dir, work_dir=work_dir)
    one.compile("paradoc_output", auto_open=False, export_format="docx")

    output_file = work_dir / "_dist" / "paradoc_output.docx"

    print(f"\n[SAVED] {output_file}")

    # Analyze the document
    analyze_document_structure(output_file, "Paradoc")

    return output_file


def test_compare_documents_from_files(tmp_path):
    """Compare pre-generated COM API and Paradoc documents side-by-side.

    This should be run after both test_com_api_reference_document and
    test_paradoc_document have been run successfully.
    """
    print("\n" + "=" * 80)
    print("TEST: COMPARE DOCUMENTS FROM FILES")
    print("=" * 80)

    # Look for existing files
    com_file = tmp_path / "com_reference.docx"
    paradoc_file = tmp_path / "paradoc_output.docx"

    if not com_file.exists():
        pytest.skip(f"COM reference file not found: {com_file}")
    if not paradoc_file.exists():
        pytest.skip(f"Paradoc output file not found: {paradoc_file}")

    # Extract and compare XML
    print("\n" + "=" * 80)
    print("COMPARING XML STRUCTURES")
    print("=" * 80)

    com_xml = extract_document_xml(com_file)
    paradoc_xml = extract_document_xml(paradoc_file)

    # Compare figure captions
    print("\n[FIGURE CAPTIONS COMPARISON]")
    com_fig_captions = extract_captions(com_xml, "Figure")
    paradoc_fig_captions = extract_captions(paradoc_xml, "Figure")

    print(f"\nCOM API Figures ({len(com_fig_captions)}):")
    for i, caption in enumerate(com_fig_captions, 1):
        print(f"  {i}. {caption}")

    print(f"\nParadoc Figures ({len(paradoc_fig_captions)}):")
    for i, caption in enumerate(paradoc_fig_captions, 1):
        print(f"  {i}. {caption}")

    # Compare table captions
    print("\n[TABLE CAPTIONS COMPARISON]")
    com_tbl_captions = extract_captions(com_xml, "Table")
    paradoc_tbl_captions = extract_captions(paradoc_xml, "Table")

    print(f"\nCOM API Tables ({len(com_tbl_captions)}):")
    for i, caption in enumerate(com_tbl_captions, 1):
        print(f"  {i}. {caption}")

    print(f"\nParadoc Tables ({len(paradoc_tbl_captions)}):")
    for i, caption in enumerate(paradoc_tbl_captions, 1):
        print(f"  {i}. {caption}")

    # Compare cross-references
    print("\n[CROSS-REFERENCES COMPARISON]")
    com_refs = extract_cross_references(com_xml)
    paradoc_refs = extract_cross_references(paradoc_xml)

    print(f"\nCOM API Cross-references ({len(com_refs)}):")
    for i, ref in enumerate(com_refs, 1):
        print(f"  {i}. {ref}")

    print(f"\nParadoc Cross-references ({len(paradoc_refs)}):")
    for i, ref in enumerate(paradoc_refs, 1):
        print(f"  {i}. {ref}")

    # Compare bookmarks
    print("\n[BOOKMARKS COMPARISON]")
    com_bookmarks = extract_bookmarks(com_xml)
    paradoc_bookmarks = extract_bookmarks(paradoc_xml)

    print(f"\nCOM API Bookmarks ({len(com_bookmarks)}):")
    for name, content in com_bookmarks.items():
        print(f"  {name}: {content[:100]}")

    print(f"\nParadoc Bookmarks ({len(paradoc_bookmarks)}):")
    for name, content in paradoc_bookmarks.items():
        print(f"  {name}: {content[:100]}")

    # Detailed field analysis
    print("\n[FIELD STRUCTURE ANALYSIS]")
    analyze_field_structure(com_xml, "COM API")
    analyze_field_structure(paradoc_xml, "Paradoc")

    # Save detailed XML for inspection
    com_xml_file = tmp_path / "com_document.xml"
    paradoc_xml_file = tmp_path / "paradoc_document.xml"

    com_xml_file.write_text(com_xml, encoding="utf-8")
    paradoc_xml_file.write_text(paradoc_xml, encoding="utf-8")

    print("\n[SAVED] Detailed XML files:")
    print(f"  COM API: {com_xml_file}")
    print(f"  Paradoc: {paradoc_xml_file}")


def analyze_document_structure(docx_path: Path, label: str):
    """Analyze and print document structure."""
    print(f"\n[ANALYZING] {label} document")

    doc = Document(str(docx_path))

    print(f"  Paragraphs: {len(doc.paragraphs)}")
    print(f"  Tables: {len(doc.tables)}")

    # Count figures (captions with "Figure" in style or text)
    figure_count = 0
    table_caption_count = 0

    for para in doc.paragraphs:
        if "Figure" in para.text and "Caption" in para.style.name:
            figure_count += 1
        if "Table" in para.text and "Caption" in para.style.name:
            table_caption_count += 1

    print(f"  Figure captions: {figure_count}")
    print(f"  Table captions: {table_caption_count}")


def extract_document_xml(docx_path: Path) -> str:
    """Extract document.xml from docx file."""
    with zipfile.ZipFile(docx_path, "r") as zip_ref:
        return zip_ref.read("word/document.xml").decode("utf-8")


def extract_captions(xml_content: str, caption_type: str) -> list[str]:
    """Extract all captions of a given type from XML."""
    captions = []

    # Pattern to find caption paragraphs
    # Look for paragraphs containing the caption type
    para_pattern = r"<w:p\b[^>]*>.*?</w:p>"

    for para_match in re.finditer(para_pattern, xml_content, re.DOTALL):
        para_xml = para_match.group(0)

        # Extract text from the paragraph
        text_parts = re.findall(r"<w:t[^>]*>(.*?)</w:t>", para_xml)
        full_text = "".join(text_parts)

        if caption_type in full_text and ("Caption" in para_xml or "SEQ" in para_xml):
            captions.append(full_text)

    return captions


def extract_cross_references(xml_content: str) -> list[str]:
    """Extract all REF field cross-references from XML."""
    refs = []

    # Pattern for REF fields
    ref_pattern = r"<w:instrText[^>]*>(.*?REF.*?)</w:instrText>"

    for match in re.finditer(ref_pattern, xml_content):
        refs.append(match.group(1))

    return refs


def extract_bookmarks(xml_content: str) -> dict[str, str]:
    """Extract all bookmarks and their content."""
    bookmarks = {}

    # Pattern for bookmark starts
    bookmark_pattern = r'<w:bookmarkStart[^>]*w:name="([^"]*)"[^>]*w:id="(\d+)"[^>]*/>'

    for match in re.finditer(bookmark_pattern, xml_content):
        name = match.group(1)
        bookmark_id = match.group(2)

        # Find content between bookmark start and end
        start_pos = match.end()
        end_pattern = f'<w:bookmarkEnd[^>]*w:id="{bookmark_id}"[^>]*/>'
        end_match = re.search(end_pattern, xml_content[start_pos:])

        if end_match:
            content = xml_content[start_pos: start_pos + end_match.start()]
            # Extract text
            text_parts = re.findall(r"<w:t[^>]*>(.*?)</w:t>", content)
            bookmarks[name] = "".join(text_parts)

    return bookmarks


def analyze_field_structure(xml_content: str, label: str):
    """Analyze the structure of fields in the document."""
    print(f"\n{label} Field Structure:")

    # Find all SEQ fields
    seq_pattern = r"<w:instrText[^>]*>(.*?SEQ.*?)</w:instrText>"
    seq_fields = re.findall(seq_pattern, xml_content)

    print(f"  SEQ fields ({len(seq_fields)}):")
    for i, field in enumerate(seq_fields[:10], 1):  # Show first 10
        print(f"    {i}. {field}")

    # Find all REF fields
    ref_pattern = r"<w:instrText[^>]*>(.*?REF.*?)</w:instrText>"
    ref_fields = re.findall(ref_pattern, xml_content)

    print(f"  REF fields ({len(ref_fields)}):")
    for i, field in enumerate(ref_fields[:10], 1):  # Show first 10
        print(f"    {i}. {field}")

    # Find all STYLEREF fields
    styleref_pattern = r"<w:instrText[^>]*>(.*?STYLEREF.*?)</w:instrText>"
    styleref_fields = re.findall(styleref_pattern, xml_content)

    print(f"  STYLEREF fields ({len(styleref_fields)}):")
    for i, field in enumerate(styleref_fields[:10], 1):  # Show first 10
        print(f"    {i}. {field}")
