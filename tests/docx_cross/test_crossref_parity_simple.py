"""Simplified cross-reference parity test focusing on figure numbering.

This test creates documents with figures only to isolate the cross-reference issue.
"""

import base64
import platform
import re
import zipfile
from pathlib import Path

import pytest
from docx import Document


@pytest.mark.skipif(platform.system() != "Windows", reason="COM automation only available on Windows")
def test_com_api_figures_only(tmp_path):
    """Create a COM API reference document with figures only.

    Structure:
    - Section 1
      - Section 1.1 (with Figure 1-1)
      - Section 1.2 (with Figure 1-2)
    - Section 2
      - Section 2.1 (with Figure 2-1)
      - Section 2.2 (with Figure 2-2)
    - Section 3
      - Section 3.1 (with Figure 3-1)
      - Section 3.2 (with Figure 3-2)
    """
    print("\n" + "=" * 80)
    print("TEST: COM API FIGURES ONLY")
    print("=" * 80)

    from paradoc.io.word.com_api import WordApplication

    # Ensure directory exists
    tmp_path.mkdir(parents=True, exist_ok=True)
    output_file = tmp_path / "com_figures.docx"

    with WordApplication(visible=False) as word_app:
        doc = word_app.create_document()

        # Store figure references for cross-referencing
        all_figs = []

        # Create 3 sections, each with 2 subsections
        for section_num in range(1, 4):
            print(f"\n[Section {section_num}]")
            doc.add_heading(f"Section {section_num}", level=1)
            doc.add_paragraph(f"This is section {section_num} introduction.")
            doc.add_paragraph()

            for subsection_num in range(1, 3):
                subsection_label = f"{section_num}.{subsection_num}"
                print(f"  [Subsection {subsection_label}]")

                doc.add_heading(f"Section {subsection_label}", level=2)

                # Add figure
                fig_ref = doc.add_figure_with_caption(
                    caption_text=f"Caption for figure in section {subsection_label}",
                    width=150,
                    height=100,
                    use_chapter_numbers=True,
                )
                all_figs.append(fig_ref)
                print(f"    Added Figure {section_num}-{subsection_num}")

                # Add cross-reference to this figure
                doc.add_text("As shown in ")
                doc.add_cross_reference(fig_ref, include_hyperlink=True)
                doc.add_text(", the data is consistent.")
                doc.add_paragraph()
                doc.add_paragraph()

        # Add a final section with cross-references to all figures
        doc.add_heading("Summary", level=1)
        doc.add_text("This document contains figures: ")
        for i, fig_ref in enumerate(all_figs):
            if i > 0:
                doc.add_text(", ")
            doc.add_cross_reference(fig_ref, include_hyperlink=True)
        doc.add_paragraph()

        # Save and update fields
        doc.save(str(output_file))
        doc.update_fields()
        doc.save(str(output_file))

    print(f"\n[SAVED] {output_file}")

    # Analyze the document
    doc = Document(str(output_file))
    print("\n[ANALYSIS]")
    print(f"  Paragraphs: {len(doc.paragraphs)}")

    # Extract figure numbers from XML
    xml_content = extract_document_xml(output_file)

    print("\n[FIGURE CAPTIONS]")
    captions = extract_captions(xml_content, "Figure")
    for i, caption in enumerate(captions, 1):
        print(f"  {i}. {caption}")

    print("\n[CROSS-REFERENCES]")
    refs = extract_cross_references(xml_content)
    for i, ref in enumerate(refs, 1):
        print(f"  {i}. {ref}")

    print("\n[SEQ FIELDS]")
    seq_fields = extract_seq_fields(xml_content)
    for i, field in enumerate(seq_fields, 1):
        print(f"  {i}. {field}")

    return output_file


def test_paradoc_figures_only(tmp_path):
    """Create a Paradoc document with figures only."""
    print("\n" + "=" * 80)
    print("TEST: PARADOC FIGURES ONLY")
    print("=" * 80)

    from paradoc import OneDoc

    # Create source directory structure
    source_dir = tmp_path / "paradoc_source"
    main_dir = source_dir / "00-main"
    main_dir.mkdir(parents=True)

    # Create test image
    png_data = base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mP8/5+hHgAHggJ/PchI7wAAAABJRU5ErkJggg=="
    )
    images_dir = main_dir / "images"
    images_dir.mkdir(parents=True)

    # Create 6 images (one for each subsection)
    for i in range(1, 7):
        img_path = images_dir / f"fig{i}.png"
        img_path.write_bytes(png_data)

    # Create markdown content
    md_content = ""
    fig_counter = 1
    all_fig_ids = []

    for section_num in range(1, 4):
        md_content += f"# Section {section_num}\n\n"
        md_content += f"This is section {section_num} introduction.\n\n"

        for subsection_num in range(1, 3):
            subsection_label = f"{section_num}.{subsection_num}"

            md_content += f"## Section {subsection_label}\n\n"

            # Add figure
            fig_id = f"fig{section_num}_{subsection_num}"
            all_fig_ids.append(fig_id)
            md_content += (
                f"![Caption for figure in section {subsection_label}](images/fig{fig_counter}.png){{#fig:{fig_id}}}\n\n"
            )

            # Add cross-reference to this figure
            md_content += f"As shown in [@fig:{fig_id}], the data is consistent.\n\n"

            fig_counter += 1

    # Add summary section with all cross-references
    md_content += "# Summary\n\n"
    md_content += "This document contains figures: "
    for i, fig_id in enumerate(all_fig_ids):
        if i > 0:
            md_content += ", "
        md_content += f"[@fig:{fig_id}]"
    md_content += ".\n\n"

    # Write markdown file
    md_file = main_dir / "document.md"
    md_file.write_text(md_content, encoding="utf-8")

    print(f"[Created] {md_file}")
    print(f"Content preview:\n{md_content[:500]}...\n")

    # Compile document
    work_dir = tmp_path / "paradoc_work"
    one = OneDoc(source_dir, work_dir=work_dir)
    one.compile("paradoc_output", auto_open=False, export_format="docx")

    output_file = work_dir / "_dist" / "paradoc_output.docx"

    print(f"\n[SAVED] {output_file}")

    # Analyze the document
    doc = Document(str(output_file))
    print("\n[ANALYSIS]")
    print(f"  Paragraphs: {len(doc.paragraphs)}")

    # Extract figure numbers from XML
    xml_content = extract_document_xml(output_file)

    print("\n[FIGURE CAPTIONS]")
    captions = extract_captions(xml_content, "Figure")
    for i, caption in enumerate(captions, 1):
        print(f"  {i}. {caption}")

    print("\n[CROSS-REFERENCES]")
    refs = extract_cross_references(xml_content)
    for i, ref in enumerate(refs, 1):
        print(f"  {i}. {ref}")

    print("\n[SEQ FIELDS]")
    seq_fields = extract_seq_fields(xml_content)
    for i, field in enumerate(seq_fields, 1):
        print(f"  {i}. {field}")

    return output_file


def extract_document_xml(docx_path: Path) -> str:
    """Extract document.xml from docx file."""
    with zipfile.ZipFile(docx_path, "r") as zip_ref:
        return zip_ref.read("word/document.xml").decode("utf-8")


def extract_captions(xml_content: str, caption_type: str) -> list[str]:
    """Extract all captions of a given type from XML."""
    captions = []

    # Pattern to find caption paragraphs
    para_pattern = r"<w:p\b[^>]*>.*?</w:p>"

    for para_match in re.finditer(para_pattern, xml_content, re.DOTALL):
        para_xml = para_match.group(0)

        # Extract text from the paragraph
        text_parts = re.findall(r"<w:t[^>]*>(.*?)</w:t>", para_xml)
        full_text = "".join(text_parts)

        if caption_type in full_text and ("Caption" in para_xml or "SEQ" in para_xml):
            captions.append(full_text)

    return captions


def extract_cross_references(xml_content: str) -> list[str]:
    """Extract all REF field cross-references from XML."""
    refs = []

    # Pattern for REF fields
    ref_pattern = r"<w:instrText[^>]*>(.*?REF.*?)</w:instrText>"

    for match in re.finditer(ref_pattern, xml_content):
        refs.append(match.group(1).strip())

    return refs


def extract_seq_fields(xml_content: str) -> list[str]:
    """Extract all SEQ fields from XML."""
    fields = []

    # Pattern for SEQ fields
    seq_pattern = r"<w:instrText[^>]*>(.*?SEQ.*?)</w:instrText>"

    for match in re.finditer(seq_pattern, xml_content):
        fields.append(match.group(1).strip())

    return fields


if __name__ == "__main__":
    import tempfile
    import sys

    with tempfile.TemporaryDirectory() as tmp:
        tmp_path = Path(tmp)

        if len(sys.argv) > 1 and sys.argv[1] == "com":
            test_com_api_figures_only(tmp_path / "com")
        elif len(sys.argv) > 1 and sys.argv[1] == "paradoc":
            test_paradoc_figures_only(tmp_path / "paradoc")
        else:
            print("Creating both documents...")
            com_file = test_com_api_figures_only(tmp_path / "com")
            paradoc_file = test_paradoc_figures_only(tmp_path / "paradoc")

            print("\n" + "=" * 80)
            print("COMPARISON")
            print("=" * 80)
            print(f"COM file: {com_file}")
            print(f"Paradoc file: {paradoc_file}")
