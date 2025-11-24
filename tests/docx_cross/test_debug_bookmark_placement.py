"""Debug test to examine bookmark placement in figure captions."""

from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from paradoc import OneDoc


def test_debug_figure_caption_structure(tmp_path):
    """Debug test to examine the XML structure of figure captions and bookmarks."""
    source_dir = tmp_path / "test_doc"
    main_dir = source_dir / "00-main"
    main_dir.mkdir(parents=True)

    # Create test image
    import base64

    png_data = base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mP8/5+hHgAHggJ/PchI7wAAAABJRU5ErkJggg=="
    )

    images_dir = main_dir / "images"
    images_dir.mkdir(parents=True)
    img_path = images_dir / "test.png"
    img_path.write_bytes(png_data)

    # Create markdown with figure in chapter 2
    md_content = """# Chapter 1

Some content in chapter 1.

# Chapter 2

This is a reference to [@fig:test_figure] in the text.

![Test Figure Caption](images/test.png){#fig:test_figure}
"""

    md_file = main_dir / "test.md"
    md_file.write_text(md_content, encoding="utf-8")

    # Compile
    work_dir = tmp_path / "work"
    one = OneDoc(source_dir, work_dir=work_dir)
    one.compile("test_output", auto_open=False, export_format="docx")

    output_file = work_dir / "_dist" / "test_output.docx"

    # Re-open document after field update
    doc = Document(str(output_file))

    # Find the figure caption paragraph
    caption_para = None
    for para in doc.paragraphs:
        if "Figure" in para.text and "Caption" in para.text:
            caption_para = para
            break

    assert caption_para is not None, "Figure caption not found"

    print(f"\n{'='*80}")
    print("FIGURE CAPTION ANALYSIS")
    print(f"{'='*80}")
    print(f"Caption text: {caption_para.text}")

    # Get the XML structure
    xml_str = caption_para._element.xml
    if isinstance(xml_str, bytes):
        xml_str = xml_str.decode("utf-8")

    print("\nCaption XML (first 2000 chars):")
    print(xml_str[:2000])

    # Find bookmark locations
    p_element = caption_para._p

    print(f"\n{'='*80}")
    print("ANALYZING RUNS IN CAPTION PARAGRAPH")
    print(f"{'='*80}")

    runs = list(p_element.findall(qn("w:r")))
    bookmark_start_idx = None
    bookmark_end_idx = None
    bookmark_name = None

    for idx, run in enumerate(runs):
        # Check what's before this run (bookmarkStart)
        prev_sibling = run.getprevious()
        if prev_sibling is not None and prev_sibling.tag == qn("w:bookmarkStart"):
            bookmark_start_idx = idx
            bookmark_name = prev_sibling.get(qn("w:name"))
            print(f"\n>>> BOOKMARK START found before run {idx}")
            print(f"    Bookmark name: {bookmark_name}")

        # Check what's after this run (bookmarkEnd)
        next_sibling = run.getnext()
        if next_sibling is not None and next_sibling.tag == qn("w:bookmarkEnd"):
            bookmark_end_idx = idx
            print(f"\n>>> BOOKMARK END found after run {idx}")

        # Analyze the run content
        run_text = ""
        for t_elem in run.findall(qn("w:t")):
            if t_elem.text:
                run_text += t_elem.text

        # Check for field characters
        fld_chars = run.findall(qn("w:fldChar"))
        instr_texts = run.findall(qn("w:instrText"))

        field_info = ""
        if fld_chars:
            for fld_char in fld_chars:
                fld_type = fld_char.get(qn("w:fldCharType"))
                field_info += f"[{fld_type}] "

        if instr_texts:
            for instr in instr_texts:
                if instr.text:
                    field_info += f"INSTR: {instr.text[:50]}"

        marker = ""
        if idx == bookmark_start_idx:
            marker = " <-- BOOKMARK STARTS HERE"
        elif idx == bookmark_end_idx:
            marker = " <-- BOOKMARK ENDS HERE"

        print(f"Run {idx}: text='{run_text}' {field_info}{marker}")

    print(f"\n{'='*80}")
    print("BOOKMARK COVERAGE ANALYSIS")
    print(f"{'='*80}")

    if bookmark_start_idx is not None and bookmark_end_idx is not None:
        print(f"Bookmark '{bookmark_name}' wraps runs {bookmark_start_idx} to {bookmark_end_idx}")
        print(f"This covers {bookmark_end_idx - bookmark_start_idx + 1} run(s)")

        # Determine what fields are covered
        print("\nField coverage:")
        has_styleref = False
        has_seq = False

        for idx in range(bookmark_start_idx, bookmark_end_idx + 1):
            run = runs[idx]
            instr_texts = run.findall(qn("w:instrText"))
            for instr in instr_texts:
                if instr.text:
                    if "STYLEREF" in instr.text:
                        has_styleref = True
                        print(f"  ✓ STYLEREF field at run {idx}")
                    if "SEQ" in instr.text and "STYLEREF" not in instr.text:
                        has_seq = True
                        print(f"  ✓ SEQ field at run {idx}")

        if has_styleref and has_seq:
            print("\n✅ Bookmark correctly covers BOTH STYLEREF and SEQ fields")
        elif has_seq and not has_styleref:
            print("\n❌ Bookmark only covers SEQ field - missing STYLEREF!")
        else:
            print(f"\n⚠️  Unexpected bookmark coverage: STYLEREF={has_styleref}, SEQ={has_seq}")
    else:
        print("❌ No bookmark found in caption!")

    # Now check the reference
    print(f"\n{'='*80}")
    print("CHECKING FIGURE REFERENCE")
    print(f"{'='*80}")

    ref_para = None
    for para in doc.paragraphs:
        if "reference to" in para.text.lower():
            ref_para = para
            break

    if ref_para:
        print(f"Reference text: {ref_para.text}")

        ref_xml = ref_para._element.xml
        if isinstance(ref_xml, bytes):
            ref_xml = ref_xml.decode("utf-8")

        # Extract REF field
        import re

        ref_match = re.search(r"REF\s+([^\s\\]+)", ref_xml)
        if ref_match:
            ref_bookmark = ref_match.group(1)
            print(f"REF field references bookmark: {ref_bookmark}")

            if bookmark_name and ref_bookmark == bookmark_name:
                print("✅ Reference uses the correct bookmark name")
            else:
                print(f"⚠️  Bookmark mismatch: caption={bookmark_name}, reference={ref_bookmark}")


if __name__ == "__main__":
    import tempfile

    with tempfile.TemporaryDirectory() as tmpdir:
        test_debug_figure_caption_structure(Path(tmpdir))
