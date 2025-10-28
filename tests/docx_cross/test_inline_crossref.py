"""Test that cross-references stay on the same line."""

import platform
from pathlib import Path

import pytest
from docx import Document


@pytest.mark.skipif(platform.system() != "Windows", reason="COM automation only available on Windows")
def test_cross_reference_inline(tmp_path):
    """Test that cross-references can be added inline without creating new paragraphs."""
    from paradoc.io.word.com_api import WordApplication

    output_file = Path(tmp_path) / "test_inline_crossref.docx"

    with WordApplication(visible=False, run_isolated=True) as word_app:
        doc = word_app.create_document()

        # Add heading
        doc.add_heading("Test Inline Cross-References", level=1)

        # Add figures with captions
        fig1 = doc.add_figure_with_caption(
            caption_text="First Example Figure", create_bookmark=True, use_chapter_numbers=False
        )
        doc.add_paragraph()

        fig2 = doc.add_figure_with_caption(
            caption_text="Second Example Figure", create_bookmark=True, use_chapter_numbers=False
        )
        doc.add_paragraph()

        # Add table with caption
        tbl1 = doc.add_table_with_caption(
            caption_text="Example Table", rows=2, cols=2, create_bookmark=True, use_chapter_numbers=False
        )
        doc.add_paragraph()

        # Test 1: Cross-reference inline with text BEFORE and AFTER on same line
        doc.add_text("As shown in ")
        doc.add_cross_reference(fig1, include_hyperlink=True)
        doc.add_text(", the architecture is modular.")
        doc.add_paragraph()  # Now end the paragraph
        doc.add_paragraph()

        # Test 2: Multiple cross-references on the same line
        doc.add_text("Comparing ")
        doc.add_cross_reference(fig1, include_hyperlink=True)
        doc.add_text(" with ")
        doc.add_cross_reference(fig2, include_hyperlink=True)
        doc.add_text(" reveals key differences.")
        doc.add_paragraph()
        doc.add_paragraph()

        # Test 3: Table cross-reference inline
        doc.add_text("According to ")
        doc.add_cross_reference(tbl1, include_hyperlink=True)
        doc.add_text(", the system meets requirements.")
        doc.add_paragraph()
        doc.add_paragraph()

        # Update fields and save
        doc.update_fields()
        doc.save(output_file)

    # Verify the file was created
    assert output_file.exists(), "Document should be created"
    assert output_file.stat().st_size > 0, "Document should not be empty"

    # Verify using python-docx that cross-references are on the same line
    docx_doc = Document(str(output_file))

    paragraphs_with_crossrefs = []
    for para in docx_doc.paragraphs:
        text = para.text.strip()
        if text and para.style.name != "Caption" and "Figure" in text or "Table" in text:
            paragraphs_with_crossrefs.append(text)

    print(f"\n✓ Found {len(paragraphs_with_crossrefs)} paragraphs with cross-references:")
    for i, para in enumerate(paragraphs_with_crossrefs, 1):
        print(f"  {i}. {para}")

    # Verify that the text appears inline (all on one paragraph)
    assert len(paragraphs_with_crossrefs) >= 3, "Should have at least 3 paragraphs with cross-references"

    # Test 1: Check that "As shown in Figure X, the architecture is modular." is on one line
    test1_found = False
    for para in paragraphs_with_crossrefs:
        if "As shown in" in para and "architecture is modular" in para:
            test1_found = True
            print(f"\n✓ Test 1 PASSED: Text before and after cross-reference on same line:\n  '{para}'")
            break
    assert test1_found, "Test 1: Should find inline text with cross-reference"

    # Test 2: Check that "Comparing Figure X with Figure Y reveals key differences." is on one line
    test2_found = False
    for para in paragraphs_with_crossrefs:
        if "Comparing" in para and "with" in para and "reveals key differences" in para:
            test2_found = True
            print(f"\n✓ Test 2 PASSED: Multiple cross-references on same line:\n  '{para}'")
            break
    assert test2_found, "Test 2: Should find multiple cross-references on same line"

    # Test 3: Check that "According to Table X, the system meets requirements." is on one line
    test3_found = False
    for para in paragraphs_with_crossrefs:
        if "According to" in para and "Table" in para and "meets requirements" in para:
            test3_found = True
            print(f"\n✓ Test 3 PASSED: Table cross-reference inline:\n  '{para}'")
            break
    assert test3_found, "Test 3: Should find table cross-reference inline"

    print("\n✓ All inline cross-reference tests passed!")
    print(f"✓ Document saved to: {output_file}")


if __name__ == "__main__":
    import tempfile

    with tempfile.TemporaryDirectory() as tmp:
        test_cross_reference_inline(Path(tmp))
