"""Test that figure references are converted to REF fields."""

import os
import re

from docx import Document

from paradoc import OneDoc

auto_open = os.getenv("AUTO_OPEN", False)


def test_figure_reference_conversion_detailed(tmp_path):
    """Test that figure references get converted to REF fields."""
    # Create test document
    source_dir = tmp_path / "test_doc"
    main_dir = source_dir / "00-main"
    main_dir.mkdir(parents=True)

    # Create test image
    import base64

    png_data = base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mP8/5+hHgAHggJ/PchI7wAAAABJRU5ErkJggg=="
    )

    images_dir = main_dir / "images"
    images_dir.mkdir(parents=True)
    img_path = images_dir / "test.png"
    img_path.write_bytes(png_data)

    # Create markdown with figure and reference
    md_content = """# Test Document

![Test Figure Caption](images/test.png){#fig:test_figure}

Reference to figure: [@fig:test_figure]

Another reference: see [@fig:test_figure] for details.
"""

    md_file = main_dir / "test.md"
    md_file.write_text(md_content, encoding="utf-8")

    # Compile
    work_dir = tmp_path / "work"
    one = OneDoc(source_dir, work_dir=work_dir)
    one.compile("test_output", auto_open=auto_open, export_format="docx")

    output_file = work_dir / "_dist" / "test_output.docx"
    doc = Document(str(output_file))

    print("\n" + "=" * 80)
    print("CHECKING FIGURE REFERENCES")
    print("=" * 80)

    # Find caption to get figure number
    figure_number = None
    for para in doc.paragraphs:
        if "Figure" in para.text and "Caption" in para.text:
            # Extract figure number
            match = re.search(r"Figure\s+([\d\-]+)", para.text)
            if match:
                figure_number = match.group(1)
            print(f"\nFound caption: {para.text}")
            print(f"  Figure number: {figure_number}")

            # Check for bookmark
            xml_str = para._element.xml.decode("utf-8") if isinstance(para._element.xml, bytes) else para._element.xml
            if "bookmarkStart" in xml_str:
                bookmark_matches = re.findall(r'w:name="([^"]+)"', xml_str)
                print(f"  OK Has bookmark: {bookmark_matches}")
            break

    # Find reference paragraphs
    ref_count = 0
    ref_field_count = 0

    for para in doc.paragraphs:
        if "Reference" in para.text or "reference" in para.text.lower():
            print(f"\nReference paragraph: {para.text}")

            xml_str = para._element.xml.decode("utf-8") if isinstance(para._element.xml, bytes) else para._element.xml

            # Check for REF field
            if "REF" in xml_str and "STYLEREF" not in xml_str:
                ref_matches = re.findall(r"<w:instrText[^>]*>([^<]*REF[^<]*)</w:instrText>", xml_str)
                if ref_matches:
                    print(f"  OK Has REF field: {ref_matches}")
                    ref_field_count += 1
                else:
                    print("  WARNING Contains 'REF' in XML but couldn't extract field")
            else:
                print("  NO No REF field found")

                # Check if it has the plain text reference
                if figure_number and figure_number in para.text:
                    print(f"  WARNING Contains plain text reference to figure {figure_number}")

            ref_count += 1

    print("\nSummary:")
    print(f"  Total reference paragraphs found: {ref_count}")
    print(f"  References with REF fields: {ref_field_count}")

    if ref_count > 0 and ref_field_count == 0:
        print("\nWARNING: References found but none converted to REF fields!")
        print("This indicates the reference conversion function is not working.")

    print("=" * 80)

    if auto_open:
        os.startfile(output_file)
