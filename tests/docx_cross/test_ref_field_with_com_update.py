"""Test that figure references work correctly with Word COM field updates."""

import base64
import os
import re

import pytest
from docx import Document

from paradoc import OneDoc
from paradoc.io.word.com_api import is_word_com_available

auto_open = os.getenv("AUTO_OPEN", False)


@pytest.mark.skipif(not is_word_com_available(), reason="COM automation only if Word COM is available")
def test_figure_reference_with_com_update(tmp_path):
    """Test that figure references work correctly after COM field update.

    This test verifies that the bookmark normalization (_Reffig_test_figure)
    works with Word's field update mechanism and doesn't produce
    "Error! Not a valid bookmark self-reference" errors.
    """
    # Create test document
    source_dir = tmp_path / "test_doc"
    main_dir = source_dir / "00-main"
    main_dir.mkdir(parents=True)

    png_data = base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mP8/5+hHgAHggJ/PchI7wAAAABJRU5ErkJggg=="
    )

    images_dir = main_dir / "images"
    images_dir.mkdir(parents=True)
    img_path = images_dir / "test.png"
    img_path.write_bytes(png_data)

    # Create markdown with figure and reference
    md_content = """# Test Document

![Test Figure Caption](images/test.png){#fig:test_figure}

Reference to figure: [@fig:test_figure]

Another reference: see [@fig:test_figure] for details.
"""

    md_file = main_dir / "test.md"
    md_file.write_text(md_content, encoding="utf-8")

    # Compile with COM automation enabled
    work_dir = tmp_path / "work"
    one = OneDoc(source_dir, work_dir=work_dir)
    one.compile("test_output_wcom", auto_open=False, export_format="docx", enable_word_com_automation=True)

    output_file = work_dir / "_dist" / "test_output_wcom.docx"

    # Reload document to check results after field update
    doc = Document(str(output_file))

    print("\n" + "=" * 80)
    print("CHECKING DOCUMENT AFTER COM UPDATE")
    print("=" * 80)

    # Find caption with bookmark
    caption_found = False
    bookmark_name = None
    for para in doc.paragraphs:
        if "Figure" in para.text and "Caption" in para.text:
            caption_found = True
            print(f"\nFound caption: {para.text}")

            # Check for bookmark
            xml_str = para._element.xml.decode("utf-8") if isinstance(para._element.xml, bytes) else para._element.xml
            if "bookmarkStart" in xml_str:
                bookmark_matches = re.findall(r'w:name="([^"]+)"', xml_str)
                if bookmark_matches:
                    bookmark_name = bookmark_matches[0]
                    print(f"  ✓ Has bookmark: {bookmark_name}")
            break

    # Find reference paragraphs and check for errors
    ref_paragraphs = []
    error_found = False

    for para in doc.paragraphs:
        if "Reference" in para.text or "reference" in para.text.lower():
            ref_paragraphs.append(para)
            print(f"\nReference paragraph: {para.text}")

            # Check for error messages
            if "Error!" in para.text or "not a valid bookmark" in para.text.lower():
                error_found = True
                print(f"  ✗ ERROR DETECTED: {para.text}")
            else:
                # Check for proper figure number (should be "Figure 1", not "Figure -")
                if "Figure" in para.text:
                    figure_match = re.search(r"Figure\s+([\d]+)", para.text)
                    if figure_match:
                        fig_num = figure_match.group(1)
                        print(f"  ✓ Contains valid figure reference: Figure {fig_num}")
                    else:
                        print("  ⚠ Figure reference format unclear")

            # Check XML for REF field structure
            xml_str = para._element.xml.decode("utf-8") if isinstance(para._element.xml, bytes) else para._element.xml
            if "REF" in xml_str and "STYLEREF" not in xml_str:
                ref_matches = re.findall(r"<w:instrText[^>]*>([^<]*REF[^<]*)</w:instrText>", xml_str)
                if ref_matches:
                    print(f"  REF field instruction: {ref_matches[0]}")

    print("\nSummary:")
    print(f"  Caption found: {caption_found}")
    print(f"  Bookmark name: {bookmark_name}")
    print(f"  Reference paragraphs found: {len(ref_paragraphs)}")
    print(f"  Errors detected: {error_found}")

    if error_found:
        print("\n❌ FAIL: Error messages found in references after COM update!")
        print("This indicates the bookmark format is not compatible with Word's REF field update.")
    else:
        print("\n✅ PASS: No errors found - bookmark format is compatible with Word COM update")

    print("=" * 80)

    if auto_open:
        os.startfile(output_file)

    # Assertions
    assert caption_found, "Caption should be found in document"
    assert bookmark_name is not None, "Caption should have bookmark"
    assert bookmark_name.startswith("_Ref"), f"Bookmark should start with '_Ref' but is '{bookmark_name}'"
    assert len(ref_paragraphs) == 2, f"Should find 2 reference paragraphs, found {len(ref_paragraphs)}"
    assert not error_found, "Should not have 'Error! Not a valid bookmark' messages in references"
