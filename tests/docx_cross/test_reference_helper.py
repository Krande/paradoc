"""Test the ReferenceHelper class for managing cross-references."""

import pytest
from docx import Document

from paradoc.io.word.reference_helper import ReferenceHelper


def test_reference_helper_initialization():
    """Test that ReferenceHelper initializes correctly."""
    helper = ReferenceHelper()

    stats = helper.get_statistics()
    assert stats["figures"] == 0
    assert stats["tables"] == 0
    assert stats["equations"] == 0
    assert stats["total"] == 0


def test_register_figure():
    """Test registering figures."""
    helper = ReferenceHelper()

    # Register first figure
    bookmark1 = helper.register_figure("test_figure_1")
    assert bookmark1.startswith("_Ref")
    assert len(bookmark1) == 13  # _Ref + 9 digits

    # Register second figure
    bookmark2 = helper.register_figure("test_figure_2")
    assert bookmark2.startswith("_Ref")
    assert bookmark1 != bookmark2  # Should be unique

    # Register same figure again - should return same bookmark
    bookmark1_again = helper.register_figure("test_figure_1")
    assert bookmark1_again == bookmark1

    stats = helper.get_statistics()
    assert stats["figures"] == 2
    assert stats["total"] == 2


def test_register_table():
    """Test registering tables."""
    helper = ReferenceHelper()

    # Register first table
    bookmark1 = helper.register_table("test_table_1")
    assert bookmark1.startswith("_Ref")

    # Register second table
    bookmark2 = helper.register_table("test_table_2")
    assert bookmark2.startswith("_Ref")
    assert bookmark1 != bookmark2

    stats = helper.get_statistics()
    assert stats["tables"] == 2
    assert stats["total"] == 2


def test_register_mixed_items():
    """Test registering mixed figures and tables."""
    helper = ReferenceHelper()

    # Register in specific order
    helper.register_figure("fig1")
    helper.register_table("tbl1")
    helper.register_figure("fig2")
    helper.register_table("tbl2")

    stats = helper.get_statistics()
    assert stats["figures"] == 2
    assert stats["tables"] == 2
    assert stats["total"] == 4

    # Check document order is preserved
    assert helper._all_items[0].semantic_id == "fig1"
    assert helper._all_items[1].semantic_id == "tbl1"
    assert helper._all_items[2].semantic_id == "fig2"
    assert helper._all_items[3].semantic_id == "tbl2"


def test_get_bookmarks_in_order():
    """Test retrieving bookmarks in document order."""
    helper = ReferenceHelper()

    fig1 = helper.register_figure("fig1")
    tbl1 = helper.register_table("tbl1")
    fig2 = helper.register_figure("fig2")
    tbl2 = helper.register_table("tbl2")
    fig3 = helper.register_figure("fig3")

    # Get figures in order (should skip tables)
    figure_bookmarks = helper.get_all_figure_bookmarks_in_order()
    assert len(figure_bookmarks) == 3
    assert figure_bookmarks == [fig1, fig2, fig3]

    # Get tables in order (should skip figures)
    table_bookmarks = helper.get_all_table_bookmarks_in_order()
    assert len(table_bookmarks) == 2
    assert table_bookmarks == [tbl1, tbl2]


def test_get_bookmark_by_semantic_id():
    """Test retrieving bookmarks by semantic ID."""
    helper = ReferenceHelper()

    fig_bookmark = helper.register_figure("my_test_figure")
    tbl_bookmark = helper.register_table("my_test_table")

    # Retrieve by semantic ID
    assert helper.get_figure_bookmark("my_test_figure") == fig_bookmark
    assert helper.get_table_bookmark("my_test_table") == tbl_bookmark

    # Non-existent IDs should return None
    assert helper.get_figure_bookmark("nonexistent") is None
    assert helper.get_table_bookmark("nonexistent") is None


def test_update_display_numbers():
    """Test updating display numbers from caption paragraphs."""
    helper = ReferenceHelper()
    doc = Document()

    # Create caption paragraphs with numbers
    para1 = doc.add_paragraph("Figure 1-1: First figure caption")
    para2 = doc.add_paragraph("Figure 1-2: Second figure caption")
    para3 = doc.add_paragraph("Table 1-1: First table caption")

    # Register items with caption paragraphs
    helper.register_figure("fig1", para1)
    helper.register_figure("fig2", para2)
    helper.register_table("tbl1", para3)

    # Update display numbers
    helper.update_display_numbers()

    # Check that numbers were extracted
    assert helper._figures["fig1"].display_number == "1-1"
    assert helper._figures["fig2"].display_number == "1-2"
    assert helper._tables["tbl1"].display_number == "1-1"


def test_print_registry(capsys):
    """Test printing the registry for debugging."""
    helper = ReferenceHelper()

    helper.register_figure("fig1")
    helper.register_table("tbl1")

    helper.print_registry()

    captured = capsys.readouterr()
    assert "[ReferenceHelper] Complete Registry:" in captured.out
    assert "Total items: 2" in captured.out
    assert "Figures: 1" in captured.out
    assert "Tables: 1" in captured.out
    assert "fig1" in captured.out
    assert "tbl1" in captured.out


def test_extract_and_convert_hyperlink_references():
    """Test extracting hyperlink references and converting them using the new method."""
    from docx.oxml import parse_xml
    from docx.oxml.ns import qn

    helper = ReferenceHelper()
    doc = Document()

    # Register some figures and tables
    fig1_bookmark = helper.register_figure("test_figure")
    tbl1_bookmark = helper.register_table("results_table")

    # Create a paragraph with hyperlinks simulating pandoc-crossref output
    para = doc.add_paragraph()
    p_element = para._p

    # Add text "See "
    run1_xml = """
    <w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
        <w:t>See fig.</w:t>
    </w:r>
    """
    run1 = parse_xml(run1_xml)
    p_element.append(run1)

    # Add hyperlink for figure
    hyperlink1_xml = """
    <w:hyperlink xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:anchor="fig:test_figure">
        <w:r>
            <w:t>1</w:t>
        </w:r>
    </w:hyperlink>
    """
    hyperlink1 = parse_xml(hyperlink1_xml)
    p_element.append(hyperlink1)

    # Add text " and tbl."
    run2_xml = """
    <w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
        <w:t> and tbl.</w:t>
    </w:r>
    """
    run2 = parse_xml(run2_xml)
    p_element.append(run2)

    # Add hyperlink for table
    hyperlink2_xml = """
    <w:hyperlink xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:anchor="tbl:results_table">
        <w:r>
            <w:t>1</w:t>
        </w:r>
    </w:hyperlink>
    """
    hyperlink2 = parse_xml(hyperlink2_xml)
    p_element.append(hyperlink2)

    # Extract hyperlink references
    hyperlink_refs = helper.extract_hyperlink_references(doc)

    # Should find 2 references
    assert len(hyperlink_refs) == 2

    # Check the first reference
    assert hyperlink_refs[0].anchor == "fig:test_figure"
    assert hyperlink_refs[0].semantic_id == "test_figure"
    assert hyperlink_refs[0].word_bookmark == fig1_bookmark
    assert hyperlink_refs[0].label == "Figure"
    assert hyperlink_refs[0].prefix_text is not None
    assert "fig." in hyperlink_refs[0].prefix_text.lower()

    # Check the second reference
    assert hyperlink_refs[1].anchor == "tbl:results_table"
    assert hyperlink_refs[1].semantic_id == "results_table"
    assert hyperlink_refs[1].word_bookmark == tbl1_bookmark
    assert hyperlink_refs[1].label == "Table"
    assert hyperlink_refs[1].prefix_text is not None
    assert "tbl." in hyperlink_refs[1].prefix_text.lower()

    # Now convert the hyperlink references using the new method
    helper.convert_hyperlink_references(hyperlink_refs)

    # Check that the paragraph has been modified
    # The hyperlinks should be replaced with REF fields
    # And the prefix text should be removed
    from lxml import etree

    paragraph_xml = etree.tostring(para._p, encoding="unicode")

    # Should contain REF field instructions
    assert "REF" in paragraph_xml
    assert fig1_bookmark in paragraph_xml
    assert tbl1_bookmark in paragraph_xml

    # Should contain the labels
    assert "Figure" in paragraph_xml or "Figure" in para.text
    assert "Table" in paragraph_xml or "Table" in para.text

    # The hyperlinks should be gone
    hyperlinks = p_element.findall(
        ".//w:hyperlink", namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    )
    # Should have no hyperlinks with our specific anchors
    for hyperlink in hyperlinks:
        anchor = hyperlink.get(qn("w:anchor"))
        assert anchor not in ["fig:test_figure", "tbl:results_table"]


if __name__ == "__main__":
    pytest.main([__file__, "-v"])
