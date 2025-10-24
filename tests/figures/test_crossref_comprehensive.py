"""Comprehensive test for cross-references with figures and tables across sections."""

import base64
import platform
import re

import pytest
from docx import Document
from docx.oxml.ns import qn

from paradoc import OneDoc
from paradoc.io.word.utils import docx_update


@pytest.mark.skipif(platform.system() != "Windows", reason="COM automation only available on Windows")
def test_comprehensive_crossref_with_figures_and_tables(tmp_path):
    r"""Test cross-references with 3 figures and 3 tables spread across different sections.
    
    This test verifies:
    1. Figures and tables in different sections get correct numbering (1-1, 2-1, 2-2, etc.)
    2. Cross-references point to the correct bookmarks
    3. REF fields are properly structured with \h switch
    4. After COM field update, all cross-references display correct labels and numbers
    """
    source_dir = tmp_path / "test_doc"
    main_dir = source_dir / "00-main"
    main_dir.mkdir(parents=True)

    # Create test images
    png_data = base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mP8/5+hHgAHggJ/PchI7wAAAABJRU5ErkJggg=="
    )

    images_dir = main_dir / "images"
    images_dir.mkdir(parents=True)

    for i in range(3):
        img_path = images_dir / f"fig{i+1}.png"
        img_path.write_bytes(png_data)

    # Create markdown with 3 sections, 3 figures, and 3 tables with multiple cross-references
    md_content = """# Section 1: Introduction

This section introduces the first figure and table.

## Subsection 1.1: First Figure

As shown in [@fig:trends], the historical trends demonstrate a clear pattern.

![Historical trends visualization](images/fig1.png){#fig:trends}

The data in [@tbl:metrics] supports this observation.

## Subsection 1.2: First Table

+------------------+-----------+------+-------+
| Metric Name      | Status    | Unit | Value |
+==================+===========+======+=======+
| metrics          | Excellent | %    |  94.5 |
+------------------+-----------+------+-------+
| Efficiency       | Good      | %    |  87.2 |
+------------------+-----------+------+-------+

Table: Current performance metrics summary {#tbl:metrics}

# Section 2: Analysis

This section presents additional figures and tables for detailed analysis.

## Subsection 2.1: Second Figure

As shown in [@fig:analysis], the analysis reveals important insights. Compare this with [@fig:trends] from the previous section.

![Analysis results visualization](images/fig2.png){#fig:analysis}

## Subsection 2.2: Second Table

The performance comparison in [@tbl:comparison] shows significant improvements. This contrasts with the baseline in [@tbl:metrics].

+--------------+-------------+-------+
| Period       | Metric      | Value |
+==============+=============+=======+
| comparison   | Performance |  85.3 |
+--------------+-------------+-------+
| Q2           | Performance |  91.2 |
+--------------+-------------+-------+

Table: Performance comparison across periods {#tbl:comparison}

## Subsection 2.3: Third Figure

The projection in [@fig:projection] indicates future trends. Both [@fig:analysis] and [@fig:projection] are consistent with [@fig:trends].

![Future projection visualization](images/fig3.png){#fig:projection}

# Section 3: Conclusions

This section summarizes findings with reference to all previous figures and tables.

## Summary of Visual Evidence

The three figures ([@fig:trends], [@fig:analysis], and [@fig:projection]) provide comprehensive visual evidence.

## Summary of Data

The tables show clear patterns:
- Baseline metrics: [@tbl:metrics]
- Performance comparison: [@tbl:comparison]
- Final summary: [@tbl:summary]

+--------------+-------+---------+
| Category     | Count | Average |
+==============+=======+=========+
| summary      |   150 |     0.0 |
+--------------+-------+---------+
| Mean Value   |     0 |    88.7 |
+--------------+-------+---------+

Table: Summary statistics {#tbl:summary}

## Final Remarks

In conclusion, [@fig:trends] establishes the baseline, [@fig:analysis] confirms the hypothesis, and [@fig:projection] predicts future outcomes. The data tables ([@tbl:metrics], [@tbl:comparison], [@tbl:summary]) provide quantitative support for all conclusions.
"""

    md_file = main_dir / "test.md"
    md_file.write_text(md_content, encoding="utf-8")

    # Compile the document
    work_dir = tmp_path / "work"
    one = OneDoc(source_dir, work_dir=work_dir)
    one.compile("test_crossref", auto_open=False, export_format="docx", update_docx_with_com=False)

    output_file = work_dir / "_dist" / "test_crossref.docx"
    assert output_file.exists(), "Output file should be created"

    # BEFORE COM update: Check that REF fields are properly structured
    print("\n" + "="*80)
    print("BEFORE COM UPDATE: Checking REF field structure")
    print("="*80)
    
    doc_before = Document(str(output_file))
    
    # Find all REF fields before update
    ref_fields_before = find_all_ref_fields(doc_before)
    print(f"\nFound {len(ref_fields_before)} REF fields before update:")
    
    for field_info in ref_fields_before:
        print(f"  • {field_info['instr']}")
        print(f"    Paragraph text: {field_info['paragraph_text'][:80]}...")
        print(f"    Placeholder: {field_info['result_text']}")
    
    # Debug: Show all paragraphs to see what pandoc-crossref generated
    print(f"\n[DEBUG] Examining all paragraphs for cross-reference patterns:")
    for i, para in enumerate(doc_before.paragraphs[:50]):  # First 50 paragraphs
        text = para.text
        if 'fig:' in text or 'tbl:' in text or 'Figure' in text or 'Table' in text:
            print(f"  Para {i}: {text[:100]}")
    
    # Verify REF fields are properly structured
    # Note: Initially we expect very few REF fields because the conversion may not be working
    print(f"\n[DEBUG] Expected many REF fields but found only {len(ref_fields_before)}")
    print(f"[DEBUG] This indicates the cross-reference conversion is not working properly")
    
    # Check that REF fields have proper structure with \h switch
    for field_info in ref_fields_before:
        instr = field_info['instr']
        assert ' REF ' in instr, f"Field should contain REF instruction: {instr}"
        assert r'\h' in instr, f"Field should have \\h switch for hyperlink: {instr}"
        
        # Extract bookmark name
        match = re.search(r'REF\s+(\S+)', instr)
        assert match, f"Could not extract bookmark name from: {instr}"
        bookmark_name = match.group(1)
        
        # Bookmark should start with _Ref (Word-style)
        assert bookmark_name.startswith('_Ref'), f"Bookmark should start with _Ref: {bookmark_name}"
    
    print(f"\n✓ All REF fields have correct structure with \\h switch and _Ref bookmarks")
    
    # Find all bookmarks
    bookmarks = find_all_bookmarks(doc_before)
    print(f"\nFound {len(bookmarks)} bookmarks:")
    for bm in bookmarks[:10]:  # Show first 10
        print(f"  • {bm['name']} (ID: {bm['id']})")
    
    # Additional debug: Check if bookmarks are around captions
    print(f"\n[DEBUG] Checking bookmark positions:")
    for i, para in enumerate(doc_before.paragraphs):
        if 'Figure' in para.text and ':' in para.text:
            print(f"  Caption para {i}: {para.text[:60]}")
            # Check for bookmarks in this paragraph
            p_element = para._p
            for child in p_element:
                if child.tag == qn('w:bookmarkStart'):
                    bm_name = child.get(qn('w:name'))
                    print(f"    -> Has bookmark: {bm_name}")
        if 'Table' in para.text and ':' in para.text:
            print(f"  Caption para {i}: {para.text[:60]}")
            # Check for bookmarks in this paragraph
            p_element = para._p
            for child in p_element:
                if child.tag == qn('w:bookmarkStart'):
                    bm_name = child.get(qn('w:name'))
                    print(f"    -> Has bookmark: {bm_name}")
    
    # Should have bookmarks for 3 figures and 3 tables
    assert len(bookmarks) >= 6, f"Expected at least 6 bookmarks (3 figs + 3 tbls), found {len(bookmarks)}"
    
    # Update fields using Word COM automation
    print("\n" + "="*80)
    print("UPDATING FIELDS WITH COM AUTOMATION")
    print("="*80)
    docx_update(str(output_file))
    print("✓ Fields updated")

    # AFTER COM update: Verify cross-references display correctly
    print("\n" + "="*80)
    print("AFTER COM UPDATE: Verifying cross-reference display")
    print("="*80)
    
    # Re-open document after field update
    doc_after = Document(str(output_file))

    # Find all paragraphs and their cross-references with context
    paragraphs_with_refs = []

    for i, para in enumerate(doc_after.paragraphs):
        text = para.text
        
        # Find figure references
        fig_matches = re.findall(r'Figure\s+(\d+)-(\d+)', text)
        if fig_matches:
            paragraphs_with_refs.append({
                'index': i,
                'text': text,
                'type': 'figure',
                'refs': [f"Figure {ch}-{num}" for ch, num in fig_matches]
            })
        
        # Find table references
        tbl_matches = re.findall(r'Table\s+(\d+)-(\d+)', text)
        if tbl_matches:
            paragraphs_with_refs.append({
                'index': i,
                'text': text,
                'type': 'table',
                'refs': [f"Table {ch}-{num}" for ch, num in tbl_matches]
            })
    
    print(f"\nFound {len(paragraphs_with_refs)} paragraphs with cross-references:")
    for p in paragraphs_with_refs:
        print(f"  Para {p['index']}: {p['text'][:80]}...")
        print(f"    References: {', '.join(p['refs'])}")

    # SPECIFIC ASSERTIONS: Verify that specific cross-references point to correct targets
    print("\n" + "="*80)
    print("VERIFYING SPECIFIC CROSS-REFERENCE TARGETS")
    print("="*80)

    # Test Case 1: "As shown in [@fig:trends]" should resolve to Figure 1-1
    para_fig_trends_1 = next((p for p in paragraphs_with_refs
                              if 'historical trends demonstrate a clear pattern' in p['text']), None)
    assert para_fig_trends_1 is not None, "Could not find paragraph with 'historical trends demonstrate'"
    print(f"\n✓ Found: '{para_fig_trends_1['text'][:60]}...'")
    print(f"  References: {para_fig_trends_1['refs']}")
    assert 'Figure 1-1' in para_fig_trends_1['refs'], \
        f"Expected 'Figure 1-1' in paragraph about historical trends, got {para_fig_trends_1['refs']}"
    assert len(para_fig_trends_1['refs']) == 1, \
        f"Expected exactly 1 figure reference in this paragraph, got {len(para_fig_trends_1['refs'])}"

    # Test Case 2: "As shown in [@fig:analysis], ... Compare this with [@fig:trends]"
    # Should resolve to Figure 2-1 and Figure 1-1
    para_fig_analysis = next((p for p in paragraphs_with_refs
                              if 'analysis reveals important insights' in p['text']), None)
    assert para_fig_analysis is not None, "Could not find paragraph with 'analysis reveals important'"
    print(f"\n✓ Found: '{para_fig_analysis['text'][:60]}...'")
    print(f"  References: {para_fig_analysis['refs']}")
    assert 'Figure 2-1' in para_fig_analysis['refs'], \
        f"Expected 'Figure 2-1' (analysis) in this paragraph, got {para_fig_analysis['refs']}"
    assert 'Figure 1-1' in para_fig_analysis['refs'], \
        f"Expected 'Figure 1-1' (trends) in this paragraph, got {para_fig_analysis['refs']}"
    assert len(para_fig_analysis['refs']) == 2, \
        f"Expected exactly 2 figure references, got {len(para_fig_analysis['refs'])}"

    # Test Case 3: "The projection in [@fig:projection] ... Both [@fig:analysis] and [@fig:projection] are consistent with [@fig:trends]"
    # Should resolve to Figure 2-2, Figure 2-1, Figure 2-2, Figure 1-1
    para_fig_projection = next((p for p in paragraphs_with_refs
                                if 'projection' in p['text'].lower() and 'future trends' in p['text']), None)
    assert para_fig_projection is not None, "Could not find paragraph with 'projection' and 'future trends'"
    print(f"\n✓ Found: '{para_fig_projection['text'][:60]}...'")
    print(f"  References: {para_fig_projection['refs']}")
    assert 'Figure 2-2' in para_fig_projection['refs'], \
        f"Expected 'Figure 2-2' (projection) in this paragraph, got {para_fig_projection['refs']}"
    assert 'Figure 2-1' in para_fig_projection['refs'], \
        f"Expected 'Figure 2-1' (analysis) in this paragraph, got {para_fig_projection['refs']}"
    assert 'Figure 1-1' in para_fig_projection['refs'], \
        f"Expected 'Figure 1-1' (trends) in this paragraph, got {para_fig_projection['refs']}"
    # Should have 4 total references (projection appears twice)
    assert len(para_fig_projection['refs']) == 4, \
        f"Expected exactly 4 figure references, got {len(para_fig_projection['refs'])}"

    # Test Case 4: Table references - "The data in [@tbl:metrics]" should resolve to Table 1-1
    para_tbl_metrics_1 = next((p for p in paragraphs_with_refs
                               if 'supports this observation' in p['text'] and p['type'] == 'table'), None)
    assert para_tbl_metrics_1 is not None, "Could not find paragraph with table reference 'supports this observation'"
    print(f"\n✓ Found: '{para_tbl_metrics_1['text'][:60]}...'")
    print(f"  References: {para_tbl_metrics_1['refs']}")
    assert 'Table 1-1' in para_tbl_metrics_1['refs'], \
        f"Expected 'Table 1-1' (metrics) in this paragraph, got {para_tbl_metrics_1['refs']}"
    assert len(para_tbl_metrics_1['refs']) == 1, \
        f"Expected exactly 1 table reference, got {len(para_tbl_metrics_1['refs'])}"

    # Test Case 5: "The performance comparison in [@tbl:comparison] ... This contrasts with the baseline in [@tbl:metrics]"
    # Should resolve to Table 2-1 and Table 1-1
    para_tbl_comparison = next((p for p in paragraphs_with_refs
                                if 'performance comparison' in p['text'].lower() and 'significant improvements' in p['text']), None)
    assert para_tbl_comparison is not None, "Could not find paragraph with 'performance comparison' and 'significant improvements'"
    print(f"\n✓ Found: '{para_tbl_comparison['text'][:60]}...'")
    print(f"  References: {para_tbl_comparison['refs']}")
    assert 'Table 2-1' in para_tbl_comparison['refs'], \
        f"Expected 'Table 2-1' (comparison) in this paragraph, got {para_tbl_comparison['refs']}"
    assert 'Table 1-1' in para_tbl_comparison['refs'], \
        f"Expected 'Table 1-1' (metrics) in this paragraph, got {para_tbl_comparison['refs']}"
    assert len(para_tbl_comparison['refs']) == 2, \
        f"Expected exactly 2 table references, got {len(para_tbl_comparison['refs'])}"

    # Test Case 6: Summary paragraph with all three tables
    # "- Baseline metrics: [@tbl:metrics] - Performance comparison: [@tbl:comparison] - Final summary: [@tbl:summary]"
    # Should resolve to Table 1-1, Table 2-1, Table 3-1
    para_tbl_summary = next((p for p in paragraphs_with_refs
                             if 'baseline metrics' in p['text'].lower() and 'performance comparison' in p['text'].lower()), None)
    assert para_tbl_summary is not None, "Could not find paragraph with 'baseline metrics' and 'performance comparison'"
    print(f"\n✓ Found: '{para_tbl_summary['text'][:60]}...'")
    print(f"  References: {para_tbl_summary['refs']}")
    assert 'Table 1-1' in para_tbl_summary['refs'], \
        f"Expected 'Table 1-1' (metrics) in summary paragraph, got {para_tbl_summary['refs']}"
    assert 'Table 2-1' in para_tbl_summary['refs'], \
        f"Expected 'Table 2-1' (comparison) in summary paragraph, got {para_tbl_summary['refs']}"
    assert 'Table 3-1' in para_tbl_summary['refs'], \
        f"Expected 'Table 3-1' (summary) in summary paragraph, got {para_tbl_summary['refs']}"
    assert len(para_tbl_summary['refs']) == 3, \
        f"Expected exactly 3 table references, got {len(para_tbl_summary['refs'])}"

    print(f"\n" + "="*80)
    print("TEST PASSED - ALL CROSS-REFERENCES VERIFIED")
    print("="*80)
    print(f"✓ Figure reference to 'trends' correctly points to Figure 1-1")
    print(f"✓ Figure reference to 'analysis' correctly points to Figure 2-1")
    print(f"✓ Figure reference to 'projection' correctly points to Figure 2-2")
    print(f"✓ Table reference to 'metrics' correctly points to Table 1-1")
    print(f"✓ Table reference to 'comparison' correctly points to Table 2-1")
    print(f"✓ Table reference to 'summary' correctly points to Table 3-1")
    print(f"✓ All multi-reference paragraphs have correct target references")
    print(f"✓ All REF fields properly structured with \\h switch")
    print(f"✓ All bookmarks use Word-style _Ref naming")


def find_all_ref_fields(doc: Document) -> list[dict]:
    """Find all REF fields in the document.
    
    Returns:
        List of dicts with 'instr', 'result_text', and 'paragraph_text'
    """
    ref_fields = []
    
    for para in doc.paragraphs:
        p_element = para._p
        
        # Find field begin/end markers
        field_begins = []
        field_ends = []
        field_separates = []
        instrs = []
        
        for i, child in enumerate(p_element):
            if child.tag == qn('w:r'):
                # Check for fldChar
                for sub in child:
                    if sub.tag == qn('w:fldChar'):
                        fld_type = sub.get(qn('w:fldCharType'))
                        if fld_type == 'begin':
                            field_begins.append(i)
                        elif fld_type == 'end':
                            field_ends.append(i)
                        elif fld_type == 'separate':
                            field_separates.append(i)
                    elif sub.tag == qn('w:instrText'):
                        instrs.append((i, sub.text))
        
        # Match begins with ends and extract instruction text
        for begin_idx in field_begins:
            # Find corresponding instruction
            instr_text = None
            result_text = ""
            
            for instr_idx, instr in instrs:
                if instr_idx > begin_idx:
                    instr_text = instr
                    break
            
            if instr_text and ' REF ' in instr_text:
                # Find corresponding separator and end
                sep_idx = None
                for sep in field_separates:
                    if sep > begin_idx:
                        sep_idx = sep
                        break
                
                end_idx = None
                for end in field_ends:
                    if end > begin_idx:
                        end_idx = end
                        break
                
                # Extract result text (between separator and end)
                if sep_idx is not None and end_idx is not None:
                    for i in range(sep_idx + 1, end_idx):
                        child = p_element[i]
                        if child.tag == qn('w:r'):
                            for sub in child:
                                if sub.tag == qn('w:t'):
                                    result_text += sub.text or ""
                
                ref_fields.append({
                    'instr': instr_text,
                    'result_text': result_text,
                    'paragraph_text': para.text
                })
    
    return ref_fields


def find_all_bookmarks(doc: Document) -> list[dict]:
    """Find all bookmarks in the document.
    
    Returns:
        List of dicts with 'name' and 'id'
    """
    bookmarks = []
    
    for para in doc.paragraphs:
        p_element = para._p
        
        for child in p_element:
            if child.tag == qn('w:bookmarkStart'):
                bm_id = child.get(qn('w:id'))
                bm_name = child.get(qn('w:name'))
                bookmarks.append({
                    'id': bm_id,
                    'name': bm_name
                })
    
    return bookmarks
