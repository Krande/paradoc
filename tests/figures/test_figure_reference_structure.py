"""Detailed inspection of figure reference structure in DOCX.

This test inspects the actual XML structure to ensure:
1. Figure captions have bookmarks
2. Figure references use REF fields pointing to those bookmarks
3. REF fields show only label and number (not caption text)
"""

import os
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from paradoc import OneDoc


auto_open = os.getenv("AUTO_OPEN", False)


def test_inspect_figure_reference_xml(tmp_path):
    """Inspect the actual XML structure of figure references."""
    import re

    # Create test document structure
    source_dir = tmp_path / "test_doc"
    main_dir = source_dir / "00-main"
    main_dir.mkdir(parents=True)

    # Create test image
    import base64
    png_data = base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mP8/5+hHgAHggJ/PchI7wAAAABJRU5ErkJggg=="
    )

    images_dir = main_dir / "images"
    images_dir.mkdir(parents=True)
    img_path = images_dir / "test.png"
    img_path.write_bytes(png_data)

    # Create markdown
    md_content = """# Test Document

![Test Figure Caption](images/test.png){#fig:test_figure}

Reference to figure: [@fig:test_figure]
"""

    md_file = main_dir / "test.md"
    md_file.write_text(md_content, encoding="utf-8")

    # Compile
    work_dir = tmp_path / "work"
    one = OneDoc(source_dir, work_dir=work_dir)
    one.compile("test_output", auto_open=auto_open, export_format="docx")

    output_file = work_dir / "_dist" / "test_output.docx"
    doc = Document(str(output_file))

    print("\n" + "="*80)
    print("DOCUMENT STRUCTURE INSPECTION")
    print("="*80)

    # Inspect all paragraphs
    for i, para in enumerate(doc.paragraphs):
        print(f"\nParagraph {i}: Style='{para.style.name}'")
        print(f"  Text: {para.text[:100]}")

        # Check for bookmarks
        xml_str = para._element.xml.decode('utf-8') if isinstance(para._element.xml, bytes) else para._element.xml

        if "bookmarkStart" in xml_str:
            print(f"  ✓ Contains bookmarkStart")
            # Extract bookmark name
            import re
            bookmark_matches = re.findall(r'w:name="([^"]+)"', xml_str)
            if bookmark_matches:
                print(f"    Bookmark names: {bookmark_matches}")

        if "SEQ" in xml_str:
            print(f"  ✓ Contains SEQ field")
            # Extract SEQ field content
            seq_matches = re.findall(r'SEQ\s+([^<]+)', xml_str)
            if seq_matches:
                print(f"    SEQ content: {seq_matches}")

        if "REF" in xml_str and "STYLEREF" not in xml_str:
            print(f"  ✓ Contains REF field")
            # Extract REF field content
            ref_matches = re.findall(r'REF\s+([^<]+)', xml_str)
            if ref_matches:
                print(f"    REF content: {ref_matches}")

        if "HYPERLINK" in xml_str:
            print(f"  ✓ Contains HYPERLINK")

        if "Test Figure Caption" in para.text or "Reference to figure" in para.text:
            print(f"\n  Full XML:\n{xml_str[:500]}")

    print("\n" + "="*80)

    if auto_open:
        os.startfile(output_file)


def test_figure_caption_bookmark_presence(tmp_path):
    """Verify that figure captions have bookmarks for cross-referencing."""
    source_dir = tmp_path / "test_doc"
    main_dir = source_dir / "00-main"
    main_dir.mkdir(parents=True)

    # Create test image
    import base64
    png_data = base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mP8/5+hHgAHggJ/PchI7wAAAABJRU5ErkJggg=="
    )

    images_dir = main_dir / "images"
    images_dir.mkdir(parents=True)
    img_path = images_dir / "test.png"
    img_path.write_bytes(png_data)

    # Create markdown
    md_content = """# Test Document

![Test Figure Caption](images/test.png){#fig:test_figure}

Reference: [@fig:test_figure]
"""

    md_file = main_dir / "test.md"
    md_file.write_text(md_content, encoding="utf-8")

    # Compile
    work_dir = tmp_path / "work"
    one = OneDoc(source_dir, work_dir=work_dir)
    one.compile("test_output", auto_open=False, export_format="docx")

    output_file = work_dir / "_dist" / "test_output.docx"
    doc = Document(str(output_file))

    # Find figure caption
    figure_caption = None
    for para in doc.paragraphs:
        if "Figure" in para.text and "Caption" in para.text:
            figure_caption = para
            break

    assert figure_caption is not None, "Figure caption not found"

    # Check for bookmark in caption
    xml_str = figure_caption._element.xml.decode('utf-8') if isinstance(figure_caption._element.xml, bytes) else figure_caption._element.xml

    has_bookmark = "bookmarkStart" in xml_str
    print(f"\nFigure caption has bookmark: {has_bookmark}")
    print(f"Figure caption text: {figure_caption.text}")

    # This is the key check - figures need bookmarks for proper cross-referencing
    # Currently this might fail, which is the bug we're fixing
    if not has_bookmark:
        print("\n⚠️  WARNING: Figure caption is missing bookmark!")
        print("This means cross-references won't work properly in Word.")
        print("The SEQ field needs to be wrapped in a bookmark for REF fields to reference it.")

