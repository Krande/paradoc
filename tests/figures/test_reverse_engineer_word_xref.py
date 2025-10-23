"""Reverse engineer Word's cross-reference system by comparing before/after."""

import os
import platform
import re
import time
import zipfile
from pathlib import Path

import pytest
from docx import Document
from docx.oxml.ns import qn


auto_open = os.getenv("AUTO_OPEN", False)


@pytest.mark.skipif(platform.system() != "Windows", reason="COM automation only available on Windows")
def test_reverse_engineer_word_crossref(tmp_path):
    """Reverse engineer how Word creates cross-references by comparing before/after.

    This test:
    1. Creates a simple Word doc with a figure caption
    2. Scans the initial structure (XML)
    3. Uses Word COM to insert a cross-reference to the figure
    4. Updates all fields
    5. Scans the modified structure to see what changed
    """
    print("\n" + "="*80)
    print("REVERSE ENGINEERING WORD CROSS-REFERENCE SYSTEM")
    print("="*80)

    # Step 1: Use Word COM to create everything properly from scratch
    print("\n[STEP 3] Using Word COM to create proper caption and cross-reference...")

    import win32com.client

    # Use late-binding (DispatchEx) instead of early-binding to avoid strict timing issues
    word = win32com.client.DispatchEx("Word.Application")
    word.Visible = False

    try:
        # Create a new document from scratch
        word_doc = word.Documents.Add()

        # Add a heading
        word.Selection.Style = "Heading 1"
        word.Selection.TypeText("Test Document Created by Word")
        word.Selection.TypeParagraph()

        # Add a dummy image placeholder (just a shape for the caption to attach to)
        word.Selection.Style = "Normal"
        word.Selection.TypeParagraph()

        # Insert a caption using Word's built-in caption feature
        # This is the OFFICIAL way Word creates captions with bookmarks
        print("  Creating caption with Word's InsertCaption...")

        # Insert an actual shape for the caption to attach to
        print("  Inserting a shape for the caption to attach to...")
        # Use Shapes.AddShape (not InlineShapes) - msoShapeRectangle = 1
        shape = word_doc.Shapes.AddShape(1, 100, 100, 100, 100)

        # Move to end of document and add caption manually
        # InsertCaption can be unreliable across different Word configurations
        word.Selection.EndKey(Unit=6)  # wdStory = 6 (end of document)
        word.Selection.TypeParagraph()
        word.Selection.Style = "Caption"

        # Create a SEQ field for the figure number (what Word does internally)
        word.Selection.TypeText("Figure ")
        word.Selection.Fields.Add(
            Range=word.Selection.Range,
            Type=-1,  # wdFieldEmpty
            Text="SEQ Figure \\* ARABIC",
            PreserveFormatting=True
        )
        word.Selection.TypeText(": Test Caption Created by Word")

        # Create a bookmark for this caption (what Word does for cross-references)
        caption_range = word.Selection.Paragraphs(1).Range
        bookmark_name = "_Ref" + str(int(time.time() * 1000) % 1000000000)
        word_doc.Bookmarks.Add(bookmark_name, caption_range)

        print("  [OK] Caption created with bookmark:", bookmark_name)

        # Add some spacing
        word.Selection.TypeParagraph()
        word.Selection.TypeParagraph()

        # Now insert a cross-reference to the figure caption
        print("  Inserting cross-reference...")
        word.Selection.TypeText("See ")

        # Use InsertCrossReference method
        # We need to set the Selection first
        try:
            # Constants for InsertCrossReference
            # wdRefTypeNumberedItem = 0 (for any numbered item including figures)
            # wdNumberNoContext = 0, wdNumberRelativeContext = 2, wdNumberFullContext = 5
            # For figures specifically, we might need wdRefTypeFigure

            # Get the number of figure captions
            # The item index is 1-based
            word.Selection.InsertCrossReference(
                ReferenceType="Figure",  # Can use string instead of constant
                ReferenceKind=2,  # wdOnlyLabelAndNumber = 2
                ReferenceItem=1,  # First (and only) figure
                InsertAsHyperlink=True,
                IncludePosition=False
            )
            print("  [OK] Cross-reference inserted successfully")
        except Exception as e:
            print(f"  [X] InsertCrossReference failed: {e}")
            print(f"    Error type: {type(e)}")
            import traceback
            traceback.print_exc()

        # Add some more text
        word.Selection.TypeText(" for details.")
        word.Selection.TypeParagraph()


        # Step 4: Update all fields
        print("\n[STEP 4] Updating all fields...")
        try:
            word_doc.Fields.Update()
            print("  [OK] Fields updated")
        except Exception as e:
            print(f"  [!] Fields.Update failed (non-critical): {e}")

        # Wait for Word to finish processing before saving
        print("\n[STEP 5] Waiting for Word to finish processing...")
        time.sleep(2)  # Give Word time to complete all operations

        # Save the document created by Word
        modified_file = tmp_path / "word_created.docx"
        try:
            word_doc.SaveAs(str(modified_file.absolute()))
            print(f"  [OK] Word-created document saved: {modified_file}")
        except Exception as e:
            print(f"  [!] SaveAs failed: {e}")
            # Try alternative save method
            time.sleep(1)
            try:
                word_doc.SaveAs(str(modified_file.absolute()))
                print(f"  [OK] SaveAs succeeded on retry")
            except:
                print(f"  [X] SaveAs failed on retry - document may not be saved")

        try:
            word_doc.Close(SaveChanges=False)
        except Exception as e:
            print(f"  [!] Close failed (non-critical): {e}")

    finally:
        word.Quit()

    # Step 5: Scan the Word-created document structure
    print("\n[STEP 5] Scanning Word-created document structure...")
    word_structure = scan_document_structure(modified_file)
    print_structure("WORD-CREATED", word_structure)

    # Step 5b: Dump full XML for detailed analysis
    print("\n[STEP 5b] Dumping full XML structure for analysis...")
    dump_full_xml(modified_file, tmp_path)

    # Step 6: Analyze the structure
    print("\n[STEP 6] Analyzing Word's cross-reference structure...")
    analyze_word_structure(word_structure)

    print("\n" + "="*80)
    print("ANALYSIS COMPLETE")
    print("="*80)

    if auto_open:
        os.startfile(modified_file)

    # Basic assertions
    assert modified_file.exists(), "Word-created file should exist"
    # Note: Word may or may not create visible bookmarks depending on how cross-refs are implemented
    # We'll check the detailed structure instead


def scan_document_structure(docx_file: Path) -> dict:
    """Scan the internal XML structure of a .docx file.

    Returns a dict with:
    - bookmarks: list of bookmark names and IDs
    - fields: list of field codes
    - captions: list of caption paragraphs
    - references: list of reference paragraphs
    """
    structure = {
        'bookmarks': [],
        'fields': [],
        'captions': [],
        'references': [],
        'raw_xml_snippets': []
    }

    # Open the docx as a zip file
    with zipfile.ZipFile(docx_file, 'r') as zip_ref:
        # Read the main document.xml
        with zip_ref.open('word/document.xml') as xml_file:
            xml_content = xml_file.read().decode('utf-8')

    # Also use python-docx for higher-level analysis
    doc = Document(str(docx_file))

    # Find bookmarks in XML
    bookmark_starts = re.findall(
        r'<w:bookmarkStart w:id="(\d+)" w:name="([^"]+)"',
        xml_content
    )
    for bm_id, bm_name in bookmark_starts:
        structure['bookmarks'].append({
            'id': bm_id,
            'name': bm_name
        })

    # Find field codes (instrText)
    field_codes = re.findall(
        r'<w:instrText[^>]*>([^<]+)</w:instrText>',
        xml_content
    )
    for code in field_codes:
        structure['fields'].append(code.strip())

    # Find caption paragraphs
    for para in doc.paragraphs:
        if para.style.name == 'Caption' or 'Figure' in para.text:
            # Get XML for this paragraph
            xml_str = para._element.xml
            if isinstance(xml_str, bytes):
                xml_str = xml_str.decode('utf-8')

            structure['captions'].append({
                'text': para.text,
                'style': para.style.name,
                'xml_snippet': xml_str[:500]  # First 500 chars
            })

    # Find reference paragraphs (containing REF fields)
    for para in doc.paragraphs:
        xml_str = para._element.xml
        if isinstance(xml_str, bytes):
            xml_str = xml_str.decode('utf-8')

        if 'REF' in xml_str and 'STYLEREF' not in xml_str:
            structure['references'].append({
                'text': para.text,
                'xml_snippet': xml_str[:500]
            })

    return structure


def print_structure(label: str, structure: dict):
    """Pretty print a document structure."""
    print(f"\n  === {label} STRUCTURE ===")

    print(f"\n  Bookmarks ({len(structure['bookmarks'])}):")
    for bm in structure['bookmarks']:
        print(f"    - ID: {bm['id']}, Name: {bm['name']}")

    print(f"\n  Field Codes ({len(structure['fields'])}):")
    for field in structure['fields']:
        print(f"    - {field}")

    print(f"\n  Captions ({len(structure['captions'])}):")
    for cap in structure['captions']:
        print(f"    - Text: {cap['text']}")
        print(f"      Style: {cap['style']}")

    print(f"\n  References ({len(structure['references'])}):")
    for ref in structure['references']:
        print(f"    - Text: {ref['text']}")


def dump_full_xml(docx_file: Path, output_dir: Path):
    """Dump the full XML structure of paragraphs containing captions and references."""
    import zipfile

    doc = Document(str(docx_file))

    # Create output file for XML dumps
    xml_dump_file = output_dir / "xml_dump.txt"

    with open(xml_dump_file, 'w', encoding='utf-8') as f:
        f.write("="*80 + "\n")
        f.write("COMPLETE XML STRUCTURE DUMP\n")
        f.write("="*80 + "\n\n")

        # Dump all paragraphs
        for idx, para in enumerate(doc.paragraphs):
            if para.text.strip():  # Only dump non-empty paragraphs
                f.write(f"\n{'='*80}\n")
                f.write(f"PARAGRAPH {idx + 1}\n")
                f.write(f"Text: {para.text}\n")
                f.write(f"Style: {para.style.name}\n")
                f.write(f"{'='*80}\n")

                xml_str = para._element.xml
                if isinstance(xml_str, bytes):
                    xml_str = xml_str.decode('utf-8')

                # Pretty print the XML
                try:
                    import xml.dom.minidom
                    dom = xml.dom.minidom.parseString(xml_str)
                    pretty_xml = dom.toprettyxml(indent="  ")
                    f.write(pretty_xml)
                except:
                    # If pretty printing fails, just write raw XML
                    f.write(xml_str)

                f.write("\n")

    print(f"  [OK] Full XML dumped to: {xml_dump_file}")

    # Also extract the raw document.xml for reference
    with zipfile.ZipFile(docx_file, 'r') as zip_ref:
        with zip_ref.open('word/document.xml') as xml_file:
            doc_xml = xml_file.read().decode('utf-8')

    doc_xml_file = output_dir / "document.xml"
    with open(doc_xml_file, 'w', encoding='utf-8') as f:
        try:
            import xml.dom.minidom
            dom = xml.dom.minidom.parseString(doc_xml)
            f.write(dom.toprettyxml(indent="  "))
        except:
            f.write(doc_xml)

    print(f"  [OK] Raw document.xml saved to: {doc_xml_file}")


def analyze_word_structure(structure: dict):
    """Analyze Word's cross-reference implementation in detail."""
    print("\n  === WORD'S CROSS-REFERENCE IMPLEMENTATION ===")

    # Analyze bookmarks
    print(f"\n   Bookmarks ({len(structure['bookmarks'])}):")
    if structure['bookmarks']:
        for bm in structure['bookmarks']:
            print(f"    • {bm['name']} (ID: {bm['id']})")
            if bm['name'].startswith('_Ref'):
                print(f"      [OK] This is a Word-generated reference bookmark")
    else:
        print("    [!] No bookmarks found (might be hidden or not in expected format)")

    # Analyze field codes
    print(f"\n   Field Codes ({len(structure['fields'])}):")
    seq_fields = []
    ref_fields = []
    styleref_fields = []

    for field in structure['fields']:
        if 'SEQ' in field and 'STYLEREF' not in field:
            seq_fields.append(field)
            print(f"    • SEQ: {field}")
        elif 'REF' in field and 'STYLEREF' not in field:
            ref_fields.append(field)
            print(f"    • REF: {field}")
        elif 'STYLEREF' in field:
            styleref_fields.append(field)
            print(f"    • STYLEREF: {field}")

    print(f"\n    Summary: {len(seq_fields)} SEQ, {len(ref_fields)} REF, {len(styleref_fields)} STYLEREF")

    # Detailed analysis of captions
    if structure['captions']:
        print(f"\n   Caption Analysis (first caption):")
        cap = structure['captions'][0]
        print(f"    Text: {cap['text']}")
        print(f"    Style: {cap['style']}")

        cap_xml = cap['xml_snippet']

        # Extract SEQ field details
        seq_match = re.search(r'<w:instrText[^>]*>([^<]*SEQ[^<]*)</w:instrText>', cap_xml)
        if seq_match:
            print(f"    SEQ Field: {seq_match.group(1)}")

        # Check for bookmarks in caption
        bm_starts = re.findall(r'<w:bookmarkStart[^>]*w:id="(\d+)"[^>]*w:name="([^"]+)"', cap_xml)
        if bm_starts:
            print(f"    Bookmarks in caption:")
            for bm_id, bm_name in bm_starts:
                print(f"      • {bm_name} (ID: {bm_id})")

        # Check bookmark positioning relative to SEQ field
        if bm_starts and seq_match:
            print(f"\n     Bookmark Placement Analysis:")
            # Find position of bookmark vs SEQ field in XML
            seq_pos = cap_xml.find('<w:instrText')
            bm_pos = cap_xml.find('<w:bookmarkStart')
            if bm_pos < seq_pos:
                print(f"      Bookmark BEFORE SEQ field")
            elif bm_pos > seq_pos:
                print(f"      Bookmark AFTER SEQ field starts")

            # Check if bookmark wraps the SEQ field
            bm_end_pos = cap_xml.find('<w:bookmarkEnd')
            if bm_pos < seq_pos < bm_end_pos:
                print(f"      [OK] Bookmark WRAPS the SEQ field")

    # Detailed analysis of references
    if structure['references']:
        print(f"\n   Reference Analysis (first reference):")
        ref = structure['references'][0]
        print(f"    Text: {ref['text']}")

        ref_xml = ref['xml_snippet']

        # Extract REF field details
        ref_match = re.search(r'<w:instrText[^>]*>([^<]*REF[^<]*)</w:instrText>', ref_xml)
        if ref_match:
            ref_instruction = ref_match.group(1).strip()
            print(f"    REF Field: {ref_instruction}")

            # Parse REF field components
            parts = ref_instruction.split()
            if len(parts) >= 2:
                bookmark_ref = parts[1]
                switches = ' '.join(parts[2:]) if len(parts) > 2 else 'none'
                print(f"      Target Bookmark: {bookmark_ref}")
                print(f"      Switches: {switches}")

        # Check for hyperlink wrapper
        if 'w:hyperlink' in ref_xml:
            hyperlink_match = re.search(r'<w:hyperlink[^>]*w:anchor="([^"]+)"', ref_xml)
            if hyperlink_match:
                anchor = hyperlink_match.group(1)
                print(f"    Hyperlink Anchor: {anchor}")

                # Compare REF field bookmark to hyperlink anchor
                if ref_match and bookmark_ref == anchor:
                    print(f"      [OK] REF bookmark matches hyperlink anchor")

    # Final summary
    print(f"\n   KEY FINDINGS:")
    print(f"    • Word creates bookmarks: {len(structure['bookmarks']) > 0}")
    print(f"    • Bookmark naming pattern: {structure['bookmarks'][0]['name'] if structure['bookmarks'] else 'N/A'}")
    print(f"    • SEQ fields for numbering: {len(seq_fields)}")
    print(f"    • REF fields for cross-refs: {len(ref_fields)}")
    print(f"    • Uses hyperlinks: {'w:hyperlink' in (structure['references'][0]['xml_snippet'] if structure['references'] else '')}")



if __name__ == "__main__":
    import tempfile
    with tempfile.TemporaryDirectory() as tmp:
        test_reverse_engineer_word_crossref(Path(tmp))

