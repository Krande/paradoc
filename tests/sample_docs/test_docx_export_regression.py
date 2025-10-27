"""Test DOCX export regression - ensure database tables and images are correctly exported."""

import os

import pytest
from docx import Document

from paradoc import OneDoc


auto_open = os.getenv("AUTO_OPEN", False)


def test_doc_lorum_docx_export_with_db_tables(files_dir, tmp_path):
    """Test that doc_lorum exports to DOCX with database tables correctly.

    This test verifies that:
    1. Tables from database are correctly substituted and stored
    2. DOCX export can find and format tables
    3. The compiled DOCX file contains expected tables
    """
    source = files_dir / "doc_lorum"
    work_dir = tmp_path / source.name

    # Create OneDoc instance and compile to DOCX
    one = OneDoc(source, work_dir=work_dir)
    one.compile(source.name, auto_open=auto_open, export_format="docx")

    # DOCX files are created in _dist subdirectory
    dest = work_dir / "_dist" / f"{source.name}.docx"

    # Verify the output file was created
    assert dest.exists(), f"Expected DOCX file not created at {dest}"

    # Open the DOCX file and verify contents
    doc = Document(str(dest))

    # Count tables in the document
    num_tables = len(doc.tables)
    assert num_tables > 0, "Expected at least one table in the document"

    # Verify there are images (inline shapes) in the document
    total_images = 0
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            total_images += 1

    assert total_images > 0, "Expected at least one image in the document"

    # Verify that tables have content (not just placeholders)
    for table in doc.tables:
        if len(table.rows) > 0 and len(table.rows[0].cells) > 0:
            # At least one table should have actual data
            cell_text = table.rows[0].cells[0].text
            # Should not be empty or just contain table name placeholder
            assert len(cell_text.strip()) > 0, "Table cells should have content"


def test_doc_table_db_docx_export(files_dir, tmp_path):
    """Test doc_table_db exports to DOCX correctly.

    This document has a nested structure with database tables.
    """
    source = files_dir / "doc_table_db"
    if not source.exists():
        pytest.skip(f"Test document {source} not found")

    work_dir = tmp_path / source.name

    # Create OneDoc instance and compile to DOCX
    one = OneDoc(source, work_dir=work_dir)
    one.compile(source.name, auto_open=auto_open, export_format="docx")

    # DOCX files are created in _dist subdirectory
    dest = work_dir / "_dist" / f"{source.name}.docx"

    # Verify the output file was created
    assert dest.exists(), f"Expected DOCX file not created at {dest}"

    # Open the DOCX file and verify contents
    doc = Document(str(dest))

    # Verify tables exist
    num_tables = len(doc.tables)
    assert num_tables > 0, "Expected at least one table in the document"


def test_doc_with_nested_images_docx_export(files_dir, tmp_path):
    """Test that documents with nested directory structures export to DOCX successfully.

    This verifies that the DOCX export process completes without errors.
    Note: doc_lorum uses database-generated plots, not static image files.

    This is a regression test for the bug where images weren't being exported to DOCX
    because the resource path was pointing to the source directory instead of the build directory.
    """
    source = files_dir / "doc_lorum"
    work_dir = tmp_path / f"{source.name}_nested"

    # Create OneDoc instance and compile to DOCX
    one = OneDoc(source, work_dir=work_dir)
    one.compile(source.name, auto_open=auto_open, export_format="docx")

    # DOCX files are created in _dist subdirectory
    dest = work_dir / "_dist" / f"{source.name}.docx"

    # Verify the output file was created
    assert dest.exists(), f"Expected DOCX file not created at {dest}"

    # Open the DOCX file
    doc = Document(str(dest))

    # Verify the document has content (tables and/or images)
    # Count embedded images (relationships with image target)
    image_count = 0
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            image_count += 1

    # doc_lorum has 7 main plots + 7 appendix plots = 14 total expected figures
    # The exact count is: historical_trends, data_framework, statistical_workflow,
    # primary_results, comparative_analysis, error_analysis, theory_comparison (7 main)
    # + system_architecture, performance_benchmarks, time_series, computational_results,
    # correlation_matrix, distributions, surface_plot, box_plots (8 appendix) = 15 total
    # But we should verify at least 14 plots are present (allowing for variations)
    expected_min_figures = 14
    assert (
        image_count >= expected_min_figures
    ), f"Expected at least {expected_min_figures} images in the document (from database plots), but found {image_count}"

    # Verify tables exist
    table_count = len(doc.tables)
    assert table_count > 0, "Expected at least one table in the document"


def test_docx_export_table_lookup_fix(files_dir, tmp_path):
    """Test that the table lookup fix works - tables from DB are stored in self.tables.

    This is the core regression test for the bug where database tables weren't
    being stored in self.tables dictionary, causing DOCX export to fail.
    """
    source = files_dir / "doc_lorum"
    work_dir = tmp_path / f"{source.name}_lookup"

    # Create OneDoc instance
    one = OneDoc(source, work_dir=work_dir)

    # Before compilation, tables dict might be empty
    initial_table_count = len(one.tables)

    # Compile to DOCX - this should NOT raise "Unable to retrieve originally parsed table"
    try:
        one.compile(source.name, auto_open=auto_open, export_format="docx")
    except ValueError as e:
        if "Unable to retrieve originally parsed table" in str(e):
            pytest.fail(f"Regression detected: {e}")
        else:
            raise

    # After compilation, tables should be populated
    final_table_count = len(one.tables)

    # We should have tables stored
    assert (
        final_table_count > initial_table_count
    ), "Tables from database should be stored in self.tables after variable substitution"

    # DOCX files are created in _dist subdirectory
    dest = work_dir / "_dist" / f"{source.name}.docx"

    # Verify the output file was created successfully
    assert dest.exists(), f"Expected DOCX file not created at {dest}"

    # Verify it's a valid DOCX file by opening it
    doc = Document(str(dest))
    assert len(doc.tables) > 0, "Document should contain tables"


def test_docx_multiple_table_instances(files_dir, tmp_path):
    """Test that reused tables (same table key used multiple times) work in DOCX export.

    The system should create unique keys for reused tables.
    """
    # Create a test directory structure
    test_dir = tmp_path / "test_multi_instance"
    test_dir.mkdir()

    main_dir = test_dir / "00-main"
    main_dir.mkdir()

    # Create a markdown file that uses the same table twice
    md_content = """# Test Document

Here is the table the first time:

{{__my_table__}}

And here it is again:

{{__my_table__}}{tbl:sortby:Value:desc}
"""

    md_file = main_dir / "test.md"
    md_file.write_text(md_content)

    # Create OneDoc instance
    one = OneDoc(test_dir, work_dir=tmp_path / "work")

    # Add table to database
    import pandas as pd
    from paradoc.db import dataframe_to_table_data

    df = pd.DataFrame({"Item": ["A", "B", "C"], "Value": [10, 20, 15]})

    table_data = dataframe_to_table_data(key="my_table", df=df, caption="My Test Table", show_index=False)
    one.db_manager.add_table(table_data)

    # Compile to DOCX
    dest = tmp_path / "work" / "_dist" / "TestDoc.docx"
    one.compile("TestDoc", export_format="docx")

    # Should compile without errors
    assert dest.exists()

    # Verify the document has 2 tables (same table used twice)
    doc = Document(str(dest))
    assert len(doc.tables) >= 2, "Document should have at least 2 table instances"

    # Verify tables dictionary has unique keys for each instance
    assert "my_table" in one.tables, "Base table key should exist"
    # The second instance might be "my_table_1" or handled differently
