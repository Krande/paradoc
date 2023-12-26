import logging
import os

import pypandoc
from docx import Document

from paradoc import MY_DOCX_TMPL, OneDoc

# Miktex installer https://miktex.org/download

logging.basicConfig(level=logging.INFO)


auto_open = os.getenv("AUTO_OPEN", False)
table_format = "Grid Table 1 Light"
pg_style = "Normal Indent"


def test_report_markdown_to_docx(files_dir, test_dir):
    source = files_dir / "doc1"
    dest = test_dir / "report_md/md_to_docx_standard.docx"

    one = OneDoc(source, work_dir=dest.parent)
    one.compile("somefile_formatted", auto_open=auto_open)

    if auto_open is True:
        os.startfile(dest)


def test_compose_markdown_to_docx(files_dir, test_dir):
    source = files_dir / "doc2"
    dest = test_dir / "md_dir_to_docx/output.docx"

    one = OneDoc(source, work_dir=dest.parent)
    one.compile("somefile_formatted", auto_open=auto_open)


def test_report_crossref(files_dir, test_dir):
    source = files_dir / "doc3/demo.md"
    dest = test_dir / "report_md/report_crossref.docx"

    os.makedirs(dest.parent, exist_ok=True)
    output = pypandoc.convert_file(
        str(source),
        to="docx",
        outputfile=str(dest),
        filters=["pandoc-crossref"],
        extra_args=[f"--resource-path={source.parent}"],
        sandbox=False,
    )
    assert output == ""
    if auto_open is True:
        os.startfile(dest)


def test_basic_document_from_template(files_dir, test_dir):
    document = Document(MY_DOCX_TMPL)
    document.add_paragraph("My Main Report", style="Heading 1")
    document.add_paragraph("A subheading", style="Heading 2")
    document.add_paragraph("A subheading of a Subheading", style="Heading 3")
    document.add_page_break()
    document.add_paragraph("My Appendix", style="Appendix")
    document.add_paragraph("A subheading of my app", style="Appendix X.1")
    document.add_paragraph("A subheading of a subheading in my app", style="Appendix X.2")
    os.makedirs(test_dir, exist_ok=True)
    document.save(test_dir / "my_demo.docx")
