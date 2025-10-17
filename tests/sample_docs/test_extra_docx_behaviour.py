import logging
import os

import pypandoc
from docx import Document

from paradoc import MY_DOCX_TMPL

# Miktex installer https://miktex.org/download

logging.basicConfig(level=logging.INFO)

auto_open = os.getenv("AUTO_OPEN", False)
table_format = "Grid Table 1 Light"
pg_style = "Normal Indent"


def test_report_crossref(files_dir, tmp_path):
    source = files_dir / "doc3/demo.md"
    dest = tmp_path / "report_md/report_crossref.docx"

    os.makedirs(dest.parent, exist_ok=True)
    output = pypandoc.convert_file(
        str(source),
        to="docx",
        outputfile=str(dest),
        filters=["pandoc-crossref"],
        extra_args=[f"--resource-path={source.parent}"],
        sandbox=False,
    )
    assert output == ""
    if auto_open is True:
        os.startfile(dest)


def test_basic_document_from_template(files_dir, tmp_path):
    document = Document(MY_DOCX_TMPL)
    document.add_paragraph("My Main Report", style="Heading 1")
    document.add_paragraph("A subheading", style="Heading 2")
    document.add_paragraph("A subheading of a Subheading", style="Heading 3")
    document.add_page_break()
    document.add_paragraph("My Appendix", style="Appendix")
    document.add_paragraph("A subheading of my app", style="Appendix X.1")
    document.add_paragraph("A subheading of a subheading in my app", style="Appendix X.2")

    document.save(tmp_path / "my_demo.docx")
