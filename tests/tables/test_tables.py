import pandas as pd
from docx import Document

from paradoc import OneDoc
from paradoc.common import TableFormat


def test_table(files_dir, tmp_path):
    """Test that tables added via add_table() are properly exported to DOCX."""
    import re
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    report_dir = files_dir / "doc_table"
    one = OneDoc(report_dir, work_dir=tmp_path / "doc_table")
    df = pd.DataFrame([(0, 0), (1, 2)], columns=["a", "b"])

    one.add_table("my_table", df, "A basic table")
    one.add_table("my_table_2", df, "A slightly smaller table", TableFormat(font_size=8))
    one.add_table("my_table_3", df, "No Space 1")
    one.add_table("my_table_4", df, "No Space 2")
    one.add_table("my_table_5", df, "No Space 3")

    one.compile("TableDoc", export_format="docx", auto_open=False)

    # Verify the DOCX was created
    output_file = tmp_path / "doc_table" / "_dist" / "TableDoc.docx"
    assert output_file.exists(), f"Output DOCX file should exist at {output_file}"

    # Load and verify the document
    doc = Document(str(output_file))

    # Count tables in the document
    # Note: my_table and my_table_3 are referenced in the markdown
    # my_table_2, my_table_4, my_table_5 are defined but not used in the markdown
    tables = doc.tables
    assert len(tables) >= 2, f"Expected at least 2 tables in document, found {len(tables)}"

    # Verify table captions exist and have proper numbering
    caption_paras = []
    for para in doc.paragraphs:
        if para.style.name == "Table Caption":
            caption_paras.append(para)

    assert len(caption_paras) >= 2, f"Expected at least 2 table captions, found {len(caption_paras)}"

    # Verify caption numbering format (should be "Table 1-1", "Table 1-2", etc.)
    for i, caption_para in enumerate(caption_paras):
        caption_text = caption_para.text

        # Check that caption contains the table name
        if i == 0:
            assert "A basic table" in caption_text, f"First caption should contain 'A basic table', got: {caption_text}"
        elif i == 1:
            assert "No Space 1" in caption_text, f"Second caption should contain 'No Space 1', got: {caption_text}"

        # Verify numbering format: Should start with "Table" followed by X-Y format (e.g., "Table 1-1")
        # The pattern should be: "Table <chapter>-<number>: Caption text"
        numbering_pattern = r"Table\s+(\d+)-(\d+)"
        match = re.search(numbering_pattern, caption_text)
        assert match, f"Caption should have 'Table X-Y' numbering format, got: '{caption_text}'"

        chapter_num = int(match.group(1))
        table_num = int(match.group(2))
        assert chapter_num >= 1, f"Chapter number should be >= 1, got {chapter_num}"
        assert table_num >= 1, f"Table number should be >= 1, got {table_num}"

        # Verify caption has SEQ field (not just static text)
        # Check XML for field codes
        xml_str = caption_para._element.xml
        if isinstance(xml_str, bytes):
            xml_str = xml_str.decode("utf-8")

        has_seq_field = "SEQ" in xml_str or "fldChar" in xml_str
        assert has_seq_field, f"Caption should use SEQ fields for numbering, not static text. Caption: '{caption_text}'"

    # Verify table content and formatting
    for idx, table in enumerate(tables[:2]):  # Check first 2 tables
        assert len(table.rows) >= 2, f"Table {idx} should have at least 2 rows (header + data)"
        assert len(table.columns) == 2, f"Table {idx} should have 2 columns (a, b)"

        # Verify headers
        header_cells = table.rows[0].cells
        assert "a" in header_cells[0].text, f"Table {idx}: First column header should be 'a'"
        assert "b" in header_cells[1].text, f"Table {idx}: Second column header should be 'b'"

        # Verify data integrity - first cell should contain actual data, not table name
        # This verifies the bookmark-based identification is working
        first_data_cell = table.rows[1].cells[0].text.strip()
        assert first_data_cell == "0", f"Table {idx}: First data cell should be '0', not table name. Got: '{first_data_cell}'"

        # Verify table formatting/style
        # Tables should have a style applied (e.g., "Grid Table 1 Light")
        if table.style is not None:
            assert table.style.name is not None, f"Table {idx} should have a style applied"

    # Verify table alignment (should be centered)
    # Check the paragraph containing the table or the table's alignment property
    for i, block in enumerate(doc.element.body):
        if block.tag.endswith('}tbl'):  # This is a table element
            # Find the table properties
            tbl_pr = block.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblPr')
            if tbl_pr is not None:
                # Check for table justification (jc element)
                jc = tbl_pr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}jc')
                if jc is not None:
                    alignment = jc.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                    # Table should be centered
                    assert alignment == 'center', f"Table should be centered, got alignment: {alignment}"


def test_regular_table(files_dir, tmp_path):
    """Test that markdown tables are properly formatted and exported to DOCX."""
    import re

    report_dir = files_dir / "doc_regular_table"
    one = OneDoc(report_dir, work_dir=tmp_path / "doc_regular_table")

    one.compile("TableDoc", export_format="docx", auto_open=False)

    # Verify the DOCX was created
    output_file = tmp_path / "doc_regular_table" / "_dist" / "TableDoc.docx"
    assert output_file.exists(), f"Output DOCX file should exist at {output_file}"

    # Load and verify the document
    doc = Document(str(output_file))

    # Count tables - should have exactly 1 table from the markdown
    tables = doc.tables
    assert len(tables) == 1, f"Expected exactly 1 table in document, found {len(tables)}"

    table = tables[0]

    # Verify table structure
    assert len(table.rows) == 5, f"Table should have 5 rows (1 header + 4 data), found {len(table.rows)}"
    assert len(table.columns) == 4, f"Table should have 4 columns, found {len(table.columns)}"

    # Verify headers
    header_cells = table.rows[0].cells
    expected_headers = ["", "cat A [unit]", "cat 2 [unitB]", "num ex [-]"]
    for i, expected in enumerate(expected_headers):
        actual = header_cells[i].text.strip()
        assert expected in actual, f"Column {i} header should contain '{expected}', got '{actual}'"

    # Verify first row of data
    row1_cells = table.rows[1].cells
    assert "example1" in row1_cells[0].text, "First data row should contain 'example1'"
    assert "4000" in row1_cells[1].text, "First data row cat A should be 4000"
    assert "1.13" in row1_cells[2].text, "First data row cat 2 should be 1.13"
    assert "6" in row1_cells[3].text, "First data row num ex should be 6"

    # Verify caption exists and is properly formatted with SEQ fields
    caption_para = None
    for para in doc.paragraphs:
        if para.style.name == "Table Caption" and "A basic table" in para.text:
            caption_para = para
            break

    assert caption_para is not None, "Should find 'A basic table' caption paragraph"

    caption_text = caption_para.text

    # Verify caption numbering format: Should be "Table X-Y: A basic table"
    # where X is the heading level and Y is the table number within that section
    numbering_pattern = r"Table\s+(\d+)-(\d+)"
    match = re.search(numbering_pattern, caption_text)
    assert match, f"Caption should have 'Table X-Y' format, got: '{caption_text}'"

    chapter_num = int(match.group(1))
    table_num = int(match.group(2))
    assert chapter_num >= 1, f"Chapter number should be >= 1, got {chapter_num}"
    assert table_num >= 1, f"Table number should be >= 1, got {table_num}"

    # Verify that caption uses SEQ fields, not static text
    xml_str = caption_para._element.xml
    if isinstance(xml_str, bytes):
        xml_str = xml_str.decode("utf-8")

    # Should have field codes (SEQ and STYLEREF)
    has_seq_field = "SEQ" in xml_str or "fldChar" in xml_str
    assert has_seq_field, f"Caption should use SEQ fields for numbering, got static text: '{caption_text}'"

    # Verify the table reference/bookmark exists in the caption
    # The caption paragraph should have a bookmark for cross-referencing
    has_bookmark = "bookmarkStart" in xml_str or "w:bookmarkStart" in xml_str
    assert has_bookmark, "Table caption should have bookmark for cross-referencing"

    # Verify table has proper style applied
    if table.style is not None:
        assert table.style.name is not None, "Table should have a style applied"
        # Should be "Grid Table 1 Light" or similar grid style
        assert "Grid" in table.style.name or "Table" in table.style.name, \
            f"Table should use a grid/table style, got: {table.style.name}"

    # Verify table alignment (should be centered)
    for i, block in enumerate(doc.element.body):
        if block.tag.endswith('}tbl'):  # This is a table element
            tbl_pr = block.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblPr')
            if tbl_pr is not None:
                jc = tbl_pr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}jc')
                if jc is not None:
                    alignment = jc.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                    assert alignment == 'center', f"Table should be centered, got alignment: {alignment}"
                    break
