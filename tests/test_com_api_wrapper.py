"""Test the Word COM API wrapper."""

import os
import platform
import tempfile
from pathlib import Path

import pytest

from paradoc import MY_DOCX_TMPL


@pytest.mark.skipif(platform.system() != "Windows", reason="COM automation only available on Windows")
def test_word_com_api_wrapper(tmp_path):
    """Test the simplified Word COM API wrapper."""
    from paradoc.io.word.com_api import WordApplication
    
    output_file = Path(tmp_path) / "test_wrapper.docx"

    # Create a document using the wrapper
    with WordApplication(visible=False) as word_app:
        doc = word_app.create_document()

        # Add heading
        doc.add_heading("Test Document with COM API Wrapper", level=1)

        # Add some content
        doc.add_paragraph("This document was created using the simplified COM API wrapper.")
        doc.add_paragraph()

        # Add a figure with caption
        fig_bookmark = doc.add_figure_with_caption(
            caption_text="Example Figure",
            create_bookmark=True
        )

        doc.add_paragraph()

        # Add a table with caption
        table_bookmark = doc.add_table_with_caption(
            caption_text="Example Table",
            rows=3,
            cols=3,
            create_bookmark=True
        )

        doc.add_paragraph()

        # Add cross-references
        doc.add_cross_reference(
            bookmark_name="figure_0",
            reference_type="figure",
            prefix_text="See "
        )
        doc.add_paragraph(" for the figure.")

        doc.add_paragraph()
        doc.add_cross_reference(
            bookmark_name="table_0",
            reference_type="table",
            prefix_text="Refer to "
        )
        doc.add_paragraph(" for the table data.")

        # Update all fields
        doc.update_fields()

        # Save the document
        doc.save(output_file)

    # Verify the file was created
    assert output_file.exists(), "Document should be created"
    assert output_file.stat().st_size > 0, "Document should not be empty"

    print(f"\nTest document created successfully: {output_file}")

    # Optionally open the file for manual inspection
    if os.getenv("AUTO_OPEN"):
        os.startfile(output_file)


@pytest.mark.skipif(platform.system() != "Windows", reason="COM automation only available on Windows")
def test_word_sections_and_breaks():
    """Test section breaks and page breaks."""
    from paradoc.io.word.com_api import WordApplication
    
    with tempfile.TemporaryDirectory() as tmp_dir:
        output_file = Path(tmp_dir) / "test_sections.docx"
        
        with WordApplication(visible=False) as word_app:
            doc = word_app.create_document()
            
            doc.add_heading("Section 1", level=1)
            doc.add_paragraph("Content in section 1.")
            
            doc.add_section_break("next_page")
            
            doc.add_heading("Section 2", level=1)
            doc.add_paragraph("Content in section 2.")
            
            doc.add_page_break()
            
            doc.add_heading("Same Section, New Page", level=2)
            doc.add_paragraph("Still in section 2 but on a new page.")
            
            doc.save(output_file)
            
        assert output_file.exists()
        print(f"\nSections test document created: {output_file}")


@pytest.mark.skipif(platform.system() != "Windows", reason="COM automation only available on Windows")
def test_word_heading_levels():
    """Test different heading levels."""
    from paradoc.io.word.com_api import WordApplication
    
    with tempfile.TemporaryDirectory() as tmp_dir:
        output_file = Path(tmp_dir) / "test_headings.docx"
        
        with WordApplication(visible=False) as word_app:
            doc = word_app.create_document()
            
            for level in range(1, 4):
                doc.add_heading(f"Heading Level {level}", level=level)
                doc.add_paragraph(f"Content under heading level {level}.")
                doc.add_paragraph()
            
            doc.save(output_file)
            
        assert output_file.exists()
        print(f"\nHeadings test document created: {output_file}")


@pytest.mark.skipif(platform.system() != "Windows", reason="COM automation only available on Windows")
def test_word_with_template(tmp_path):
    """Test creating a document from a template."""
    from paradoc.io.word.com_api import WordApplication
    
    # First, create a simple template
    template_file = MY_DOCX_TMPL

    # Now create a new document based on the template
    output_file = Path(tmp_path) / "from_template.docx"
    
    with WordApplication(visible=False) as word_app:
        doc = word_app.create_document(template=template_file)

        doc.add_page_break()
        # Add content to the template-based document
        doc.add_heading("New Document", level=1)
        doc.add_heading("A Subsection", level=2)

        doc.add_paragraph("This document uses styles from the template.")
        
        doc.add_figure_with_caption("Figure in templated doc")
        doc.add_table_with_caption("Table in templated doc", rows=2, cols=2)

        doc.add_heading("Another Document", level=1)
        doc.add_heading("Another Subsection", level=2)

        doc.add_heading("Third Document", level=1)
        doc.add_heading("Third Subsection", level=2)

        doc.add_figure_with_caption("Second Figure in templated doc")
        doc.add_table_with_caption("Second Table in templated doc", rows=2, cols=2)

        doc.update_fields()
        doc.save(output_file)
    
    # Verify both files were created
    assert template_file.exists(), "Template should be created"
    assert output_file.exists(), "Document from template should be created"
    assert output_file.stat().st_size > 0, "Template-based document should not be empty"
    
    print(f"\nTemplate: {template_file}")
    print(f"Document from template: {output_file}")
    
    if os.getenv("AUTO_OPEN"):
        os.startfile(output_file)


@pytest.mark.skipif(platform.system() != "Windows", reason="COM automation only available on Windows")
def test_table_with_data(tmp_path):
    """Test creating a table with data."""
    from paradoc.io.word.com_api import WordApplication
    
    output_file = Path(tmp_path) / "test_table_data.docx"
    
    # Create test data
    table_data = [
        ["Header 1", "Header 2", "Header 3"],
        ["Row 1, Col 1", "Row 1, Col 2", "Row 1, Col 3"],
        ["Row 2, Col 1", "Row 2, Col 2", "Row 2, Col 3"],
    ]
    
    with WordApplication(visible=False) as word_app:
        doc = word_app.create_document()
        
        doc.add_heading("Table with Data Test", level=1)
        doc.add_paragraph("Testing table data population feature.")
        doc.add_paragraph()
        
        # Add table with data
        table_bookmark = doc.add_table_with_caption(
            caption_text="Test Data Table",
            rows=3,
            cols=3,
            data=table_data,
            create_bookmark=True
        )
        
        doc.add_paragraph()
        
        # Test with numeric data
        numeric_data = [
            ["Item", "Quantity", "Price"],
            ["Apple", 10, 1.50],
            ["Banana", 20, 0.75],
            ["Orange", 15, 2.00],
        ]
        
        doc.add_table_with_caption(
            caption_text="Numeric Data Table",
            rows=4,
            cols=3,
            data=numeric_data
        )
        
        doc.update_fields()
        doc.save(output_file)
    
    assert output_file.exists(), "Document should be created"
    assert output_file.stat().st_size > 0, "Document should not be empty"
    
    print(f"\nTable with data test document created: {output_file}")
    
    if os.getenv("AUTO_OPEN"):
        os.startfile(output_file)


@pytest.mark.skipif(platform.system() != "Windows", reason="COM automation only available on Windows")
def test_table_data_validation(tmp_path):
    """Test that table data validation works correctly."""
    from paradoc.io.word.com_api import WordApplication
    
    with WordApplication(visible=False) as word_app:
        doc = word_app.create_document()
        
        # Test: Wrong number of rows
        with pytest.raises(ValueError, match="Data has 2 rows but table has 3 rows"):
            doc.add_table_with_caption(
                caption_text="Invalid Table",
                rows=3,
                cols=2,
                data=[["A", "B"], ["C", "D"]]  # Only 2 rows instead of 3
            )
        
        # Test: Wrong number of columns
        with pytest.raises(ValueError, match="Data row 0 has 3 columns but table has 2 columns"):
            doc.add_table_with_caption(
                caption_text="Invalid Table",
                rows=2,
                cols=2,
                data=[["A", "B", "C"], ["D", "E", "F"]]  # 3 columns instead of 2
            )
    
    print("\nTable data validation tests passed")


@pytest.mark.skipif(platform.system() != "Windows", reason="COM automation only available on Windows")
def test_figure_layouts(tmp_path):
    """Test different figure layout options."""
    from paradoc.io.word.com_api import WordApplication, FigureLayout
    
    output_file = Path(tmp_path) / "test_figure_layouts.docx"
    
    with WordApplication(visible=False) as word_app:
        doc = word_app.create_document()
        
        doc.add_heading("Figure Layout Tests", level=1)
        doc.add_paragraph("Testing different figure layout options.")
        doc.add_paragraph()
        
        # Test inline (default)
        doc.add_paragraph("Figure with inline layout (default):")
        doc.add_figure_with_caption(
            caption_text="Inline Figure",
            layout=FigureLayout.INLINE
        )
        doc.add_paragraph()
        
        # Test square wrapping
        doc.add_paragraph("Figure with square wrapping:")
        doc.add_figure_with_caption(
            caption_text="Square Wrapped Figure",
            layout=FigureLayout.SQUARE
        )
        doc.add_paragraph()
        
        # Test tight wrapping
        doc.add_paragraph("Figure with tight wrapping:")
        doc.add_figure_with_caption(
            caption_text="Tight Wrapped Figure",
            layout=FigureLayout.TIGHT
        )
        doc.add_paragraph()
        
        # Test top/bottom wrapping
        doc.add_paragraph("Figure with top/bottom wrapping:")
        doc.add_figure_with_caption(
            caption_text="Top/Bottom Wrapped Figure",
            layout=FigureLayout.TOP_BOTTOM
        )
        doc.add_paragraph()
        
        # Test behind text
        doc.add_paragraph("Figure behind text:")
        doc.add_figure_with_caption(
            caption_text="Behind Text Figure",
            layout=FigureLayout.BEHIND_TEXT
        )
        doc.add_paragraph()
        
        # Test in front of text
        doc.add_paragraph("Figure in front of text:")
        doc.add_figure_with_caption(
            caption_text="In Front of Text Figure",
            layout=FigureLayout.IN_FRONT_OF_TEXT
        )
        doc.add_paragraph()
        
        # Test with string parameter
        doc.add_paragraph("Figure using string parameter:")
        doc.add_figure_with_caption(
            caption_text="String Parameter Figure",
            layout="square"  # Pass as string instead of enum
        )
        
        doc.update_fields()
        doc.save(output_file)
    
    assert output_file.exists(), "Document should be created"
    assert output_file.stat().st_size > 0, "Document should not be empty"
    
    print(f"\nFigure layouts test document created: {output_file}")
    
    if os.getenv("AUTO_OPEN"):
        os.startfile(output_file)


@pytest.mark.skipif(platform.system() != "Windows", reason="COM automation only available on Windows")
def test_chapter_numbering(tmp_path):
    """Test chapter-based numbering for figures and tables."""
    from paradoc.io.word.com_api import WordApplication
    from docx import Document
    
    output_file = Path(tmp_path) / "test_chapter_numbering.docx"
    
    with WordApplication(visible=False) as word_app:
        doc = word_app.create_document(template=MY_DOCX_TMPL)
        
        # Chapter 1
        doc.add_heading("Chapter 1: Introduction", level=1)
        doc.add_paragraph("This is the first chapter.")
        doc.add_paragraph()
        
        # Add figures with chapter numbering in Chapter 1
        doc.add_figure_with_caption(
            caption_text="First figure in chapter 1",
            use_chapter_numbers=True,
            create_bookmark=False  # Skip bookmarks to avoid cross-reference issues
        )
        doc.add_paragraph()
        
        doc.add_figure_with_caption(
            caption_text="Second figure in chapter 1",
            use_chapter_numbers=True,
            create_bookmark=False
        )
        doc.add_paragraph()
        
        # Add table with chapter numbering in Chapter 1
        doc.add_table_with_caption(
            caption_text="First table in chapter 1",
            rows=2,
            cols=2,
            use_chapter_numbers=True,
            create_bookmark=False
        )
        doc.add_paragraph()
        
        # Chapter 2
        doc.add_heading("Chapter 2: Methods", level=1)
        doc.add_paragraph("This is the second chapter.")
        doc.add_paragraph()
        
        # Add figures with chapter numbering in Chapter 2
        doc.add_figure_with_caption(
            caption_text="First figure in chapter 2",
            use_chapter_numbers=True,
            create_bookmark=False
        )
        doc.add_paragraph()
        
        doc.add_figure_with_caption(
            caption_text="Second figure in chapter 2",
            use_chapter_numbers=True,
            create_bookmark=False
        )
        doc.add_paragraph()
        
        # Add table with chapter numbering in Chapter 2
        doc.add_table_with_caption(
            caption_text="First table in chapter 2",
            rows=2,
            cols=2,
            use_chapter_numbers=True,
            create_bookmark=False
        )
        doc.add_paragraph()
        
        # Chapter 3
        doc.add_heading("Chapter 3: Results", level=1)
        doc.add_paragraph("This is the third chapter.")
        doc.add_paragraph()
        
        # Add mixed numbering - with and without chapter numbers
        doc.add_figure_with_caption(
            caption_text="Figure with chapter numbering",
            use_chapter_numbers=True,
            create_bookmark=False
        )
        doc.add_paragraph()
        
        doc.add_table_with_caption(
            caption_text="Table with simple numbering",
            rows=2,
            cols=2,
            use_chapter_numbers=False,  # This should continue simple numbering
            create_bookmark=False
        )
        doc.add_paragraph()
        
        doc.update_fields()
        doc.save(output_file)
    
    assert output_file.exists(), "Document should be created"
    assert output_file.stat().st_size > 0, "Document should not be empty"
    
    # Read the document using python-docx and verify figure/table numbering
    docx_doc = Document(str(output_file))
    
    # Find all caption paragraphs
    figure_captions = []
    table_captions = []
    
    for para in docx_doc.paragraphs:
        if para.style.name == 'Caption':
            text = para.text
            if text.startswith('Figure'):
                figure_captions.append(text)
            elif text.startswith('Table'):
                table_captions.append(text)
    
    # Assert figure numbering
    print(f"\nFound {len(figure_captions)} figure captions:")
    for caption in figure_captions:
        print(f"  {caption}")
    
    assert len(figure_captions) >= 5, f"Expected at least 5 figures, found {len(figure_captions)}"
    assert "Figure 1-1:" in figure_captions[0], f"Expected 'Figure 1-1:' in first caption, got: {figure_captions[0]}"
    assert "Figure 1-2:" in figure_captions[1], f"Expected 'Figure 1-2:' in second caption, got: {figure_captions[1]}"
    assert "Figure 2-1:" in figure_captions[2], f"Expected 'Figure 2-1:' in third caption, got: {figure_captions[2]}"
    assert "Figure 2-2:" in figure_captions[3], f"Expected 'Figure 2-2:' in fourth caption, got: {figure_captions[3]}"
    assert "Figure 3-1:" in figure_captions[4], f"Expected 'Figure 3-1:' in fifth caption, got: {figure_captions[4]}"
    
    # Assert table numbering
    print(f"\nFound {len(table_captions)} table captions:")
    for caption in table_captions:
        print(f"  {caption}")
    
    assert len(table_captions) >= 3, f"Expected at least 3 tables, found {len(table_captions)}"
    assert "Table 1-1:" in table_captions[0], f"Expected 'Table 1-1:' in first caption, got: {table_captions[0]}"
    assert "Table 2-1:" in table_captions[1], f"Expected 'Table 2-1:' in second caption, got: {table_captions[1]}"
    assert "Table 2:" in table_captions[2], f"Expected 'Table 2:' in third caption (simple numbering), got: {table_captions[2]}"
    
    print(f"\n[PASS] All figure and table numbering assertions passed!")
    print(f"  Chapter 1: Figure 1-1, Figure 1-2, Table 1-1")
    print(f"  Chapter 2: Figure 2-1, Figure 2-2, Table 2-1")
    print(f"  Chapter 3: Figure 3-1, Table 2 (simple)")
    
    if os.getenv("AUTO_OPEN"):
        os.startfile(output_file)


if __name__ == "__main__":
    # Run tests directly
    test_word_com_api_wrapper()
    test_word_sections_and_breaks()
    test_word_heading_levels()
    test_word_with_template()
    test_table_with_data()
    test_table_data_validation()
    test_figure_layouts()
    test_chapter_numbering()
    print("\nAll tests passed!")
